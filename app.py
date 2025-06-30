from flask import Flask, render_template, request, redirect, url_for, session, flash, jsonify, send_file, make_response
from subscription_system import check_user_limits, increment_usage
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
import psycopg2
from psycopg2.extras import RealDictCursor
import os
import openai
import PyPDF2
from docx import Document
from datetime import datetime
import json
import uuid
import tempfile
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

# Cargar variables de entorno desde .env
from dotenv import load_dotenv
load_dotenv()  # Carga las variables de entorno desde .env

# Configuración de OpenAI
# Configurar cliente OpenAI
openai.api_key = os.getenv('OPENAI_API_KEY')

# Configurar cliente OpenAI moderno
try:
    from openai import OpenAI
    OPENAI_CLIENT = OpenAI(api_key=os.getenv('OPENAI_API_KEY')) if os.getenv('OPENAI_API_KEY') else None
except ImportError:
    OPENAI_CLIENT = None

# Configuración de email
EMAIL_USER = os.getenv('EMAIL_USER')
EMAIL_PASSWORD = os.getenv('EMAIL_PASSWORD')
EMAIL_HOST = os.getenv('EMAIL_HOST', 'smtp.gmail.com')
EMAIL_PORT = int(os.getenv('EMAIL_PORT', 587))

# Validar configuración de email
if EMAIL_USER and EMAIL_PASSWORD:
    print(f"✅ Email configurado: {EMAIL_USER}")
    EMAIL_CONFIG_COMPLETE = True
else:
    print("❌ ADVERTENCIA: Configuración de email incompleta")
    print("   Configura las variables de entorno EMAIL_USER, EMAIL_PASSWORD, etc.")
    EMAIL_CONFIG_COMPLETE = False

# Intentar usar pdfkit como alternativa a WeasyPrint
try:
    import pdfkit
    PDFKIT_AVAILABLE = True
except ImportError:
    PDFKIT_AVAILABLE = False

# weasyprint deshabilitado en Windows por problemas de dependencias
WEASYPRINT_AVAILABLE = False

# Intentar usar reportlab como alternativa
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

from bs4 import BeautifulSoup  # Uncommented for HTML processing
import re
import tempfile  # Added for temporary file handling
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user

app = Flask(__name__)

# Configurar Flask
app.secret_key = os.getenv('SECRET_KEY', 'dev-key-change-in-production')
app.config['UPLOAD_FOLDER'] = os.getenv('UPLOAD_FOLDER', 'uploads')
app.config['MAX_CONTENT_LENGTH'] = int(os.getenv('MAX_CONTENT_LENGTH', 16 * 1024 * 1024))  # 16MB

# Configurar Flask-Login
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'
login_manager.login_message = 'Por favor inicia sesión para acceder a esta página.'
login_manager.login_message_category = 'info'

# Hacer datetime disponible en todas las plantillas
app.jinja_env.globals['datetime'] = datetime

# Importar y registrar blueprint de suscripción
from subscription_routes import subscription_bp
app.register_blueprint(subscription_bp, url_prefix='/subscription')

# Importar y registrar blueprint de sistema de ventas admin
from admin_sales_routes import register_admin_sales_routes
register_admin_sales_routes(app)

# Importar y registrar sistema de registro con suscripción
from registration_with_subscription import RegistrationWithSubscription
registration_system = RegistrationWithSubscription(app)
registration_system.register_routes()

# Verificar configuración de OpenAI
if openai.api_key:
    print("✅ OpenAI API configurada")
else:
    print("⚠️ OpenAI API Key no configurada")

# Configuración de base de datos
DB_CONFIG = {
    'host': os.getenv('DB_HOST', 'localhost'),
    'database': os.getenv('DB_NAME', 'armind_db'),
    'user': os.getenv('DB_USER', 'postgres'),
    'password': os.getenv('DB_PASSWORD', ''),
    'port': int(os.getenv('DB_PORT', 5432))
}

# Crear directorio de uploads si no existe
if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])

# Clase User para Flask-Login
class User(UserMixin):
    def __init__(self, user_id, email, name, role='user'):
        self.id = str(user_id)
        self.email = email
        self.name = name
        self.role = role
    
    def is_admin(self):
        return self.role == 'admin'
    
    def get_id(self):
        return self.id

@login_manager.user_loader
def load_user(user_id):
    """Cargar usuario para Flask-Login"""
    try:
        connection = get_db_connection()
        if not connection:
            return None
        
        cursor = connection.cursor()
        cursor.execute("SELECT id, email, username, role FROM users WHERE id = %s", (user_id,))
        user_data = cursor.fetchone()
        cursor.close()
        connection.close()
        
        if user_data:
            return User(user_data['id'], user_data['email'], user_data['username'], user_data.get('role', 'user'))
        return None
    except Exception as e:
        print(f"Error cargando usuario: {e}")
        return None

def generate_pdf_with_reportlab(html_content, title):
    """Generar PDF usando ReportLab desde contenido HTML"""
    if not REPORTLAB_AVAILABLE:
        return None
    
    try:
        from io import BytesIO
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.colors import black, blue
        
        # Crear buffer para el PDF
        buffer = BytesIO()
        
        # Crear documento
        doc = SimpleDocTemplate(buffer, pagesize=A4, 
                              rightMargin=72, leftMargin=72,
                              topMargin=72, bottomMargin=18)
        
        # Obtener estilos
        styles = getSampleStyleSheet()
        
        # Crear estilos personalizados
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1,  # Centrado
            textColor=black
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            textColor=blue
        )
        
        # Extraer texto del HTML usando BeautifulSoup
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Construir contenido del PDF
        story = []
        
        # Título principal
        title_text = soup.find('h1')
        if title_text:
            story.append(Paragraph(title_text.get_text().strip(), title_style))
        else:
            story.append(Paragraph(title, title_style))
        
        story.append(Spacer(1, 12))
        
        # Procesar el contenido
        print(f"DEBUG: HTML content: {html_content[:500]}...")  # Debug
        
        content_div = soup.find('div', class_='content')
        if content_div:
            # Buscar todos los párrafos dentro del div de contenido
            paragraphs = content_div.find_all('p')
            print(f"DEBUG: Found {len(paragraphs)} paragraphs")  # Debug
            
            if paragraphs:
                for p in paragraphs:
                    para_text = p.get_text().strip()
                    print(f"DEBUG: Paragraph text: {para_text[:100]}...")  # Debug
                    if para_text:
                        story.append(Paragraph(para_text, styles['Normal']))
                        story.append(Spacer(1, 12))
            else:
                # Si no hay párrafos, obtener todo el texto del div
                content_text = content_div.get_text().strip()
                print(f"DEBUG: Content text fallback: {content_text[:200]}...")  # Debug
                if content_text:
                    # Dividir por saltos de línea dobles para párrafos
                    paragraphs_text = content_text.split('\n\n')
                    for para in paragraphs_text:
                        para = para.strip().replace('\n', ' ')  # Reemplazar saltos simples con espacios
                        if para:
                            story.append(Paragraph(para, styles['Normal']))
                            story.append(Spacer(1, 12))
                else:
                    story.append(Paragraph("Error: No se pudo extraer el contenido de la carta.", styles['Normal']))
        else:
            # Fallback: procesar todos los elementos
            print("DEBUG: Using fallback processing")  # Debug
            for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'div']):
                text = element.get_text().strip()
                if text and text not in ['Carta de Presentación', 'Cover Letter', 'Carta de Apresentação', 'Anschreiben', 'Lettre de Motivation']:  # Evitar duplicar títulos
                    if element.name in ['h1', 'h2', 'h3']:
                        story.append(Paragraph(text, heading_style))
                    else:
                        story.append(Paragraph(text, styles['Normal']))
                    story.append(Spacer(1, 6))
        
        # Construir PDF
        doc.build(story)
        
        # Obtener contenido del buffer
        pdf_content = buffer.getvalue()
        buffer.close()
        
        return pdf_content
        
    except Exception as e:
        print(f"Error en generate_pdf_with_reportlab: {str(e)}")
        return None

def get_db_connection():
    """Obtener conexión a la base de datos PostgreSQL con manejo de errores mejorado"""
    try:
        # Usar opciones adicionales para manejar problemas de codificación
        os.environ['PGCLIENTENCODING'] = 'UTF8'
        
        connection = psycopg2.connect(
            host=DB_CONFIG['host'],
            database=DB_CONFIG['database'],
            user=DB_CONFIG['user'],
            password=DB_CONFIG['password'],
            port=DB_CONFIG['port'],
            cursor_factory=RealDictCursor,
            client_encoding='UTF8',
            options="-c client_encoding=UTF8"
        )
        
        return connection
    except psycopg2.Error as err:
        app.logger.error(f"Error de conexión a la base de datos: {err}")
        return None
    except Exception as err:
        app.logger.error(f"Error inesperado de base de datos: {err}")
        return None

def get_db():
    """Obtener conexión a la base de datos con cursor de diccionario"""
    return get_db_connection()

def generate_verification_token():
    """Generar token único para verificación de email"""
    return str(uuid.uuid4())

def send_verification_email(email, username, token):
    """Enviar email de verificación"""
    try:
        # Crear mensaje
        msg = MIMEMultipart('alternative')
        msg['Subject'] = 'Verifica tu cuenta de ARMind CVs'
        msg['From'] = EMAIL_USER
        msg['To'] = email
        
        # Versión texto plano
        text = f"""Hola {username},
        
Gracias por registrarte en ARMind CVs. Para verificar tu cuenta, haz clic en el siguiente enlace:
        
http://localhost:5000/verify_email/{token}
        
Si no solicitaste esta verificación, puedes ignorar este mensaje.
        
Saludos,
El equipo de ARMind CVs
        """
        
        # Versión HTML
        html = f"""
        <html>
        <head>
            <style>
                body {{ font-family: Arial, sans-serif; line-height: 1.6; color: #333; }}
                .container {{ max-width: 600px; margin: 0 auto; padding: 20px; }}
                .header {{ background-color: #4A90E2; color: white; padding: 20px; text-align: center; }}
                .content {{ padding: 20px; background-color: #f9f9f9; }}
                .button {{ display: inline-block; background-color: #4A90E2; color: white; padding: 10px 20px; text-decoration: none; border-radius: 5px; }}
                .footer {{ text-align: center; margin-top: 20px; font-size: 12px; color: #777; }}
            </style>
        </head>
        <body>
            <div class="container">
                <div class="header">
                    <h1>ARMind CVs</h1>
                </div>
                <div class="content">
                    <h2>Hola {username},</h2>
                    <p>Gracias por registrarte en ARMind CVs. Para verificar tu cuenta, haz clic en el siguiente botón:</p>
                    <p style="text-align: center;">
                        <a href="http://localhost:5000/verify_email/{token}" class="button">Verificar mi cuenta</a>
                    </p>
                    <p>Si el botón no funciona, copia y pega el siguiente enlace en tu navegador:</p>
                    <p>http://localhost:5000/verify_email/{token}</p>
                    <p>Si no solicitaste esta verificación, puedes ignorar este mensaje.</p>
                </div>
                <div class="footer">
                    <p>© 2023 ARMind CVs. Todos los derechos reservados.</p>
                </div>
            </div>
        </body>
        </html>
        """
        
        # Adjuntar partes al mensaje
        part1 = MIMEText(text, 'plain')
        part2 = MIMEText(html, 'html')
        msg.attach(part1)
        msg.attach(part2)
        
        # Conectar al servidor SMTP
        server = smtplib.SMTP(EMAIL_HOST, EMAIL_PORT)
        server.starttls()
        
        # Iniciar sesión
        server.login(EMAIL_USER, EMAIL_PASSWORD)
        
        # Enviar email
        server.sendmail(EMAIL_USER, email, msg.as_string())
        server.quit()
        
        print(f"✅ Email de verificación enviado exitosamente a: {email}")
        print(f"   Usuario: {username}")
        print(f"   Token: {token}")
        return True
    except Exception as e:
        print(f"Error al enviar email de verificación: {e}")
        print(f"Email destino: {email}")
        print(f"Username: {username}")
        print(f"Token: {token}")
        print(f"Configuración SMTP: {EMAIL_HOST}:{EMAIL_PORT}")
        import traceback
        traceback.print_exc()
        return False

def generate_reset_token():
    """Generar token único para reset de contraseña"""
    return str(uuid.uuid4())

def send_password_reset_email(email, username, token):
    """Enviar email de recuperación de contraseña"""
    try:
        # Crear mensaje
        msg = MIMEMultipart('alternative')
        msg['Subject'] = 'Recuperar contraseña - ARMind CVs'
        msg['From'] = EMAIL_USER
        msg['To'] = email
        
        # Versión texto plano
        text = f"""Hola {username},
        
Hemos recibido una solicitud para restablecer tu contraseña en ARMind CVs.
        
Para restablecer tu contraseña, haz clic en el siguiente enlace:
        
http://localhost:5000/reset_password/{token}
        
Este enlace expirará en 1 hora por seguridad.
        
Si no solicitaste este restablecimiento, puedes ignorar este mensaje.
        
Saludos,
El equipo de ARMind CVs
        """
        
        # Versión HTML
        html = f"""
        <html>
        <head>
            <style>
                body {{ font-family: Arial, sans-serif; line-height: 1.6; color: #333; }}
                .container {{ max-width: 600px; margin: 0 auto; padding: 20px; }}
                .header {{ background-color: #dc3545; color: white; padding: 20px; text-align: center; }}
                .content {{ padding: 20px; background-color: #f9f9f9; }}
                .button {{ display: inline-block; background-color: #dc3545; color: white; padding: 12px 24px; text-decoration: none; border-radius: 5px; font-weight: bold; }}
                .footer {{ text-align: center; margin-top: 20px; font-size: 12px; color: #777; }}
                .warning {{ background-color: #fff3cd; border: 1px solid #ffeaa7; padding: 10px; border-radius: 5px; margin: 15px 0; }}
            </style>
        </head>
        <body>
            <div class="container">
                <div class="header">
                    <h1>🔐 Recuperar Contraseña</h1>
                </div>
                <div class="content">
                    <h2>Hola {username},</h2>
                    <p>Hemos recibido una solicitud para restablecer tu contraseña en ARMind CVs.</p>
                    <p style="text-align: center;">
                        <a href="http://localhost:5000/reset_password/{token}" class="button">Restablecer mi contraseña</a>
                    </p>
                    <div class="warning">
                        <strong>⚠️ Importante:</strong> Este enlace expirará en 1 hora por seguridad.
                    </div>
                    <p>Si el botón no funciona, copia y pega el siguiente enlace en tu navegador:</p>
                    <p style="word-break: break-all;">http://localhost:5000/reset_password/{token}</p>
                    <p>Si no solicitaste este restablecimiento, puedes ignorar este mensaje de forma segura.</p>
                </div>
                <div class="footer">
                    <p>© 2023 ARMind CVs. Todos los derechos reservados.</p>
                </div>
            </div>
        </body>
        </html>
        """
        
        # Adjuntar partes al mensaje
        part1 = MIMEText(text, 'plain')
        part2 = MIMEText(html, 'html')
        msg.attach(part1)
        msg.attach(part2)
        
        # Conectar al servidor SMTP
        server = smtplib.SMTP(EMAIL_HOST, EMAIL_PORT)
        server.starttls()
        
        # Iniciar sesión
        server.login(EMAIL_USER, EMAIL_PASSWORD)
        
        # Enviar email
        server.sendmail(EMAIL_USER, email, msg.as_string())
        server.quit()
        
        print(f"✅ Email de recuperación enviado exitosamente a: {email}")
        print(f"   Usuario: {username}")
        print(f"   Token: {token}")
        return True
    except Exception as e:
        print(f"Error al enviar email de recuperación: {e}")
        print(f"Email destino: {email}")
        print(f"Username: {username}")
        print(f"Token: {token}")
        import traceback
        traceback.print_exc()
        return False

def validate_password_strength(password):
    """Validar que la contraseña cumpla con los requisitos de seguridad"""
    errors = []
    
    if len(password) < 8:
        errors.append("La contraseña debe tener al menos 8 caracteres")
    
    if not re.search(r'[a-z]', password):
        errors.append("La contraseña debe contener al menos una letra minúscula")
    
    if not re.search(r'[A-Z]', password):
        errors.append("La contraseña debe contener al menos una letra mayúscula")
    
    if not re.search(r'\d', password):
        errors.append("La contraseña debe contener al menos un número")
    
    if not re.search(r'[!@#$%^&*(),.?":{}|<>\-_+=\[\]\\;/~`]', password):
        errors.append("La contraseña debe contener al menos un carácter especial")
    
    return errors

def init_default_content(cursor):
    """Inicializar contenido por defecto del sitio"""
    # Contenido del dashboard principal
    default_content = [
        ('hero', 'title', 'Optimiza tu Currículum con Inteligencia Artificial'),
        ('hero', 'subtitle', 'ARMind CVs es una plataforma de análisis de currículums impulsada por IA para equipos de desarrollo ocupados.'),
        ('hero', 'cta_primary', 'COMENZAR'),
        ('hero', 'cta_secondary', 'DOCUMENTACIÓN'),
        ('features', 'section_title', '¿Por qué elegir ARMind CVs?'),
        ('features', 'section_subtitle', 'Herramientas profesionales para impulsar tu carrera'),
        ('features', 'ai_analysis_title', 'Análisis con IA'),
        ('features', 'ai_analysis_desc', 'Nuestro sistema de IA analiza tu CV como lo haría un ATS real, proporcionando puntuaciones y recomendaciones precisas.'),
        ('features', 'cv_builder_title', 'Constructor de CV'),
        ('features', 'cv_builder_desc', 'Crea currículums profesionales en formato Harvard con nuestro constructor guiado paso a paso.'),
        ('features', 'job_search_title', 'Búsqueda de Empleos'),
        ('features', 'job_search_desc', 'Encuentra oportunidades laborales relevantes en múltiples plataformas con nuestro motor de búsqueda inteligente.'),
        ('how_it_works', 'section_title', '¿Cómo funciona?'),
        ('how_it_works', 'section_subtitle', 'Tres simples pasos para optimizar tu currículum'),
        ('how_it_works', 'step1_title', 'Sube tu CV'),
        ('how_it_works', 'step1_desc', 'Carga tu currículum en formato PDF o Word. Nuestro sistema extraerá automáticamente toda la información.'),
        ('how_it_works', 'step2_title', 'Análisis IA'),
        ('how_it_works', 'step2_desc', 'Nuestra IA analiza tu CV como un sistema ATS profesional, evaluando cada sección y palabra clave.'),
        ('how_it_works', 'step3_title', 'Mejora y Aplica'),
        ('how_it_works', 'step3_desc', 'Recibe recomendaciones específicas, mejora tu CV y encuentra empleos que se ajusten a tu perfil.'),
        ('stats', 'section_title', 'Resultados que Hablan por Sí Solos'),
        ('cta_final', 'title', '¿Listo para Optimizar tu Carrera?'),
        ('cta_final', 'subtitle', 'Únete a miles de profesionales que ya están usando ARMind CVs para mejorar sus oportunidades laborales.')
    ]
    
    for section, key, value in default_content:
        cursor.execute("""
            INSERT INTO site_content (section, content_key, content_value) 
            VALUES (%s, %s, %s) 
            ON CONFLICT (section, content_key) DO NOTHING
        """, (section, key, value))
    
    # Consejos del día por defecto - solo insertar si no existen
    try:
        cursor.execute("SELECT COUNT(*) FROM daily_tips")
        tips_count = cursor.fetchone()[0]
        
        if tips_count == 0:
            default_tips = [
                ('Palabras Clave', 'Incluye palabras clave relevantes del sector en tu CV para mejorar tu puntuación ATS.', 'fas fa-key', 'primary'),
                ('Cuantifica Logros', 'Usa números y porcentajes para demostrar el impacto de tus logros profesionales.', 'fas fa-chart-bar', 'success')
            ]
            
            for title, description, icon, color in default_tips:
                cursor.execute("""
                    INSERT INTO daily_tips (title, description, icon, color) 
                    VALUES (%s, %s, %s, %s)
                """, (title, description, icon, color))
    except Exception as e:
        # Si hay error, continuar sin insertar consejos
        pass

def init_database():
    """Inicializar la base de datos y crear las tablas necesarias"""
    connection = get_db_connection()
    if not connection:
        app.logger.error("No se pudo conectar a la base de datos para inicializar")
        return False
        
    try:
        cursor = connection.cursor()
        
        # Crear tabla de usuarios
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS users (
                id SERIAL PRIMARY KEY,
                username VARCHAR(255) UNIQUE NOT NULL,
                first_name VARCHAR(255),
                last_name VARCHAR(255),
                phone VARCHAR(20),
                email VARCHAR(255) UNIQUE NOT NULL,
                password_hash VARCHAR(255) NOT NULL,
                email_verified BOOLEAN DEFAULT FALSE,
                verification_token VARCHAR(255),
                reset_token VARCHAR(255),
                reset_token_expires TIMESTAMP,
                role VARCHAR(50) DEFAULT 'user',
                is_banned BOOLEAN DEFAULT FALSE,
                ban_until TIMESTAMP NULL,
                ban_reason TEXT,
                last_login TIMESTAMP,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        
        # Agregar columnas nuevas si no existen
        # Agregar nuevas columnas individualmente con transacciones separadas
        new_columns = [
            ('first_name', 'VARCHAR(255)'),
            ('last_name', 'VARCHAR(255)'),
            ('phone', 'VARCHAR(20)')
        ]
        
        for column_name, column_def in new_columns:
            try:
                # Verificar si la columna ya existe
                cursor.execute("""
                    SELECT column_name 
                    FROM information_schema.columns 
                    WHERE table_name='users' AND column_name=%s
                """, (column_name,))
                
                if cursor.fetchone():
                    print(f"ℹ️  Columna {column_name} ya existe en la tabla users")
                else:
                    cursor.execute(f"ALTER TABLE users ADD COLUMN {column_name} {column_def}")
                    connection.commit()  # Commit individual para cada columna
                    print(f"✅ Columna {column_name} agregada a la tabla users")
            except Exception as e:
                connection.rollback()  # Rollback solo para esta operación
                print(f"⚠️  Error agregando columna {column_name}: {e}")
        
        # Crear tabla de curriculums
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS resumes (
                id SERIAL PRIMARY KEY,
                user_id INTEGER REFERENCES users(id) ON DELETE CASCADE,
                filename VARCHAR(255) NOT NULL,
                file_path VARCHAR(500) NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        
        # Crear tabla de empleos
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS jobs (
                id SERIAL PRIMARY KEY,
                title VARCHAR(500),
                company VARCHAR(255),
                location VARCHAR(255),
                description TEXT,
                url VARCHAR(1000),
                source VARCHAR(100),
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        
        # Crear tabla de feedback de IA
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS feedback (
                id SERIAL PRIMARY KEY,
                resume_id INTEGER REFERENCES resumes(id) ON DELETE CASCADE,
                score INTEGER,
                strengths TEXT,
                weaknesses TEXT,
                recommendations TEXT,
                keywords TEXT,
                s3_key VARCHAR(500),
                analysis_type VARCHAR(100),
                ai_provider VARCHAR(50),
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        
        # Crear tabla para guardar información del CV del usuario (múltiples CVs)
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS user_cvs (
                id SERIAL PRIMARY KEY,
                user_id INTEGER REFERENCES users(id) ON DELETE CASCADE,
                cv_name VARCHAR(255) NOT NULL DEFAULT 'Mi CV',
                personal_info TEXT,
                professional_summary TEXT,
                education TEXT,
                experience TEXT,
                skills TEXT,
                languages TEXT,
                certificates TEXT,
                format_options TEXT,
                is_active BOOLEAN DEFAULT TRUE,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                CONSTRAINT unique_user_cv_name UNIQUE(user_id, cv_name)
            )
        """)
        
        # Agregar columna certificates si no existe
        cursor.execute("""
            ALTER TABLE user_cvs 
            ADD COLUMN IF NOT EXISTS certificates TEXT
        """)
        
        # Agregar columna ai_methodologies si no existe
        cursor.execute("""
            ALTER TABLE user_cvs 
            ADD COLUMN IF NOT EXISTS ai_methodologies TEXT
        """)
        
        # Migrar datos de user_cv_data a user_cvs si existe
        cursor.execute("""
            INSERT INTO user_cvs (user_id, cv_name, personal_info, professional_summary, education, experience, skills, languages, format_options, updated_at)
            SELECT user_id, 'Mi CV Principal', personal_info, professional_summary, education, experience, skills, languages, format_options, updated_at
            FROM user_cv_data
            WHERE NOT EXISTS (
                SELECT 1 FROM user_cvs WHERE user_cvs.user_id = user_cv_data.user_id
            )
        """)
        
        # Crear tabla para contenido editable del sitio
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS site_content (
                id SERIAL PRIMARY KEY,
                section VARCHAR(100) NOT NULL,
                content_key VARCHAR(100) NOT NULL,
                content_value TEXT NOT NULL,
                content_type VARCHAR(50) DEFAULT 'text',
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_by INTEGER REFERENCES users(id),
                UNIQUE(section, content_key)
            )
        """)
        
        # Crear tabla para consejos del día
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS daily_tips (
                id SERIAL PRIMARY KEY,
                title VARCHAR(255) NOT NULL,
                description TEXT NOT NULL,
                icon VARCHAR(100) DEFAULT 'fas fa-lightbulb',
                color VARCHAR(50) DEFAULT 'primary',
                is_active BOOLEAN DEFAULT TRUE,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_by INTEGER REFERENCES users(id)
            )
        """)
        
        # Crear tabla para blog de tips y sugerencias
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS blog_posts (
                id SERIAL PRIMARY KEY,
                title VARCHAR(255) NOT NULL,
                content TEXT NOT NULL,
                image_url VARCHAR(500),
                author_id INTEGER REFERENCES users(id),
                is_published BOOLEAN DEFAULT TRUE,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        
        # Crear tabla para reacciones de usuarios a los posts del blog
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS blog_reactions (
                id SERIAL PRIMARY KEY,
                post_id INTEGER REFERENCES blog_posts(id) ON DELETE CASCADE,
                user_id INTEGER REFERENCES users(id),
                emoji VARCHAR(10) NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                UNIQUE(post_id, user_id, emoji)
            )
        """)
        
        # Crear tabla para almacenar imágenes
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS uploaded_images (
                id SERIAL PRIMARY KEY,
                filename VARCHAR(255) NOT NULL,
                original_filename VARCHAR(255) NOT NULL,
                content_type VARCHAR(100) NOT NULL,
                file_size INTEGER NOT NULL,
                image_data BYTEA NOT NULL,
                uploaded_by INTEGER REFERENCES users(id),
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        
        # Crear tabla para cartas de presentación
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS cover_letters (
                id SERIAL PRIMARY KEY,
                user_id INTEGER REFERENCES users(id) ON DELETE CASCADE,
                job_title VARCHAR(255) NOT NULL,
                company_name VARCHAR(255) NOT NULL,
                content TEXT NOT NULL,
                language VARCHAR(10) DEFAULT 'es',
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        
        # Crear tabla para análisis de CV
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS cv_analyses (
                id VARCHAR(255) PRIMARY KEY,
                user_id INTEGER REFERENCES users(id) ON DELETE CASCADE,
                filename VARCHAR(255) NOT NULL,
                analysis_type VARCHAR(100) NOT NULL,
                analysis_result TEXT NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        
        # Verificar si la columna professional_summary existe (PostgreSQL)
        cursor.execute("""
            SELECT column_name 
            FROM information_schema.columns 
            WHERE table_name='user_cv_data' AND column_name='professional_summary'
        """)
        if not cursor.fetchone():
            cursor.execute("ALTER TABLE user_cv_data ADD COLUMN professional_summary TEXT")
        
        # Verificar y agregar nuevas columnas de administración y reset de contraseña
        admin_columns = [
            ('role', 'VARCHAR(50) DEFAULT \'user\''),
            ('is_banned', 'BOOLEAN DEFAULT FALSE'),
            ('ban_until', 'TIMESTAMP NULL'),
            ('ban_reason', 'TEXT'),
            ('last_login', 'TIMESTAMP'),
            ('reset_token', 'VARCHAR(255)'),
            ('reset_token_expires', 'TIMESTAMP')
        ]
        
        for column_name, column_def in admin_columns:
            cursor.execute("""
                SELECT column_name 
                FROM information_schema.columns 
                WHERE table_name='users' AND column_name=%s
            """, (column_name,))
            if not cursor.fetchone():
                cursor.execute(f"ALTER TABLE users ADD COLUMN {column_name} {column_def}")
        
        # Insertar contenido por defecto si no existe
        init_default_content(cursor)
        
        connection.commit()
        cursor.close()
        connection.close()
        print("Base de datos inicializada correctamente")
        return True
        
    except psycopg2.Error as e:
        app.logger.error(f"Error de PostgreSQL al inicializar la base de datos: {e}")
        if connection:
            connection.rollback()
            connection.close()
        return False
    except Exception as e:
        app.logger.error(f"Error inesperado al inicializar la base de datos: {e}")
        if connection:
            connection.rollback()
            connection.close()
        return False

def save_image_to_database(image_file):
    """Guardar imagen en la base de datos y retornar URL"""
    try:
        import uuid
        import os
        
        app.logger.info(f"Iniciando guardado de imagen: {image_file.filename}")
        
        # Generar nombre único para la imagen
        file_extension = os.path.splitext(image_file.filename)[1].lower()
        unique_filename = f"{uuid.uuid4()}{file_extension}"
        
        app.logger.info(f"Nombre único generado: {unique_filename}")
        
        # Leer datos del archivo
        image_file.seek(0)
        image_data = image_file.read()
        
        app.logger.info(f"Datos leídos: {len(image_data)} bytes")
        
        connection = get_db_connection()
        if not connection:
            app.logger.error("No se pudo obtener conexión a la base de datos")
            return None
            
        try:
            cursor = connection.cursor()
            
            app.logger.info("Ejecutando INSERT en uploaded_images")
            
            cursor.execute("""
                INSERT INTO uploaded_images (filename, original_filename, content_type, file_size, image_data, uploaded_by)
                VALUES (%s, %s, %s, %s, %s, %s)
                RETURNING id
            """, (
                unique_filename,
                image_file.filename,
                image_file.content_type,
                len(image_data),
                image_data,
                session.get('user_id')
            ))
            
            result = cursor.fetchone()
            if result:
                # Manejar tanto tuplas como diccionarios
                if isinstance(result, dict):
                    image_id = result['id']
                else:
                    image_id = result[0]
                app.logger.info(f"Imagen guardada con ID: {image_id}")
            else:
                app.logger.error("No se obtuvo ID de la imagen insertada")
                return None
                
            connection.commit()
            cursor.close()
            connection.close()
            
            # Retornar URL para acceder a la imagen
            return f"/image/{image_id}"
            
        except Exception as db_error:
            error_msg = str(db_error)
            app.logger.error(f"Error en base de datos: {error_msg}")
            app.logger.error(f"Tipo de error: {type(db_error).__name__}")
            if connection:
                try:
                    connection.rollback()
                    connection.close()
                except Exception as close_error:
                    app.logger.error(f"Error cerrando conexión: {str(close_error)}")
            return None
            
    except Exception as general_error:
        error_msg = str(general_error)
        app.logger.error(f"Error general procesando imagen: {error_msg}")
        app.logger.error(f"Tipo de error general: {type(general_error).__name__}")
        return None

@app.route('/image/<int:image_id>')
def serve_image(image_id):
    """Servir imagen desde la base de datos"""
    try:
        connection = get_db_connection()
        if not connection:
            abort(404)
            
        cursor = connection.cursor()
        cursor.execute("""
            SELECT image_data, content_type, filename
            FROM uploaded_images
            WHERE id = %s
        """, (image_id,))
        
        result = cursor.fetchone()
        cursor.close()
        connection.close()
        
        if not result:
            abort(404)
            
        # Manejar tanto tuplas como diccionarios
        if isinstance(result, dict):
            image_data = result['image_data']
            content_type = result['content_type']
            filename = result['filename']
        else:
            image_data, content_type, filename = result
        
        from flask import Response
        return Response(
            image_data,
            mimetype=content_type,
            headers={
                'Content-Disposition': f'inline; filename="{filename}"',
                'Cache-Control': 'public, max-age=31536000'  # Cache por 1 año
            }
        )
        
    except Exception as e:
        app.logger.error(f"Error sirviendo imagen: {e}")
        abort(404)

@app.route('/')
def index():
    """Página principal"""
    return render_template('index.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    """Redirige al nuevo sistema de registro con suscripción obligatoria"""
    return redirect(url_for('register_with_subscription'))

@app.route('/register-legacy', methods=['GET', 'POST'])
def register_legacy():
    """Registro de usuarios (sistema anterior - solo para referencia)"""
    if request.method == 'POST':
        first_name = request.form.get('first_name', '').strip()
        last_name = request.form.get('last_name', '').strip()
        username = request.form.get('username', '').strip()
        country_code = request.form.get('country_code', '+56').strip()
        phone = request.form.get('phone', '').strip()
        email = request.form['email']
        password = request.form['password']
        
        # Combinar código de país con número de teléfono
        full_phone = f"{country_code}{phone}" if phone else ''
        
        add_console_log('INFO', f'Intento de registro para usuario: {username} ({email})', 'AUTH')
        
        # Validaciones básicas
        if not first_name or not last_name or not username or not email or not password or not phone:
            add_console_log('WARNING', f'Registro fallido - campos incompletos para: {username}', 'AUTH')
            flash('Todos los campos son obligatorios', 'error')
            return render_template('register.html')
        
        # Validar username
        if len(username) < 3:
            flash('El nombre de usuario debe tener al menos 3 caracteres', 'error')
            return render_template('register.html')
        
        if len(username) > 30:
            flash('El nombre de usuario no puede tener más de 30 caracteres', 'error')
            return render_template('register.html')
        
        if not re.match(r'^[a-zA-Z0-9._-]+$', username):
            flash('El nombre de usuario solo puede contener letras, números, puntos, guiones y guiones bajos', 'error')
            return render_template('register.html')
        
        # Verificar que el username no esté ya en uso
        try:
            connection = get_db_connection()
            if connection:
                cursor = connection.cursor()
                cursor.execute("SELECT id FROM users WHERE username = %s", (username,))
                if cursor.fetchone():
                    flash('Este nombre de usuario ya está en uso', 'error')
                    cursor.close()
                    connection.close()
                    return render_template('register.html')
                cursor.close()
                connection.close()
        except Exception as e:
            add_console_log('ERROR', f'Error verificando username: {str(e)}', 'AUTH')
            flash('Error verificando disponibilidad del nombre de usuario', 'error')
            return render_template('register.html')
        
        # Validar fortaleza de la contraseña
        password_errors = validate_password_strength(password)
        if password_errors:
            add_console_log('WARNING', f'Registro fallido - contraseña débil para: {username}', 'AUTH')
            for error in password_errors:
                flash(error, 'error')
            return render_template('register.html')
        
        # Hash de la contraseña
        password_hash = generate_password_hash(password)
        
        # Generar token de verificación
        verification_token = generate_verification_token()
        
        connection = get_db_connection()
        if connection:
            cursor = connection.cursor()
            try:
                cursor.execute(
                    "INSERT INTO users (username, first_name, last_name, phone, email, password_hash, email_verified, verification_token) VALUES (%s, %s, %s, %s, %s, %s, %s, %s) RETURNING id",
                    (username, first_name, last_name, full_phone, email, password_hash, False, verification_token)
                )
                user_id = cursor.fetchone()['id']
                connection.commit()
                
                add_console_log('INFO', f'Usuario registrado exitosamente: {username} (ID: {user_id})', 'AUTH')
                
                # Enviar correo de verificación
                print(f"🔄 Intentando enviar email de verificación...")
                email_sent = send_verification_email(email, username, verification_token)
                
                if email_sent:
                    add_console_log('INFO', f'Email de verificación enviado a: {email}', 'EMAIL')
                    print(f"✅ Email de verificación procesado correctamente")
                    flash('Usuario registrado exitosamente. Por favor, verifica tu correo electrónico para activar tu cuenta.', 'success')
                else:
                    add_console_log('ERROR', f'Error al enviar email de verificación a: {email}', 'EMAIL')
                    print(f"❌ Error al enviar email de verificación")
                    flash('Usuario registrado, pero hubo un problema al enviar el email de verificación. Contacta al administrador.', 'warning')
                return redirect(url_for('login'))
            except Exception as e:
                add_console_log('ERROR', f'Error al registrar usuario {username}: {str(e)}', 'AUTH')
                flash(f'Error al registrar usuario: {e}', 'error')
            finally:
                cursor.close()
                connection.close()
    
    return render_template('register.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    """Inicio de sesión"""
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        
        add_console_log('INFO', f'Intento de login para usuario: {username}', 'AUTH')
        
        connection = get_db_connection()
        if connection:
            cursor = connection.cursor()
            cursor.execute(
                "SELECT id, username, password_hash, email, email_verified, role, is_banned, ban_until, ban_reason FROM users WHERE username = %s",
                (username,)
            )
            user = cursor.fetchone()
            
            if user and check_password_hash(user['password_hash'], password):
                # Verificar si el correo está verificado
                if not user['email_verified']:
                    add_console_log('WARNING', f'Usuario no verificado intentó acceder: {username}', 'AUTH')
                    cursor.close()
                    connection.close()
                    flash('Por favor, verifica tu correo electrónico antes de iniciar sesión. Si no has recibido el correo de verificación, puedes solicitar uno nuevo.', 'warning')
                    return render_template('login.html', unverified_email=user['email'])
                
                # Verificar si el usuario está baneado
                if user['is_banned']:
                    ban_message = 'Tu cuenta ha sido suspendida.'
                    if user['ban_until']:
                        from datetime import datetime
                        if datetime.now() < user['ban_until']:
                            ban_message += f" Suspensión hasta: {user['ban_until'].strftime('%d/%m/%Y %H:%M')}"
                        else:
                            # El ban ha expirado, desbanearlo
                            cursor.execute(
                                "UPDATE users SET is_banned = FALSE, ban_until = NULL, ban_reason = NULL WHERE id = %s",
                                (user['id'],)
                            )
                            connection.commit()
                    else:
                        ban_message += ' Suspensión permanente.'
                    
                    if user['ban_reason']:
                        ban_message += f" Razón: {user['ban_reason']}"
                    
                    if user['is_banned'] and (not user['ban_until'] or datetime.now() < user['ban_until']):
                        add_console_log('WARNING', f'Usuario baneado intentó acceder: {username} - {ban_message}', 'AUTH')
                        cursor.close()
                        connection.close()
                        flash(ban_message, 'error')
                        return render_template('login.html')
                
                # Actualizar último login
                cursor.execute(
                    "UPDATE users SET last_login = CURRENT_TIMESTAMP WHERE id = %s",
                    (user['id'],)
                )
                connection.commit()
                
                # Crear objeto User y hacer login con Flask-Login
                user_obj = User(user['id'], user['email'], user.get('username', ''), user.get('role', 'user'))
                login_user(user_obj)
                
                # Mantener sesión para compatibilidad con código existente
                session['user_id'] = user['id']
                session['username'] = user['username']
                session['user_role'] = user.get('role', 'user')
                
                add_console_log('INFO', f'Login exitoso para {user.get("role", "user")}: {username}', 'AUTH')
                
                cursor.close()
                connection.close()
                
                flash('Inicio de sesión exitoso', 'success')
                
                # Redirigir según el rol
                if user.get('role') == 'admin':
                    return redirect(url_for('admin_dashboard'))
                else:
                    return redirect(url_for('dashboard'))
            else:
                add_console_log('WARNING', f'Login fallido para usuario: {username}', 'AUTH')
                cursor.close()
                connection.close()
                flash('Credenciales inválidas', 'error')
    
    return render_template('login.html')

@app.route('/logout')
def logout():
    """Cerrar sesión"""
    username = session.get('username', 'unknown')
    add_console_log('INFO', f'Usuario cerró sesión: {username}', 'AUTH')
    logout_user()  # Flask-Login logout
    session.clear()
    flash('Sesión cerrada', 'info')
    return redirect(url_for('index'))

@app.route('/forgot_password', methods=['POST'])
def forgot_password():
    """Solicitar recuperación de contraseña"""
    email = request.form.get('resetEmail')
    
    if not email:
        return jsonify({'success': False, 'message': 'Email es requerido'}), 400
    
    # Verificar si la configuración de email está completa
    if not EMAIL_CONFIG_COMPLETE:
        add_console_log('ERROR', 'Intento de recuperación de contraseña sin configuración de email', 'AUTH')
        return jsonify({
            'success': False, 
            'message': 'El servicio de recuperación de contraseña no está disponible. Contacta al administrador.'
        }), 503
    
    add_console_log('INFO', f'Solicitud de recuperación de contraseña para: {email}', 'AUTH')
    
    connection = get_db_connection()
    if connection:
        cursor = connection.cursor()
        try:
            # Verificar si el email existe
            cursor.execute("SELECT id, username, email FROM users WHERE email = %s", (email,))
            user = cursor.fetchone()
            
            if user:
                # Generar token de reset
                reset_token = generate_reset_token()
                
                # Establecer expiración del token (1 hora)
                from datetime import datetime, timedelta
                expires_at = datetime.now() + timedelta(hours=1)
                
                # Guardar token en la base de datos
                cursor.execute(
                    "UPDATE users SET reset_token = %s, reset_token_expires = %s WHERE id = %s",
                    (reset_token, expires_at, user['id'])
                )
                connection.commit()
                
                # Enviar email de recuperación
                email_sent = send_password_reset_email(user['email'], user['username'], reset_token)
                
                if email_sent:
                    add_console_log('INFO', f'Email de recuperación enviado a: {email}', 'EMAIL')
                    return jsonify({'success': True, 'message': 'Se ha enviado un enlace de recuperación a tu correo electrónico.'})
                else:
                    add_console_log('ERROR', f'Error al enviar email de recuperación a: {email}', 'EMAIL')
                    return jsonify({'success': False, 'message': 'Error al enviar el email. Inténtalo más tarde.'}), 500
            else:
                # Por seguridad, no revelamos si el email existe o no
                add_console_log('WARNING', f'Intento de recuperación con email inexistente: {email}', 'AUTH')
                return jsonify({'success': True, 'message': 'Si el email existe, se ha enviado un enlace de recuperación.'})
                
        except Exception as e:
            add_console_log('ERROR', f'Error en recuperación de contraseña para {email}: {str(e)}', 'AUTH')
            return jsonify({'success': False, 'message': 'Error interno del servidor'}), 500
        finally:
            cursor.close()
            connection.close()
    
    return jsonify({'success': False, 'message': 'Error de conexión a la base de datos'}), 500

@app.route('/reset_password/<token>', methods=['GET', 'POST'])
def reset_password(token):
    """Restablecer contraseña con token"""
    if request.method == 'GET':
        # Verificar si el token es válido
        connection = get_db_connection()
        if connection:
            cursor = connection.cursor()
            try:
                from datetime import datetime
                cursor.execute(
                    "SELECT id, username, email FROM users WHERE reset_token = %s AND reset_token_expires > %s",
                    (token, datetime.now())
                )
                user = cursor.fetchone()
                
                if user:
                    return render_template('reset_password.html', token=token, username=user['username'])
                else:
                    add_console_log('WARNING', f'Token de reset inválido o expirado: {token}', 'AUTH')
                    flash('El enlace de recuperación es inválido o ha expirado. Solicita uno nuevo.', 'error')
                    return redirect(url_for('login'))
                    
            except Exception as e:
                add_console_log('ERROR', f'Error al verificar token de reset: {str(e)}', 'AUTH')
                flash('Error al procesar la solicitud', 'error')
                return redirect(url_for('login'))
            finally:
                cursor.close()
                connection.close()
    
    elif request.method == 'POST':
        new_password = request.form.get('new_password')
        confirm_password = request.form.get('confirm_password')
        
        # Validaciones
        if not new_password or not confirm_password:
            flash('Todos los campos son obligatorios', 'error')
            return render_template('reset_password.html', token=token)
        
        if new_password != confirm_password:
            flash('Las contraseñas no coinciden', 'error')
            return render_template('reset_password.html', token=token)
        
        # Validar fortaleza de la contraseña
        password_errors = validate_password_strength(new_password)
        if password_errors:
            for error in password_errors:
                flash(error, 'error')
            return render_template('reset_password.html', token=token)
        
        # Actualizar contraseña
        connection = get_db_connection()
        if connection:
            cursor = connection.cursor()
            try:
                from datetime import datetime
                # Verificar token nuevamente
                cursor.execute(
                    "SELECT id, username, email FROM users WHERE reset_token = %s AND reset_token_expires > %s",
                    (token, datetime.now())
                )
                user = cursor.fetchone()
                
                if user:
                    # Actualizar contraseña y limpiar token
                    password_hash = generate_password_hash(new_password)
                    cursor.execute(
                        "UPDATE users SET password_hash = %s, reset_token = NULL, reset_token_expires = NULL WHERE id = %s",
                        (password_hash, user['id'])
                    )
                    connection.commit()
                    
                    add_console_log('INFO', f'Contraseña restablecida exitosamente para: {user["username"]}', 'AUTH')
                    flash('Tu contraseña ha sido restablecida exitosamente. Ya puedes iniciar sesión.', 'success')
                    return redirect(url_for('login'))
                else:
                    add_console_log('WARNING', f'Token de reset inválido o expirado en POST: {token}', 'AUTH')
                    flash('El enlace de recuperación es inválido o ha expirado.', 'error')
                    return redirect(url_for('login'))
                    
            except Exception as e:
                add_console_log('ERROR', f'Error al restablecer contraseña: {str(e)}', 'AUTH')
                flash('Error al restablecer la contraseña', 'error')
                return render_template('reset_password.html', token=token)
            finally:
                cursor.close()
                connection.close()
    
    flash('Error de conexión a la base de datos', 'error')
    return redirect(url_for('login'))

@app.route('/resend_verification', methods=['POST'])
def resend_verification():
    """Reenviar correo de verificación"""
    if request.method == 'POST':
        email = request.form.get('email')
        
        if not email:
            flash('Por favor, introduce tu correo electrónico', 'danger')
            return redirect(url_for('login'))
        
        # Verificar si el usuario existe y no está verificado
        conn = get_db_connection()
        if conn:
            try:
                cursor = conn.cursor()
                cursor.execute("SELECT * FROM users WHERE email = %s", (email,))
                user = cursor.fetchone()
                
                if not user:
                    flash('No existe una cuenta con ese correo electrónico', 'danger')
                    return redirect(url_for('login'))
                
                if user['email_verified']:
                    flash('Esta cuenta ya está verificada. Puedes iniciar sesión', 'info')
                    return redirect(url_for('login'))
                
                # Generar nuevo token
                token = generate_verification_token()
                
                # Actualizar token en la base de datos
                cursor.execute(
                    "UPDATE users SET verification_token = %s WHERE id = %s",
                    (token, user['id'])
                )
                conn.commit()
                
                # Enviar correo de verificación
                if send_verification_email(user['email'], user['username'], token):
                    flash('Se ha enviado un nuevo correo de verificación. Por favor, revisa tu bandeja de entrada', 'success')
                else:
                    flash('Error al enviar el correo de verificación. Por favor, inténtalo de nuevo más tarde', 'danger')
                
                return redirect(url_for('login'))
            
            except Exception as e:
                conn.rollback()
                flash(f'Error: {str(e)}', 'danger')
                return redirect(url_for('login'))
            
            finally:
                cursor.close()
                conn.close()
        
        flash('Error de conexión a la base de datos', 'danger')
        return redirect(url_for('login'))

@app.route('/verify_email/<token>')
def verify_email(token):
    """Verificar correo electrónico con token"""
    conn = get_db_connection()
    if conn:
        try:
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM users WHERE verification_token = %s", (token,))
            user = cursor.fetchone()
            
            if not user:
                flash('El enlace de verificación no es válido o ha expirado', 'danger')
                return redirect(url_for('login'))
            
            if user['email_verified']:
                flash('Esta cuenta ya está verificada. Puedes iniciar sesión', 'info')
                return redirect(url_for('login'))
            
            # Marcar la cuenta como verificada
            cursor.execute(
                "UPDATE users SET email_verified = TRUE, verification_token = NULL WHERE id = %s",
                (user['id'],)
            )
            conn.commit()
            
            flash('¡Tu cuenta ha sido verificada exitosamente! Ahora puedes iniciar sesión', 'success')
            return redirect(url_for('login'))
        
        except Exception as e:
            conn.rollback()
            flash(f'Error: {str(e)}', 'danger')
            return redirect(url_for('login'))
        
        finally:
            cursor.close()
            conn.close()
    
    flash('Error de conexión a la base de datos', 'danger')
    return redirect(url_for('login'))

@app.route('/dashboard')
def dashboard():
    """Panel principal del usuario"""
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    # Importar funciones de suscripción
    # Las importaciones se manejan en la sección de suscripción más abajo
    
    # Obtener estadísticas del usuario
    connection = get_db_connection()
    stats = {
        'cvs_analyzed': 0,
        'jobs_found': 0
    }
    daily_tips = []
    site_content = {}
    user_subscription = None
    user_usage = None
    
    if connection:
        cursor = connection.cursor()
        
        # Contar CVs analizados
        cursor.execute(
            "SELECT COUNT(*) FROM resumes WHERE user_id = %s",
            (session['user_id'],)
        )
        result = cursor.fetchone()
        stats['cvs_analyzed'] = result['count'] if result else 0
        
        # Contar empleos encontrados (esto es un placeholder, ajustar según la lógica de negocio)
        cursor.execute(
            "SELECT COUNT(*) FROM jobs"
        )
        result = cursor.fetchone()
        stats['jobs_found'] = result['count'] if result else 0
        
        # Obtener contenido del sitio
        cursor.execute(
            "SELECT section, content_key, content_value FROM site_content"
        )
        content_rows = cursor.fetchall()
        for row in content_rows:
            section = row['section']
            if section not in site_content:
                site_content[section] = {}
            site_content[section][row['content_key']] = row['content_value']
        
        cursor.close()
        connection.close()
    
    # Obtener consejos del día desde la tabla daily_tips
    tips_data = []
    try:
        connection = get_db_connection()
        if connection:
            cursor = connection.cursor()
            
            # Obtener todos los consejos del día desde la base de datos
            cursor.execute("""
                SELECT id, title, description, icon, color, is_active 
                FROM daily_tips 
                WHERE is_active = true
                ORDER BY id
            """)
            
            daily_tips_results = cursor.fetchall()
            
            # Convertir a formato compatible con el template
            for tip in daily_tips_results:
                tips_data.append({
                    'id': tip['id'],
                    'title': tip['title'],
                    'description': tip['description'],
                    'icon': tip['icon'],
                    'icon_color': tip['color']
                })
            
            cursor.close()
            connection.close()
    except Exception as e:
        print(f"Error cargando consejos del día desde BD: {e}")
        # Valores por defecto si hay error
        tips_data = [
            {
                'id': 1,
                'title': 'Formato Profesional',
                'description': 'Usa un formato limpio y profesional para tu CV.',
                'icon': 'fas fa-file-alt',
                'icon_color': 'text-info'
            },
            {
                'id': 2,
                'title': 'Palabras Clave',
                'description': 'Incluye palabras clave relevantes para tu industria.',
                'icon': 'fas fa-key',
                'icon_color': 'text-primary'
            }
        ]
    
    # Obtener información de suscripción del usuario
    from subscription_helpers import get_complete_user_usage
    from subscription_system import get_user_subscription, SUBSCRIPTION_PLANS
    
    user_subscription = get_user_subscription(session['user_id'])
    user_usage = get_complete_user_usage(session['user_id'])
    
    return render_template('dashboard.html', 
                         stats=stats, 
                         tips_data=tips_data, 
                         site_content=site_content,
                         user_subscription=user_subscription,
                         user_usage=user_usage,
                         subscription_plans=SUBSCRIPTION_PLANS)

def cleanup_cv_temp_file(cv_temp_id):
    """Función auxiliar para limpiar archivos temporales de CV"""
    if cv_temp_id:
        temp_cv_path = os.path.join(os.getcwd(), 'temp', f'cv_{cv_temp_id}.txt')
        try:
            if os.path.exists(temp_cv_path):
                os.remove(temp_cv_path)
                print(f"[DEBUG CV] ✅ Archivo temporal eliminado: {cv_temp_id}")
        except Exception as e:
            print(f"[DEBUG CV] ⚠️ Error eliminando archivo temporal: {str(e)}")

@app.route('/analyze_cv', methods=['GET', 'POST'])
def analyze_cv():
    """Analizador de CV con IA - Paso 1: Subir archivo (usando archivos temporales)"""
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    # Verificar restricciones de suscripción
    user_id = session.get('user_id')
    can_analyze, message = check_user_limits(user_id, 'cv_analysis')
    
    if not can_analyze:
        flash(f'Restricción de plan: {message}', 'error')
        return redirect(url_for('dashboard'))
    
    username = session.get('username', 'unknown')
    
    if request.method == 'POST':
        add_console_log('INFO', f'Usuario inició análisis de CV: {username}', 'CV')
        
        if 'file' not in request.files:
            add_console_log('WARNING', f'Análisis CV fallido - sin archivo: {username}', 'CV')
            flash('No se seleccionó ningún archivo', 'error')
            return redirect(request.url)
        
        file = request.files['file']
        if file.filename == '':
            add_console_log('WARNING', f'Análisis CV fallido - archivo vacío: {username}', 'CV')
            flash('No se seleccionó ningún archivo', 'error')
            return redirect(request.url)
        
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            
            # Usar archivo temporal en lugar de guardarlo permanentemente
            with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                filepath = temp_file.name
                file.save(filepath)
                
                add_console_log('INFO', f'Archivo CV procesado temporalmente: {filename} por {username}', 'CV')
                
                # Extraer texto del archivo (pasar la extensión del archivo original)
                original_extension = filename.rsplit('.', 1)[1].lower() if '.' in filename else None
                
                # LOGGING DETALLADO PARA DEBUG
                add_console_log('INFO', f'Iniciando extracción de texto - Archivo: {filename}, Extensión: {original_extension}, Usuario: {username}', 'CV')
                print(f"[DEBUG CV] Archivo: {filename}, Tamaño: {os.path.getsize(filepath)} bytes, Extensión: {original_extension}")
                
                text_content = extract_text_from_file(filepath, original_extension)
                
                # LOGGING DETALLADO DEL RESULTADO
                if text_content:
                    text_length = len(text_content)
                    text_preview = text_content[:200] if text_content else 'None'
                    add_console_log('INFO', f'Texto extraído exitosamente - {text_length} caracteres de {filename} por {username}', 'CV')
                    print(f"[DEBUG CV] ✅ Extracción exitosa: {text_length} caracteres")
                    print(f"[DEBUG CV] Vista previa: {repr(text_preview)}")
                    
                    # Validaciones adicionales para debug
                    if text_length < 50:
                        add_console_log('WARNING', f'Texto muy corto extraído ({text_length} chars) de {filename} por {username}', 'CV')
                        print(f"[DEBUG CV] ⚠️ ADVERTENCIA: Texto muy corto ({text_length} caracteres)")
                    
                    # Verificar palabras clave de CV
                    cv_keywords = ['experiencia', 'educación', 'habilidades', 'trabajo', 'universidad', 'empresa', 'proyecto', 'experience', 'education', 'skills', 'work', 'university', 'company', 'project']
                    found_keywords = sum(1 for keyword in cv_keywords if keyword.lower() in text_content.lower())
                    
                    if found_keywords < 2:
                        add_console_log('WARNING', f'Posible contenido no-CV detectado ({found_keywords} palabras clave) en {filename} por {username}', 'CV')
                        print(f"[DEBUG CV] ⚠️ ADVERTENCIA: Solo {found_keywords} palabras clave de CV encontradas")
                    else:
                        print(f"[DEBUG CV] ✅ Contenido parece ser CV ({found_keywords} palabras clave)")
                else:
                    add_console_log('ERROR', f'Extracción de texto falló completamente para {filename} por {username}', 'CV')
                    print(f"[DEBUG CV] ❌ ERROR: No se extrajo texto del archivo")
                
                print(f"[DEBUG CV] Texto extraído (primeros 200 caracteres): {text_content[:200] if text_content else 'None'}")
                
                # Eliminar el archivo temporal después de extraer el texto
                try:
                    os.unlink(filepath)
                    print(f"[DEBUG CV] ✅ Archivo temporal eliminado: {filepath}")
                except PermissionError:
                    # En Windows, a veces el archivo sigue en uso
                    # Registrar el error pero continuar con el proceso
                    add_console_log('WARNING', f'No se pudo eliminar el archivo temporal: {filepath}', 'CV')
                    print(f"[DEBUG CV] ⚠️ No se pudo eliminar archivo temporal: {filepath}")
                    pass
                
                if text_content:
                    # LOGGING DETALLADO DE SESIÓN
                    print(f"[DEBUG CV] Guardando en sesión - Usuario: {user_id}, Archivo: {filename}")
                    
                    # Generar ID único para el CV temporal
                    import uuid
                    cv_temp_id = str(uuid.uuid4())
                    
                    # Guardar CV temporalmente en el servidor
                    temp_dir = os.path.join(os.getcwd(), 'temp')
                    os.makedirs(temp_dir, exist_ok=True)
                    
                    temp_cv_path = os.path.join(temp_dir, f'cv_{cv_temp_id}.txt')
                    
                    try:
                        with open(temp_cv_path, 'w', encoding='utf-8') as f:
                            f.write(text_content)
                        
                        # Guardar solo el ID y filename en sesión
                        session['cv_temp_id'] = cv_temp_id
                        session['cv_filename'] = filename
                        
                        add_console_log('INFO', f'CV guardado temporalmente - {len(text_content)} chars para {username}', 'CV')
                        print(f"[DEBUG CV] ✅ CV guardado temporalmente")
                        print(f"[DEBUG CV] ID temporal: {cv_temp_id}")
                        print(f"[DEBUG CV] Contenido: {len(text_content)} caracteres")
                        print(f"[DEBUG CV] Archivo: {filename}")
                        
                    except Exception as e:
                        add_console_log('ERROR', f'Error guardando CV temporal para {username}: {str(e)}', 'CV')
                        print(f"[DEBUG CV] ❌ Error guardando temporalmente: {str(e)}")
                        flash('Error interno: No se pudo procesar el CV', 'error')
                        return redirect(request.url)
                    
                    # Incrementar contador de uso
                    increment_usage(session.get('user_id'), 'cv_analysis')
                    
                    # LOGGING ANTES DE REDIRECCIÓN
                    add_console_log('INFO', f'Redirigiendo a selección de IA - {filename} por {username}', 'CV')
                    print(f"[DEBUG CV] 🔄 Redirigiendo a select_ai_provider")
                    
                    # Redirigir a la selección de IA
                    return redirect(url_for('select_ai_provider'))
                else:
                    add_console_log('ERROR', f'Error extrayendo texto de: {filename} por {username}', 'CV')
                    print(f"[DEBUG CV] ❌ Error: No se pudo extraer texto del archivo")
                    flash('No se pudo extraer texto del archivo. Verifica que el archivo no esté dañado o encriptado.', 'error')
        else:
            add_console_log('WARNING', f'Archivo no permitido subido: {file.filename} por {username}', 'CV')
            flash('Tipo de archivo no permitido. Solo se permiten archivos PDF, DOC y DOCX.', 'error')
    
    # Definir consejos de análisis con valores por defecto
    tips_data = {
        'tip_format': {
            'title': 'Formato Claro',
            'description': 'Asegúrate de que tu CV tenga un formato limpio y sea fácil de leer.',
            'icon': 'fas fa-check-circle',
            'icon_color': 'text-success'
        },
        'tip_keywords': {
            'title': 'Palabras Clave',
            'description': 'Incluye palabras clave relevantes para tu industria y puesto objetivo.',
            'icon': 'fas fa-key',
            'icon_color': 'text-primary'
        },
        'tip_achievements': {
            'title': 'Logros Cuantificados',
            'description': 'Usa números y porcentajes para demostrar tus logros.',
            'icon': 'fas fa-chart-bar',
            'icon_color': 'text-info'
        },
        'tip_errors': {
            'title': 'Sin Errores',
            'description': 'Revisa la ortografía y gramática antes de subir tu CV.',
            'icon': 'fas fa-spell-check',
            'icon_color': 'text-warning'
        }
    }
    
    # Intentar cargar desde base de datos para sobrescribir valores por defecto
    try:
        connection = get_db_connection()
        if connection:
            cursor = connection.cursor()
            
            # Obtener todos los consejos de análisis desde la base de datos
            cursor.execute("""
                SELECT content_key, content_value 
                FROM site_content 
                WHERE section = 'analysis_tips'
            """)
            
            results = cursor.fetchall()
            analysis_tips = {}
            for row in results:
                analysis_tips[row['content_key']] = row['content_value']
            
            # Sobrescribir valores por defecto con los de la base de datos si existen
            if analysis_tips:
                for tip_type in ['tip_format', 'tip_keywords', 'tip_achievements', 'tip_errors']:
                    if f'{tip_type}_title' in analysis_tips:
                        tips_data[tip_type]['title'] = analysis_tips[f'{tip_type}_title']
                    if f'{tip_type}_description' in analysis_tips:
                        tips_data[tip_type]['description'] = analysis_tips[f'{tip_type}_description']
                    if f'{tip_type}_icon' in analysis_tips:
                        tips_data[tip_type]['icon'] = analysis_tips[f'{tip_type}_icon']
                    if f'{tip_type}_icon_color' in analysis_tips:
                        tips_data[tip_type]['icon_color'] = analysis_tips[f'{tip_type}_icon_color']
                
                print(f"Consejos cargados desde BD: {len(analysis_tips)} elementos")
            else:
                print("No hay consejos personalizados en BD, usando valores por defecto")
            
            cursor.close()
            connection.close()
    except Exception as e:
        print(f"Error cargando desde BD, usando valores por defecto: {e}")
    
    print("Consejos de análisis cargados correctamente")
    
    return render_template('analyze_cv.html', tips_data=tips_data)

@app.route('/select_ai_provider')
def select_ai_provider():
    """Paso 2: Seleccionar proveedor de IA"""
    # LOGGING DETALLADO PARA DEBUG
    user_id = session.get('user_id')
    username = session.get('username', 'unknown')
    
    print(f"[DEBUG AI_SELECT] Acceso a select_ai_provider - Usuario: {username} (ID: {user_id})")
    
    if 'user_id' not in session:
        print(f"[DEBUG AI_SELECT] ❌ Sin user_id en sesión")
        return redirect(url_for('login'))
    
    # VERIFICACIÓN DETALLADA DE SESIÓN
    print(f"[DEBUG AI_SELECT] Verificando contenido de sesión...")
    print(f"[DEBUG AI_SELECT] Claves en sesión: {list(session.keys())}")
    
    if 'cv_temp_id' not in session:
        print(f"[DEBUG AI_SELECT] ❌ ERROR CRÍTICO: cv_temp_id no está en sesión")
        print(f"[DEBUG AI_SELECT] Sesión actual: {dict(session)}")
        add_console_log('ERROR', f'FALLO CRÍTICO: cv_temp_id perdido en sesión para {username}', 'CV')
        flash('Primero debes subir un CV', 'error')
        return redirect(url_for('analyze_cv'))
    
    # VERIFICACIONES ADICIONALES
    cv_temp_id = session.get('cv_temp_id')
    cv_filename = session.get('cv_filename', 'NO_FILENAME')
    
    # Cargar CV desde archivo temporal
    temp_cv_path = os.path.join(os.getcwd(), 'temp', f'cv_{cv_temp_id}.txt')
    
    try:
        with open(temp_cv_path, 'r', encoding='utf-8') as f:
            cv_content = f.read()
    except FileNotFoundError:
        print(f"[DEBUG AI_SELECT] ❌ ERROR: Archivo temporal no encontrado")
        add_console_log('ERROR', f'Archivo temporal CV no encontrado para {username}', 'CV')
        cleanup_cv_temp_file(cv_temp_id)
        flash('El contenido del CV se perdió. Por favor, sube el archivo nuevamente.', 'error')
        return redirect(url_for('analyze_cv'))
    except Exception as e:
        print(f"[DEBUG AI_SELECT] ❌ ERROR: Error leyendo archivo temporal: {str(e)}")
        add_console_log('ERROR', f'Error leyendo CV temporal para {username}: {str(e)}', 'CV')
        cleanup_cv_temp_file(cv_temp_id)
        flash('Error interno al procesar el CV. Por favor, sube el archivo nuevamente.', 'error')
        return redirect(url_for('analyze_cv'))
    
    if not cv_content:
        print(f"[DEBUG AI_SELECT] ❌ ERROR: cv_content está vacío")
        add_console_log('ERROR', f'cv_content vacío en select_ai_provider para {username}', 'CV')
        cleanup_cv_temp_file(cv_temp_id)
        flash('El contenido del CV se perdió. Por favor, sube el archivo nuevamente.', 'error')
        return redirect(url_for('analyze_cv'))
    
    cv_length = len(cv_content)
    print(f"[DEBUG AI_SELECT] ✅ CV encontrado en sesión")
    print(f"[DEBUG AI_SELECT] Archivo: {cv_filename}")
    print(f"[DEBUG AI_SELECT] Contenido: {cv_length} caracteres")
    print(f"[DEBUG AI_SELECT] Vista previa: {repr(cv_content[:100])}")
    
    add_console_log('INFO', f'Acceso exitoso a select_ai_provider - {cv_filename} ({cv_length} chars) por {username}', 'CV')
    
    return render_template('select_ai_provider.html')

@app.route('/select_analysis_type/<ai_provider>')
def select_analysis_type(ai_provider):
    """Paso 3: Seleccionar tipo de análisis"""
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    if 'cv_temp_id' not in session:
        flash('Primero debes subir un CV', 'error')
        return redirect(url_for('analyze_cv'))
    
    # Validar proveedor de IA
    valid_providers = ['openai', 'anthropic', 'gemini']
    if ai_provider not in valid_providers:
        flash('Proveedor de IA no válido', 'error')
        return redirect(url_for('select_ai_provider'))
    
    session['selected_ai'] = ai_provider
    
    return render_template('select_analysis_type.html', ai_provider=ai_provider)

@app.route('/perform_analysis/<analysis_type>')
def perform_analysis(analysis_type):
    """Paso 4: Realizar análisis específico"""
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    if 'cv_temp_id' not in session or 'selected_ai' not in session:
        flash('Sesión expirada. Reinicia el proceso', 'error')
        return redirect(url_for('analyze_cv'))
    
    username = session.get('username', 'unknown')
    cv_temp_id = session.get('cv_temp_id')
    filename = session['cv_filename']
    ai_provider = session['selected_ai']
    
    # Cargar CV desde archivo temporal
    temp_cv_path = os.path.join(os.getcwd(), 'temp', f'cv_{cv_temp_id}.txt')
    
    try:
        with open(temp_cv_path, 'r', encoding='utf-8') as f:
            cv_content = f.read()
    except FileNotFoundError:
        cleanup_cv_temp_file(cv_temp_id)
        flash('El contenido del CV se perdió. Por favor, sube el archivo nuevamente.', 'error')
        return redirect(url_for('analyze_cv'))
    except Exception as e:
        cleanup_cv_temp_file(cv_temp_id)
        flash('Error interno al procesar el CV. Por favor, sube el archivo nuevamente.', 'error')
        return redirect(url_for('analyze_cv'))
    
    # Validar tipo de análisis
    valid_analysis_types = [
        'general_health_check',
        'content_quality_analysis', 
        'job_tailoring_optimization',
        'ats_compatibility_verification',
        'tone_style_evaluation',
        'industry_role_feedback',
        'benchmarking_comparison',
        'ai_improvement_suggestions',
        'visual_design_assessment',
        'comprehensive_score'
    ]
    
    if analysis_type not in valid_analysis_types:
        flash('Tipo de análisis no válido', 'error')
        return redirect(url_for('select_analysis_type', ai_provider=ai_provider))
    
    try:
        add_console_log('INFO', f'Iniciando análisis {analysis_type} con {ai_provider} para: {filename}', 'CV')
        
        # Realizar análisis según el proveedor y tipo seleccionado
        analysis = perform_cv_analysis(cv_content, ai_provider, analysis_type)
        
        # Guardar en la base de datos
        save_cv_analysis(session['user_id'], filename, cv_content, analysis)
        add_console_log('INFO', f'Análisis CV completado exitosamente: {filename} por {username}', 'CV')
        
        # Limpiar sesión y archivo temporal
        cv_temp_id = session.pop('cv_temp_id', None)
        session.pop('cv_filename', None)
        session.pop('selected_ai', None)
        
        # Eliminar archivo temporal
        cleanup_cv_temp_file(cv_temp_id)
        
        return render_template('cv_analysis_result.html', analysis=analysis, analysis_type=analysis_type, ai_provider=ai_provider)
        
    except Exception as e:
        add_console_log('ERROR', f'Error en análisis: {str(e)}', 'CV')
        flash(f'Error durante el análisis: {str(e)}', 'error')
        return redirect(url_for('select_analysis_type', ai_provider=ai_provider))

def allowed_file(filename):
    """Verificar si el archivo tiene una extensión permitida"""
    ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_file(filepath, file_extension=None):
    """Extraer texto de archivos PDF o Word con manejo robusto de errores"""
    
    # Validaciones iniciales
    if not filepath or not os.path.exists(filepath):
        print(f"Error: Archivo no existe: {filepath}")
        return None
    
    if not os.access(filepath, os.R_OK):
        print(f"Error: No se puede leer el archivo: {filepath}")
        return None
    
    file_size = os.path.getsize(filepath)
    if file_size == 0:
        print(f"Error: El archivo está vacío: {filepath}")
        return None
    
    if file_size > 50 * 1024 * 1024:  # 50MB límite
        print(f"Error: Archivo demasiado grande: {file_size} bytes")
        return None
    
    print(f"Procesando archivo: {os.path.basename(filepath)} ({file_size} bytes)")
    
    # Determinar extensión del archivo
    if file_extension is None:
        if '.' not in filepath:
            # Detectar tipo por header
            try:
                with open(filepath, 'rb') as f:
                    header = f.read(16)
                    if header.startswith(b'%PDF'):
                        file_extension = 'pdf'
                        print("Detectado como PDF por header")
                    elif header.startswith(b'PK\x03\x04'):
                        file_extension = 'docx'
                        print("Detectado como DOCX por header")
                    elif header.startswith(b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'):
                        file_extension = 'doc'
                        print("Detectado como DOC por header")
                    else:
                        print(f"Error: Tipo de archivo no reconocido: {header[:8]}")
                        return None
            except Exception as e:
                print(f"Error leyendo header del archivo: {e}")
                return None
        else:
            file_extension = filepath.rsplit('.', 1)[1].lower()
    
    text = ""
    
    try:
        if file_extension == 'pdf':
            text = _extract_from_pdf_robust(filepath)
        elif file_extension in ['doc', 'docx']:
            text = _extract_from_word_robust(filepath, file_extension)
        else:
            print(f"Error: Extensión no soportada: {file_extension}")
            return None
    
    except Exception as e:
        print(f"Error general extrayendo texto: {e}")
        return None
    
    if not text or len(text.strip()) == 0:
        print("Advertencia: No se extrajo texto del archivo")
        return None
    
    # Validar calidad del texto
    text = text.strip()
    if len(text) < 10:
        print(f"Advertencia: Texto extraído muy corto: {len(text)} caracteres")
    
    # Verificar caracteres imprimibles
    printable_ratio = sum(1 for c in text if c.isprintable() or c.isspace()) / len(text)
    if printable_ratio < 0.8:
        print(f"Advertencia: Texto contiene muchos caracteres no imprimibles: {printable_ratio:.2%}")
    
    print(f"Texto extraído exitosamente: {len(text)} caracteres")
    return text

def _extract_from_pdf_robust(filepath):
    """Extraer texto de archivo PDF con manejo robusto de errores"""
    text = ""
    
    try:
        with open(filepath, 'rb') as file:
            # Verificar que es un PDF válido
            file.seek(0)
            header = file.read(4)
            if not header.startswith(b'%PDF'):
                raise ValueError("No es un archivo PDF válido")
            
            file.seek(0)
            pdf_reader = PyPDF2.PdfReader(file)
            
            # Verificar si está encriptado
            if pdf_reader.is_encrypted:
                print("PDF está encriptado, intentando desencriptar")
                try:
                    pdf_reader.decrypt('')  # Intentar con contraseña vacía
                    print("PDF desencriptado exitosamente")
                except:
                    print("Error: No se puede desencriptar el PDF")
                    return None
            
            num_pages = len(pdf_reader.pages)
            if num_pages == 0:
                print("Error: PDF no tiene páginas")
                return None
            
            print(f"PDF tiene {num_pages} páginas")
            
            # Extraer texto de cada página
            pages_with_text = 0
            for i, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text += page_text + "\n"
                        pages_with_text += 1
                    else:
                        print(f"Advertencia: Página {i+1} no tiene texto extraíble")
                except Exception as page_error:
                    print(f"Advertencia: Error en página {i+1}: {page_error}")
                    continue
            
            print(f"Texto extraído de {pages_with_text}/{num_pages} páginas")
    
    except Exception as e:
        print(f"Error procesando PDF: {e}")
        raise
    
    return text

def _extract_from_word_robust(filepath, file_extension):
    """Extraer texto de archivo Word con manejo robusto de errores"""
    text = ""
    
    try:
        # Verificar header del archivo
        with open(filepath, 'rb') as f:
            header = f.read(8)
            if file_extension == 'docx' and not header.startswith(b'PK'):
                print("Advertencia: El archivo no parece ser un DOCX válido")
            elif file_extension == 'doc' and not header.startswith(b'\xd0\xcf'):
                print("Advertencia: El archivo no parece ser un DOC válido")
        
        # Abrir documento
        doc = Document(filepath)
        
        # Extraer texto de párrafos
        paragraph_count = 0
        for paragraph in doc.paragraphs:
            para_text = paragraph.text.strip()
            if para_text:
                text += para_text + '\n'
                paragraph_count += 1
        
        print(f"Extraídos {paragraph_count} párrafos")
        
        # Extraer texto de tablas
        table_count = 0
        if hasattr(doc, 'tables') and doc.tables:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            text += cell_text + ' '
                text += '\n'
                table_count += 1
            
            print(f"Extraídas {table_count} tablas")
    
    except Exception as e:
        print(f"Error procesando documento Word: {e}")
        raise
    
    return text

def perform_cv_analysis(cv_text, ai_provider, analysis_type):
    """Realizar análisis de CV según el proveedor de IA y tipo de análisis seleccionado"""
    if ai_provider == 'openai':
        return analyze_cv_with_openai(cv_text, analysis_type)
    elif ai_provider == 'anthropic':
        return analyze_cv_with_anthropic(cv_text, analysis_type)
    elif ai_provider == 'gemini':
        return analyze_cv_with_gemini(cv_text, analysis_type)
    else:
        raise ValueError(f"Proveedor de IA no soportado: {ai_provider}")

def get_analysis_prompt(analysis_type, cv_text):
    """Obtener el prompt específico según el tipo de análisis"""
    
    analysis_prompts = {
        'general_health_check': f"""
        Realiza una revisión general exhaustiva del estado del currículum como un experto en recursos humanos con 15 años de experiencia. Evalúa minuciosamente la ortografía, gramática, formato, integridad de secciones y longitud del currículum.
        
        Proporciona un análisis EXTREMADAMENTE DETALLADO que incluya:
        1. Puntaje general (0-100) con justificación específica
        2. Lista detallada de errores de ortografía y gramática encontrados con ejemplos exactos
        3. Problemas de formato específicos con ubicaciones precisas
        4. Recomendaciones de mejora con ejemplos concretos de cómo implementarlas
        5. Análisis de estructura y organización
        6. Evaluación de la longitud y densidad de información
        7. Ejemplos específicos de mejoras con texto "antes" y "después"
        8. Comparación con estándares de la industria
        
        IMPORTANTE: Responde ÚNICAMENTE en formato JSON válido con esta estructura exacta:
        {{
            "score": número_entre_0_y_100,
            "strengths": ["fortaleza1 con ejemplo específico", "fortaleza2 con detalle", "fortaleza3 con contexto", "fortaleza4 con justificación", "fortaleza5 con evidencia"],
            "weaknesses": ["debilidad1 con ejemplo específico del CV", "debilidad2 con ubicación exacta", "debilidad3 con impacto explicado", "debilidad4 con consecuencias", "debilidad5 con contexto"],
            "recommendations": ["recomendación1 con ejemplo práctico de implementación", "recomendación2 con texto sugerido", "recomendación3 con pasos específicos", "recomendación4 con plantilla", "recomendación5 con mejores prácticas"],
            "keywords": ["palabra1", "palabra2", "palabra3", "palabra4", "palabra5", "palabra6", "palabra7", "palabra8"],
            "analysis_type": "general_health_check",
            "detailed_feedback": "Análisis exhaustivo de 500+ palabras que incluya: evaluación detallada de cada sección del CV, ejemplos específicos de errores encontrados con citas textuales, sugerencias de mejora con ejemplos concretos de texto mejorado, comparación con mejores prácticas de la industria, análisis de impacto en sistemas ATS, recomendaciones de formato específicas con justificación, evaluación de la coherencia y flujo narrativo, análisis de la efectividad comunicativa, sugerencias de reorganización si es necesario, y proyección del impacto de las mejoras sugeridas en la empleabilidad del candidato.",
            "examples": {{
                "spelling_errors": ["Error encontrado: 'texto_original' → Corrección: 'texto_corregido'"],
                "format_improvements": ["Problema: descripción del problema → Solución: ejemplo específico de mejora"],
                "before_after": ["Antes: 'texto original problemático' → Después: 'texto mejorado y optimizado'"]
            }},
            "metrics": {{
                "readability_score": número_0_100,
                "ats_compatibility": número_0_100,
                "professional_impact": número_0_100,
                "structure_quality": número_0_100
            }}
        }}
        
        CV: {cv_text}
        """,
        
        'content_quality_analysis': f"""
        Evalúa exhaustivamente verbos de acción, logros cuantificables, claridad y lenguaje profesional del currículum como un experto en comunicación corporativa y desarrollo profesional.
        
        Analiza DETALLADAMENTE:
        1. Uso de verbos de acción efectivos con análisis de impacto y alternativas
        2. Logros cuantificados con números/porcentajes y contexto de relevancia
        3. Claridad en la comunicación con ejemplos específicos de mejora
        4. Profesionalismo del lenguaje con evaluación de tono y registro
        5. Estructura narrativa y storytelling profesional
        6. Densidad informativa y eficiencia comunicativa
        7. Diferenciación competitiva en la expresión
        8. Adaptabilidad del mensaje a diferentes audiencias
        9. Coherencia estilística y consistencia terminológica
        10. Impacto emocional y persuasivo del contenido
        
        IMPORTANTE: Responde ÚNICAMENTE en formato JSON válido con esta estructura exacta:
        {{
            "score": número_entre_0_y_100,
            "strengths": ["fortaleza1 con ejemplo específico del CV", "fortaleza2 con análisis de impacto", "fortaleza3 con contexto profesional", "fortaleza4 con diferenciación", "fortaleza5 con efectividad comunicativa", "fortaleza6 con valor agregado"],
            "weaknesses": ["debilidad1 con ejemplo textual específico", "debilidad2 con impacto en percepción", "debilidad3 con oportunidad perdida", "debilidad4 con comparación de mejores prácticas", "debilidad5 con consecuencias en empleabilidad", "debilidad6 con análisis de efectividad"],
            "recommendations": ["recomendación1 con ejemplo antes/después", "recomendación2 con verbos alternativos específicos", "recomendación3 con métricas sugeridas", "recomendación4 con reformulación completa", "recomendación5 con estrategia de storytelling", "recomendación6 con optimización de impacto"],
            "keywords": ["palabra1", "palabra2", "palabra3", "palabra4", "palabra5", "palabra6", "palabra7", "palabra8", "palabra9", "palabra10"],
            "analysis_type": "content_quality_analysis",
            "detailed_feedback": "Análisis exhaustivo de 600+ palabras que incluya: evaluación detallada de cada verbo de acción utilizado con sugerencias de alternativas más impactantes, análisis específico de logros cuantificados con contexto de relevancia sectorial, evaluación de claridad comunicativa con ejemplos de reformulación, análisis de profesionalismo del lenguaje con comparación de registros, evaluación de estructura narrativa y coherencia del storytelling, análisis de densidad informativa y eficiencia del mensaje, evaluación de diferenciación competitiva en la expresión, análisis de adaptabilidad del contenido a diferentes audiencias, evaluación de consistencia terminológica y estilística, y proyección del impacto persuasivo en reclutadores y sistemas ATS.",
            "communication_analysis": {{
                "action_verbs_effectiveness": número_0_100,
                "quantified_achievements": número_0_100,
                "clarity_score": número_0_100,
                "professionalism_level": número_0_100,
                "storytelling_quality": número_0_100,
                "persuasive_impact": número_0_100
            }},
            "detailed_examples": {{
                "weak_verbs": ["Verbo débil encontrado: 'responsable de' → Alternativa impactante: 'lideró/optimizó/transformó' con contexto específico"],
                "strong_achievements": ["Logro bien cuantificado identificado con análisis de por qué es efectivo"],
                "clarity_improvements": ["Frase confusa: 'texto original' → Versión clara: 'texto mejorado' con explicación"],
                "professionalism_upgrades": ["Expresión informal: 'texto original' → Versión profesional: 'texto mejorado'"]
            }},
            "optimization_suggestions": {{
                "high_impact_verbs": ["Lista de verbos de alto impacto específicos para el perfil"],
                "quantification_opportunities": ["Oportunidades específicas de cuantificación identificadas"],
                "storytelling_improvements": ["Sugerencias específicas de mejora narrativa"]
            }}
        }}
        
        CV: {cv_text}
        """,
        
        'job_tailoring_optimization': f"""
        Analiza exhaustivamente si el currículum coincide con descripciones de trabajo específicas para una mejor inserción laboral, actuando como un especialista en reclutamiento con experiencia en múltiples industrias.
        
        Evalúa COMPREHENSIVAMENTE:
        1. Alineación estratégica con roles objetivo y análisis de fit cultural
        2. Palabras clave relevantes para la industria con análisis de densidad y posicionamiento
        3. Habilidades transferibles con mapeo de competencias cross-funcionales
        4. Sugerencias de personalización específicas por sector y nivel jerárquico
        5. Análisis de competitividad frente a perfiles similares del mercado
        6. Evaluación de gaps críticos y cómo compensarlos estratégicamente
        7. Optimización de narrativa profesional para diferentes tipos de empleadores
        8. Análisis de tendencias del mercado laboral y adaptación del perfil
        9. Estrategias de diferenciación competitiva por industria
        10. Proyección de empleabilidad y potencial de contratación por sector
        
        IMPORTANTE: Responde ÚNICAMENTE en formato JSON válido con esta estructura exacta:
        {{
            "score": número_entre_0_y_100,
            "strengths": ["fortaleza1 con análisis de mercado específico", "fortaleza2 con ventaja competitiva identificada", "fortaleza3 con alineación sectorial", "fortaleza4 con transferibilidad de skills", "fortaleza5 con diferenciación estratégica", "fortaleza6 con potencial de crecimiento"],
            "weaknesses": ["debilidad1 con impacto en empleabilidad sectorial", "debilidad2 con gap crítico identificado", "debilidad3 con desalineación de mercado", "debilidad4 con competencia desfavorable", "debilidad5 con barrera de entrada", "debilidad6 con limitación de oportunidades"],
            "recommendations": ["recomendación1 con estrategia de personalización específica", "recomendación2 con keywords sectoriales exactas", "recomendación3 con reformulación para industria target", "recomendación4 con compensación de gaps", "recomendación5 con diferenciación competitiva", "recomendación6 con optimización de narrativa"],
            "keywords": ["keyword1_industria", "keyword2_técnica", "keyword3_soft_skill", "keyword4_herramienta", "keyword5_metodología", "keyword6_certificación", "keyword7_sector", "keyword8_nivel", "keyword9_función", "keyword10_tendencia"],
            "analysis_type": "job_tailoring_optimization",
            "detailed_feedback": "Análisis exhaustivo de 650+ palabras que incluya: evaluación detallada de alineación con roles target específicos, análisis de competitividad frente a perfiles similares del mercado, mapeo de habilidades transferibles con ejemplos de aplicación sectorial, identificación de gaps críticos con estrategias de compensación, análisis de keywords sectoriales con densidad óptima y posicionamiento estratégico, evaluación de narrativa profesional con adaptaciones por tipo de empleador, análisis de tendencias del mercado laboral relevantes, estrategias de diferenciación competitiva específicas por industria, proyección de empleabilidad por sector con probabilidades de éxito, y recomendaciones de personalización con ejemplos concretos de reformulación para diferentes oportunidades laborales.",
            "market_analysis": {{
                "industry_alignment": número_0_100,
                "keyword_optimization": número_0_100,
                "skills_transferability": número_0_100,
                "competitive_positioning": número_0_100,
                "market_readiness": número_0_100,
                "growth_potential": número_0_100
            }},
            "tailoring_examples": {{
                "industry_specific_keywords": ["Keyword crítico para sector X con justificación de importancia"],
                "role_adaptations": ["Adaptación específica: 'descripción original' → 'versión optimizada para rol Y'"],
                "skills_repositioning": ["Skill reposicionado: 'presentación actual' → 'enfoque estratégico para industria Z'"],
                "narrative_adjustments": ["Ajuste narrativo: 'versión genérica' → 'versión personalizada para empleador tipo A'"]
            }},
            "competitive_analysis": {{
                "market_advantages": ["Ventaja competitiva específica con contexto de mercado"],
                "improvement_priorities": ["Prioridad de mejora con impacto proyectado en empleabilidad"],
                "differentiation_strategies": ["Estrategia de diferenciación específica con implementación práctica"]
            }}
        }}
        
        CV: {cv_text}
        """,
        
        'ats_compatibility_verification': f"""
        Verifica exhaustivamente la compatibilidad con sistemas ATS (Applicant Tracking Systems) actuando como un especialista en tecnología de reclutamiento con conocimiento profundo de múltiples plataformas ATS.
        
        Asegúrate de que el currículum pase exitosamente a través de los sistemas de seguimiento de solicitantes más utilizados en el mercado.
        
        Evalúa METICULOSAMENTE:
        1. Formato compatible con ATS con análisis de parsing y legibilidad automática
        2. Uso de palabras clave estándar con análisis de densidad y posicionamiento óptimo
        3. Estructura de secciones con evaluación de headers y jerarquía informativa
        4. Elementos que podrían causar problemas con identificación específica de conflictos
        5. Análisis de fuentes, espaciado y elementos gráficos problemáticos
        6. Evaluación de compatibilidad con diferentes versiones de ATS populares
        7. Análisis de metadata y información estructurada
        8. Verificación de campos estándar y formatos de fecha/contacto
        9. Evaluación de longitud y densidad de contenido para parsing óptimo
        10. Análisis de probabilidad de ranking alto en búsquedas automatizadas
        
        IMPORTANTE: Responde ÚNICAMENTE en formato JSON válido con esta estructura exacta:
        {{
            "score": número_entre_0_y_100,
            "strengths": ["fortaleza1 con compatibilidad específica de ATS", "fortaleza2 con ventaja en parsing automático", "fortaleza3 con optimización de keywords", "fortaleza4 con estructura favorable", "fortaleza5 con metadata correcta", "fortaleza6 con ranking potencial alto"],
            "weaknesses": ["debilidad1 con riesgo específico de parsing", "debilidad2 con elemento problemático identificado", "debilidad3 con incompatibilidad de formato", "debilidad4 con pérdida de información", "debilidad5 con ranking desfavorable", "debilidad6 con barrera técnica"],
            "recommendations": ["recomendación1 con solución técnica específica", "recomendación2 con reformateo exacto", "recomendación3 con optimización de keywords", "recomendación4 con ajuste de estructura", "recomendación5 con corrección de metadata", "recomendación6 con mejora de compatibilidad"],
            "keywords": ["keyword1_ats_friendly", "keyword2_standard", "keyword3_industry", "keyword4_technical", "keyword5_role", "keyword6_skill", "keyword7_certification", "keyword8_tool", "keyword9_methodology", "keyword10_level"],
            "analysis_type": "ats_compatibility_verification",
            "detailed_feedback": "Análisis exhaustivo de 700+ palabras que incluya: evaluación detallada de compatibilidad con sistemas ATS principales (Workday, Taleo, iCIMS, Greenhouse, etc.), análisis específico de elementos que causan problemas de parsing con ejemplos concretos, evaluación de estructura de headers y secciones con recomendaciones de optimización, análisis de densidad y posicionamiento de keywords para ranking automático, verificación de formatos de fecha, contacto y campos estándar, evaluación de fuentes y elementos gráficos problemáticos, análisis de metadata y información estructurada, proyección de probabilidad de paso exitoso por filtros automáticos, recomendaciones específicas de reformateo con ejemplos antes/después, y estrategias de optimización para diferentes tipos de ATS con consideraciones técnicas específicas.",
            "ats_analysis": {{
                "parsing_compatibility": número_0_100,
                "keyword_optimization": número_0_100,
                "structure_quality": número_0_100,
                "format_compliance": número_0_100,
                "metadata_accuracy": número_0_100,
                "ranking_potential": número_0_100
            }},
            "technical_issues": {{
                "problematic_elements": ["Elemento problemático específico: 'descripción del problema' → Solución: 'corrección técnica'"],
                "parsing_risks": ["Riesgo de parsing identificado con probabilidad de fallo y solución"],
                "format_conflicts": ["Conflicto de formato: 'problema específico' → Alternativa compatible: 'solución'"],
                "optimization_opportunities": ["Oportunidad de optimización: 'situación actual' → 'mejora sugerida con impacto'"]
            }},
            "ats_compatibility_matrix": {{
                "workday_score": número_0_100,
                "taleo_score": número_0_100,
                "icims_score": número_0_100,
                "greenhouse_score": número_0_100,
                "general_ats_score": número_0_100
            }}
        }}
        
        CV: {cv_text}
        """,
        
        'tone_style_evaluation': f"""
        Evalúa el tono profesional, la idoneidad del idioma y la legibilidad del currículum con análisis exhaustivo.
        
        Analiza en detalle:
        1. **Consistencia del tono profesional**: Evalúa uniformidad, autoridad, confianza
        2. **Apropiación del lenguaje para la industria**: Terminología técnica, jerga profesional, nivel de formalidad
        3. **Legibilidad y fluidez**: Estructura de oraciones, transiciones, claridad
        4. **Impacto y persuasión**: Poder de convencimiento, llamadas a la acción, diferenciación
        5. **Comunicación efectiva**: Concisión, precisión, engagement
        6. **Adaptación al público objetivo**: Alineación con expectativas del reclutador
        
        Proporciona ejemplos específicos de:
        - Frases que demuestran tono profesional vs. informal
        - Terminología técnica bien/mal utilizada
        - Oraciones que necesitan mejora en fluidez
        - Expresiones con alto/bajo impacto persuasivo
        
        IMPORTANTE: Responde ÚNICAMENTE en formato JSON válido con esta estructura exacta:
        {{
            "score": número_entre_0_y_100,
            "strengths": ["fortaleza1 con ejemplo específico", "fortaleza2 con ejemplo específico", "fortaleza3"],
            "weaknesses": ["debilidad1 con ejemplo específico", "debilidad2 con ejemplo específico", "debilidad3"],
            "recommendations": ["recomendación1 con ejemplo antes/después", "recomendación2 con ejemplo", "recomendación3"],
            "keywords": ["palabra_clave_profesional1", "palabra_clave_industria2", "término_técnico3"],
            "analysis_type": "tone_style_evaluation",
            "detailed_feedback": "análisis detallado de tono y estilo con ejemplos específicos",
            "tone_metrics": {{
                "professionalism_score": número_0_100,
                "industry_language_score": número_0_100,
                "readability_score": número_0_100,
                "persuasion_impact_score": número_0_100,
                "consistency_score": número_0_100
            }},
            "language_analysis": {{
                "formal_expressions": ["expresión formal 1", "expresión formal 2"],
                "informal_expressions": ["expresión informal 1", "expresión informal 2"],
                "technical_terms_used": ["término técnico 1", "término técnico 2"],
                "missing_industry_terms": ["término faltante 1", "término faltante 2"]
            }},
            "improvement_examples": {{
                "weak_phrases": [
                    {{"original": "frase débil", "improved": "frase mejorada", "reason": "razón de mejora"}},
                    {{"original": "otra frase débil", "improved": "otra frase mejorada", "reason": "razón de mejora"}}
                ],
                "tone_adjustments": [
                    {{"section": "sección del CV", "current_tone": "tono actual", "recommended_tone": "tono recomendado", "example": "ejemplo de mejora"}}
                ]
            }}
        }}
        
        CV: {cv_text}
        """,
        
        'industry_role_feedback': f"""
        Proporciona asesoramiento personalizado exhaustivo basado en la industria y rol específico identificados en el CV.
        
        Analiza en profundidad:
        1. **Relevancia para la industria específica**: Alineación con estándares, certificaciones, experiencia sectorial
        2. **Competencias clave del rol**: Skills técnicos, soft skills, competencias emergentes
        3. **Tendencias del mercado laboral**: Demanda actual, skills en crecimiento, tecnologías emergentes
        4. **Recomendaciones específicas del sector**: Certificaciones valoradas, experiencias clave, networking
        5. **Posicionamiento competitivo**: Ventajas diferenciales, gaps vs. competencia
        6. **Proyección de carrera**: Próximos pasos, roles objetivo, desarrollo profesional
        
        Identifica y proporciona ejemplos específicos de:
        - Experiencias que demuestran expertise sectorial
        - Skills técnicos específicos de la industria presentes/ausentes
        - Logros cuantificados relevantes para el sector
        - Terminología y keywords específicas de la industria
        - Certificaciones y formación valoradas en el sector
        
        IMPORTANTE: Responde ÚNICAMENTE en formato JSON válido con esta estructura exacta:
        {{
            "score": número_entre_0_y_100,
            "strengths": ["fortaleza1 con ejemplo específico del sector", "fortaleza2 con contexto industrial", "fortaleza3"],
            "weaknesses": ["debilidad1 con impacto en la industria", "debilidad2 con comparación sectorial", "debilidad3"],
            "recommendations": ["recomendación1 específica del sector con ejemplo", "recomendación2 con certificación sugerida", "recomendación3"],
            "keywords": ["keyword_industria1", "skill_técnico2", "certificación3", "herramienta4"],
            "analysis_type": "industry_role_feedback",
            "detailed_feedback": "análisis detallado específico de la industria con ejemplos y contexto sectorial",
            "industry_analysis": {{
                "identified_industry": "industria identificada",
                "target_role": "rol objetivo identificado",
                "industry_alignment_score": número_0_100,
                "role_readiness_score": número_0_100,
                "market_competitiveness_score": número_0_100
            }},
            "sector_requirements": {{
                "essential_skills": ["skill esencial 1", "skill esencial 2", "skill esencial 3"],
                "preferred_skills": ["skill preferido 1", "skill preferido 2"],
                "missing_skills": ["skill faltante 1", "skill faltante 2"],
                "relevant_certifications": ["certificación 1", "certificación 2"]
            }},
            "market_insights": {{
                "current_trends": ["tendencia 1", "tendencia 2", "tendencia 3"],
                "emerging_technologies": ["tecnología 1", "tecnología 2"],
                "salary_range": "rango salarial estimado",
                "growth_outlook": "perspectiva de crecimiento del sector"
            }},
            "career_development": {{
                "next_steps": ["paso 1 específico", "paso 2 con timeline", "paso 3"],
                "target_companies": ["tipo de empresa 1", "tipo de empresa 2"],
                "networking_opportunities": ["evento/plataforma 1", "asociación profesional 2"],
                "skill_development_priority": ["skill prioritario 1", "skill prioritario 2"]
            }}
        }}
        
        CV: {cv_text}
        """,
        
        'benchmarking_comparison': f"""
        Compara exhaustivamente el currículum con ejemplos exitosos y estándares de la industria para roles similares.
        
        Realiza análisis comparativo detallado:
        1. **Comparación con estándares de la industria**: Benchmarks sectoriales, métricas de rendimiento, niveles esperados
        2. **Identificación de brechas críticas**: Gaps en experiencia, skills, logros, formación
        3. **Mejores prácticas aplicables**: Estructuras exitosas, formatos optimizados, contenido efectivo
        4. **Posicionamiento competitivo**: Ranking vs. competencia, ventajas diferenciales, áreas de mejora
        5. **Análisis de perfiles top-tier**: Comparación con candidatos exitosos del mismo nivel
        6. **Evaluación de mercado**: Posición relativa, competitividad, oportunidades
        
        Proporciona ejemplos específicos y comparaciones concretas:
        - Logros cuantificados vs. estándares del mercado
        - Estructura y formato vs. mejores prácticas
        - Skills y certificaciones vs. perfiles exitosos
        - Experiencia y progresión vs. trayectorias típicas
        - Lenguaje y terminología vs. estándares profesionales
        
        IMPORTANTE: Responde ÚNICAMENTE en formato JSON válido con esta estructura exacta:
        {{
            "score": número_entre_0_y_100,
            "strengths": ["fortaleza1 vs. benchmark específico", "fortaleza2 con comparación cuantificada", "fortaleza3"],
            "weaknesses": ["debilidad1 vs. estándar industrial", "debilidad2 con gap específico", "debilidad3"],
            "recommendations": ["recomendación1 basada en mejores prácticas", "recomendación2 con ejemplo de mejora", "recomendación3"],
            "keywords": ["keyword_benchmark1", "término_estándar2", "skill_competitivo3"],
            "analysis_type": "benchmarking_comparison",
            "detailed_feedback": "análisis detallado de comparación y benchmarking con ejemplos específicos",
            "benchmark_analysis": {{
                "industry_percentile": número_0_100,
                "experience_level_match": "junior/mid/senior/executive",
                "competitive_position": "below_average/average/above_average/top_tier",
                "market_readiness_score": número_0_100
            }},
            "comparison_metrics": {{
                "format_vs_standard": {{
                    "current_score": número_0_100,
                    "industry_average": número_0_100,
                    "top_performers": número_0_100,
                    "gap_analysis": "descripción del gap"
                }},
                "content_vs_benchmark": {{
                    "achievements_quality": número_0_100,
                    "skills_relevance": número_0_100,
                    "experience_depth": número_0_100,
                    "industry_alignment": número_0_100
                }}
            }},
            "best_practices_analysis": {{
                "successful_patterns": ["patrón exitoso 1", "patrón exitoso 2", "patrón exitoso 3"],
                "missing_elements": ["elemento faltante 1", "elemento faltante 2"],
                "optimization_opportunities": [
                    {{"area": "área de mejora", "current_state": "estado actual", "best_practice": "mejor práctica", "expected_impact": "impacto esperado"}}
                ]
            }},
            "competitive_analysis": {{
                "strengths_vs_market": ["ventaja competitiva 1", "ventaja competitiva 2"],
                "weaknesses_vs_market": ["desventaja 1", "desventaja 2"],
                "differentiation_opportunities": ["oportunidad 1", "oportunidad 2"],
                "market_positioning": "posicionamiento en el mercado"
            }}
        }}
        
        CV: {cv_text}
        """,
        
        'ai_improvement_suggestions': f"""
        Proporciona consejos de mejora avanzados basados en IA y machine learning para optimizar el currículum de manera integral.
        
        Analiza y proporciona sugerencias inteligentes sobre:
        1. **Optimizaciones basadas en IA**: Análisis de patrones exitosos, predicciones de rendimiento, scoring automático
        2. **Sugerencias específicas de mejora**: Recomendaciones personalizadas, ajustes de contenido, optimización de formato
        3. **Tendencias actuales del mercado**: Insights de datos de reclutamiento, skills emergentes, demandas del mercado
        4. **Estrategias de diferenciación**: Elementos únicos, propuesta de valor, posicionamiento competitivo
        5. **Optimización para ATS y IA**: Compatibilidad con sistemas automatizados, keywords estratégicas
        6. **Predicciones de éxito**: Probabilidad de éxito, áreas de alto impacto, ROI de mejoras
        
        Incluye ejemplos específicos y transformaciones concretas:
        - Antes/después de secciones optimizadas
        - Keywords con mayor impacto según IA
        - Estructuras de frases más efectivas
        - Métricas y logros optimizados
        - Formatos que maximizan el engagement
        
        IMPORTANTE: Responde ÚNICAMENTE en formato JSON válido con esta estructura exacta:
        {{
            "score": número_entre_0_y_100,
            "strengths": ["fortaleza1 identificada por IA", "fortaleza2 con potencial de optimización", "fortaleza3"],
            "weaknesses": ["debilidad1 detectada por análisis automático", "debilidad2 con impacto cuantificado", "debilidad3"],
            "recommendations": ["recomendación1 basada en ML con ejemplo", "recomendación2 con predicción de impacto", "recomendación3"],
            "keywords": ["keyword_ai_optimized1", "trending_skill2", "high_impact_term3"],
            "analysis_type": "ai_improvement_suggestions",
            "detailed_feedback": "sugerencias detalladas de mejora basadas en IA con ejemplos específicos",
            "ai_insights": {{
                "optimization_score": número_0_100,
                "market_alignment_score": número_0_100,
                "ats_compatibility_prediction": número_0_100,
                "success_probability": número_0_100
            }},
            "smart_optimizations": {{
                "high_impact_changes": [
                    {{"section": "sección", "current": "contenido actual", "optimized": "contenido optimizado", "impact_score": número_0_100, "reasoning": "razón basada en IA"}}
                ],
                "keyword_optimization": [
                    {{"current_keyword": "keyword actual", "optimized_keyword": "keyword optimizado", "frequency_recommendation": "frecuencia recomendada", "context": "contexto de uso"}}
                ],
                "structure_improvements": [
                    {{"area": "área de mejora", "current_structure": "estructura actual", "recommended_structure": "estructura recomendada", "ai_reasoning": "razón basada en datos"}}
                ]
            }},
            "market_intelligence": {{
                "trending_skills": ["skill emergente 1", "skill emergente 2", "skill emergente 3"],
                "declining_skills": ["skill en declive 1", "skill en declive 2"],
                "industry_predictions": ["predicción 1", "predicción 2"],
                "salary_impact_factors": ["factor 1", "factor 2"]
            }},
            "personalization": {{
                "career_stage_optimization": "optimización específica para nivel de carrera",
                "industry_customization": "personalización para industria específica",
                "role_targeting": "enfoque para rol objetivo",
                "geographic_considerations": "consideraciones geográficas del mercado"
            }},
            "predictive_analysis": {{
                "interview_probability": número_0_100,
                "salary_negotiation_strength": número_0_100,
                "career_advancement_potential": número_0_100,
                "market_competitiveness": número_0_100
            }}
        }}
        
        CV: {cv_text}
        """,
        
        'visual_design_assessment': f"""
        Evalúa exhaustivamente el atractivo visual, la legibilidad y la presentación profesional del currículum desde una perspectiva de diseño UX/UI.
        
        Analiza en detalle:
        1. **Diseño visual y layout**: Jerarquía visual, balance, alineación, consistencia tipográfica
        2. **Legibilidad y organización**: Flujo de lectura, espaciado, contraste, accesibilidad
        3. **Uso efectivo del espacio**: Distribución, márgenes, densidad de información, respiración visual
        4. **Impresión profesional general**: Primera impresión, credibilidad visual, modernidad
        5. **Experiencia de usuario**: Facilidad de navegación, escaneabilidad, jerarquía de información
        6. **Adaptabilidad**: Compatibilidad con diferentes formatos, impresión, visualización digital
        
        Proporciona análisis específico de elementos visuales:
        - Tipografía: fuentes, tamaños, jerarquías, legibilidad
        - Color y contraste: esquema cromático, accesibilidad, profesionalismo
        - Espaciado: márgenes, padding, line-height, secciones
        - Estructura: grid, alineación, balance visual
        - Elementos gráficos: iconos, líneas, separadores, bullets
        
        IMPORTANTE: Responde ÚNICAMENTE en formato JSON válido con esta estructura exacta:
        {{
            "score": número_entre_0_y_100,
            "strengths": ["fortaleza1 de diseño específica", "fortaleza2 visual con detalle", "fortaleza3"],
            "weaknesses": ["debilidad1 de layout específica", "debilidad2 de legibilidad", "debilidad3"],
            "recommendations": ["recomendación1 de diseño con ejemplo", "recomendación2 de mejora visual", "recomendación3"],
            "keywords": ["término_diseño1", "concepto_visual2", "elemento_ux3"],
            "analysis_type": "visual_design_assessment",
            "detailed_feedback": "análisis detallado de diseño visual con ejemplos específicos",
            "design_metrics": {{
                "visual_hierarchy_score": número_0_100,
                "readability_score": número_0_100,
                "professional_appearance_score": número_0_100,
                "space_utilization_score": número_0_100,
                "consistency_score": número_0_100,
                "modern_design_score": número_0_100
            }},
            "visual_elements_analysis": {{
                "typography": {{
                    "font_choices": "evaluación de fuentes",
                    "hierarchy_effectiveness": número_0_100,
                    "readability_assessment": "evaluación de legibilidad",
                    "size_consistency": número_0_100
                }},
                "layout_structure": {{
                    "grid_system": "evaluación del sistema de grid",
                    "alignment_quality": número_0_100,
                    "balance_assessment": "evaluación del balance visual",
                    "flow_effectiveness": número_0_100
                }},
                "spacing_whitespace": {{
                    "margin_usage": "evaluación de márgenes",
                    "section_separation": número_0_100,
                    "breathing_room": "evaluación del espacio en blanco",
                    "density_optimization": número_0_100
                }}
            }},
            "ux_assessment": {{
                "scannability": número_0_100,
                "information_hierarchy": número_0_100,
                "user_journey": "evaluación del flujo de lectura",
                "accessibility_score": número_0_100
            }},
            "design_improvements": {{
                "priority_fixes": [
                    {{"issue": "problema visual", "solution": "solución específica", "impact": "impacto esperado", "difficulty": "fácil/medio/difícil"}}
                ],
                "enhancement_suggestions": [
                    {{"area": "área de mejora", "current_state": "estado actual", "recommended_change": "cambio recomendado", "visual_impact": número_0_100}}
                ],
                "modern_trends": ["tendencia de diseño 1", "tendencia de diseño 2", "tendencia de diseño 3"]
            }},
            "format_compatibility": {{
                "print_readiness": número_0_100,
                "digital_optimization": número_0_100,
                "ats_visual_compatibility": número_0_100,
                "mobile_friendliness": número_0_100
            }}
        }}
        
        CV: {cv_text}
        """,
        
        'comprehensive_score': f"""
        Proporciona una puntuación completa del currículum con desglose detallado y retroalimentación procesable.
        
        Incluye:
        1. Puntuación general detallada (0-100)
        2. Desglose por categorías (formato, contenido, ATS, etc.)
        3. Fortalezas principales identificadas
        4. Áreas de mejora prioritarias
        5. Plan de acción específico y procesable
        
        IMPORTANTE: Responde ÚNICAMENTE en formato JSON válido con esta estructura exacta:
        {{
            "score": número_entre_0_y_100,
            "strengths": ["fortaleza1", "fortaleza2", "fortaleza3"],
            "weaknesses": ["debilidad1", "debilidad2", "debilidad3"],
            "recommendations": ["recomendación1", "recomendación2", "recomendación3"],
            "keywords": ["palabra_clave1", "palabra_clave2", "palabra_clave3"],
            "analysis_type": "comprehensive_score",
            "detailed_feedback": "Análisis completo: [Puntuación general: X/100] [Formato: Y/100] [Contenido: Z/100] [Compatibilidad ATS: W/100] - Detalles específicos y plan de acción aquí"
        }}
        
        CV: {cv_text}
        """
    }
    
    return analysis_prompts.get(analysis_type, analysis_prompts['general_health_check'])

def analyze_cv_with_openai(cv_text, analysis_type):
    """Analizar CV usando OpenAI"""
    prompt = get_analysis_prompt(analysis_type, cv_text)
    
    system_prompt = """Eres un experto en recursos humanos, reclutamiento y sistemas ATS (Applicant Tracking System). 
    Tu objetivo es ayudar a los candidatos a optimizar sus currículums para maximizar sus posibilidades de pasar los filtros ATS y llegar a la entrevista.
    
    Responde SIEMPRE en formato JSON con la siguiente estructura:
    {
        "score": número (0-100),
        "strengths": ["fortaleza1", "fortaleza2", ...],
        "weaknesses": ["debilidad1", "debilidad2", ...],
        "recommendations": ["recomendación1", "recomendación2", ...],
        "keywords": ["palabra_clave1", "palabra_clave2", ...],
        "analysis_type": "tipo_de_análisis",
        "detailed_feedback": "retroalimentación detallada específica del tipo de análisis"
    }"""
    
    try:
        response = OPENAI_CLIENT.chat.completions.create(
            model="gpt-3.5-turbo",  # Usar gpt-3.5-turbo que es más estable
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": prompt}
            ],
            max_tokens=2000,
            temperature=0.7
        )
        
        analysis_text = response.choices[0].message.content
        analysis = json.loads(analysis_text)
        
        # Asegurar campos requeridos
        analysis.setdefault("keywords", [])
        analysis.setdefault("analysis_type", analysis_type)
        analysis.setdefault("ai_provider", "openai")
            
        return analysis
    
    except Exception as e:
        print(f"Error al analizar con OpenAI: {e}")
        return get_error_analysis(analysis_type, "openai", str(e))

def analyze_cv_with_anthropic(cv_text, analysis_type):
    """Analizar CV usando Anthropic Claude"""
    try:
        import anthropic
        
        # Verificar API key
        api_key = os.getenv('ANTHROPIC_API_KEY')
        if not api_key or api_key == 'your_anthropic_api_key_here':
            return get_error_analysis(analysis_type, "anthropic", "API Key de Anthropic no configurada")
        
        # Configurar cliente de Anthropic v0.25.0 - sin proxies
        client = anthropic.Anthropic(api_key=api_key)
        
        prompt = get_analysis_prompt(analysis_type, cv_text)
        
        system_prompt = """Eres un experto en recursos humanos, reclutamiento y sistemas ATS. 
        Responde SIEMPRE en formato JSON válido con la estructura especificada."""
        
        response = client.messages.create(
            model="claude-3-5-sonnet-20241022",  # Usar modelo disponible
            max_tokens=2000,
            temperature=0.7,
            system=system_prompt,
            messages=[
                {"role": "user", "content": prompt + "\n\nResponde en formato JSON con: score, strengths, weaknesses, recommendations, keywords, analysis_type, detailed_feedback"}
            ]
        )
        
        analysis_text = response.content[0].text
        
        # Limpiar respuesta si contiene markdown
        if "```json" in analysis_text:
            analysis_text = analysis_text.split("```json")[1].split("```")[0]
        elif "```" in analysis_text:
            analysis_text = analysis_text.split("```")[1]
        
        analysis = json.loads(analysis_text.strip())
        
        # Asegurar campos requeridos
        analysis.setdefault("keywords", [])
        analysis.setdefault("analysis_type", analysis_type)
        analysis.setdefault("ai_provider", "anthropic")
        
        return analysis
        
    except ImportError:
        return get_error_analysis(analysis_type, "anthropic", "Librería de Anthropic no instalada")
    except json.JSONDecodeError as e:
        print(f"Error de JSON en Anthropic: {e}")
        print(f"Respuesta recibida: {analysis_text if 'analysis_text' in locals() else 'No disponible'}")
        return get_error_analysis(analysis_type, "anthropic", f"Error de formato JSON: {str(e)}")
    except Exception as e:
        print(f"Error al analizar con Anthropic: {e}")
        return get_error_analysis(analysis_type, "anthropic", str(e))

def analyze_cv_with_gemini(cv_text, analysis_type):
    """Analizar CV usando Google Gemini"""
    try:
        # Intentar importar con diferentes métodos
        try:
            import google.generativeai as genai
        except ImportError:
            try:
                from google import generativeai as genai
            except ImportError:
                return get_error_analysis(analysis_type, "gemini", "Librería de Google Generative AI no disponible. Ejecute: pip install google-generativeai")
        
        # Verificar API key
        api_key = os.getenv('GEMINI_API_KEY')
        if not api_key or api_key == 'your_gemini_api_key_here':
            return get_error_analysis(analysis_type, "gemini", "API Key de Gemini no configurada")
        
        # Configurar Gemini
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-1.5-flash')  # Usar modelo disponible
        
        prompt = get_analysis_prompt(analysis_type, cv_text)
        
        system_instruction = """Eres un experto en recursos humanos y sistemas ATS. 
        Responde SIEMPRE en formato JSON válido con la estructura especificada."""
        
        full_prompt = f"{system_instruction}\n\n{prompt}\n\nResponde en formato JSON con: score, strengths, weaknesses, recommendations, keywords, analysis_type, detailed_feedback"
        
        response = model.generate_content(full_prompt)
        analysis_text = response.text
        
        # Limpiar respuesta si contiene markdown
        if "```json" in analysis_text:
            analysis_text = analysis_text.split("```json")[1].split("```")[0]
        elif "```" in analysis_text:
            analysis_text = analysis_text.split("```")[1]
        
        analysis = json.loads(analysis_text.strip())
        
        # Asegurar campos requeridos
        analysis.setdefault("keywords", [])
        analysis.setdefault("analysis_type", analysis_type)
        analysis.setdefault("ai_provider", "gemini")
        
        return analysis
        
    except ImportError:
        return get_error_analysis(analysis_type, "gemini", "Librería de Google Generative AI no instalada")
    except json.JSONDecodeError as e:
        print(f"Error de JSON en Gemini: {e}")
        print(f"Respuesta recibida: {analysis_text if 'analysis_text' in locals() else 'No disponible'}")
        return get_error_analysis(analysis_type, "gemini", f"Error de formato JSON: {str(e)}")
    except Exception as e:
        print(f"Error al analizar con Gemini: {e}")
        return get_error_analysis(analysis_type, "gemini", str(e))

def get_error_analysis(analysis_type, ai_provider, error_message):
    """Retornar análisis de error estándar"""
    return {
        "score": 0,
        "strengths": [f"Error al procesar con {ai_provider}"],
        "weaknesses": ["No se pudo completar el análisis"],
        "recommendations": ["Intente nuevamente más tarde o use otro proveedor de IA"],
        "keywords": [],
        "analysis_type": analysis_type,
        "ai_provider": ai_provider,
        "detailed_feedback": f"Error: {error_message}",
        "error": True
    }

def analyze_cv_with_ai(cv_text):
    """Función legacy para compatibilidad - usar OpenAI por defecto"""
    return analyze_cv_with_openai(cv_text, 'general_health_check')

def save_cv_analysis(user_id, filename, content, analysis):
    """Guardar análisis de CV en S3 y referencia en la base de datos"""
    from s3_utils import save_analysis_to_s3, delete_old_analysis_for_section
    
    # Obtener información del análisis
    analysis_type = analysis.get('analysis_type', 'general_health_check')
    ai_provider = analysis.get('ai_provider', 'openai')
    
    # Eliminar análisis anterior del mismo tipo y proveedor
    delete_old_analysis_for_section(user_id, analysis_type, ai_provider)
    
    # Guardar análisis completo en S3
    s3_key = save_analysis_to_s3(user_id, analysis, analysis_type, ai_provider)
    
    connection = get_db_connection()
    if connection:
        cursor = connection.cursor()
        
        # Eliminar análisis anterior del mismo tipo y proveedor de la base de datos
        cursor.execute("""
            DELETE FROM feedback 
            WHERE resume_id IN (
                SELECT id FROM resumes WHERE user_id = %s
            ) AND analysis_type = %s AND ai_provider = %s
        """, (user_id, analysis_type, ai_provider))
        
        # Insertar un registro mínimo en resumes (sin el contenido completo)
        cursor.execute(
            "INSERT INTO resumes (user_id, filename, content) VALUES (%s, %s, %s) RETURNING id",
            (user_id, filename, "Análisis almacenado en S3 - contenido no local")
        )
        resume_id = cursor.fetchone()['id']
        
        # Insertar feedback con referencia a S3
        cursor.execute("""
            INSERT INTO feedback (resume_id, score, strengths, weaknesses, recommendations, keywords, s3_key, analysis_type, ai_provider) 
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
        """, (
            resume_id, 
            analysis['score'], 
            json.dumps(analysis['strengths']), 
            json.dumps(analysis['weaknesses']), 
            json.dumps(analysis['recommendations']),
            json.dumps(analysis['keywords']),
            s3_key,
            analysis_type,
            ai_provider
        ))
        
        connection.commit()
        cursor.close()
        connection.close()
        
        print(f"Análisis guardado - S3: {s3_key}, DB: resume_id {resume_id}")

def get_latest_cv_analysis(user_id):
    """Obtener el análisis de CV más reciente del usuario desde S3"""
    from s3_utils import get_analysis_from_s3
    
    connection = get_db_connection()
    if connection:
        cursor = connection.cursor()
        cursor.execute("""
            SELECT f.s3_key, f.analysis_type, f.ai_provider, f.created_at
            FROM feedback f
            INNER JOIN resumes r ON f.resume_id = r.id
            WHERE r.user_id = %s AND f.s3_key IS NOT NULL
            ORDER BY f.created_at DESC
            LIMIT 1
        """, (user_id,))
        
        result = cursor.fetchone()
        cursor.close()
        connection.close()
        
        if result and result['s3_key']:
            # Recuperar análisis completo desde S3
            analysis_data = get_analysis_from_s3(result['s3_key'])
            if analysis_data:
                return analysis_data
        
        # Fallback: buscar en base de datos local (para compatibilidad)
        cursor = connection.cursor()
        cursor.execute("""
            SELECT f.score, f.strengths, f.weaknesses, f.recommendations, f.keywords, r.content
            FROM feedback f
            INNER JOIN resumes r ON f.resume_id = r.id
            WHERE r.user_id = %s
            ORDER BY f.created_at DESC
            LIMIT 1
        """, (user_id,))
        
        result = cursor.fetchone()
        cursor.close()
        connection.close()
        
        if result:
            return {
                'score': result['score'],
                'strengths': json.loads(result['strengths']) if result['strengths'] else [],
                'weaknesses': json.loads(result['weaknesses']) if result['weaknesses'] else [],
                'recommendations': json.loads(result['recommendations']) if result['recommendations'] else [],
                'keywords': json.loads(result['keywords']) if result['keywords'] else [],
                'content': result['content']
            }
    return None

def get_user_cv_analyses(user_id):
    """Obtener todos los análisis de CV de un usuario organizados por tipo y proveedor"""
    from s3_utils import get_analysis_from_s3
    
    connection = get_db_connection()
    if not connection:
        return {}
    
    cursor = connection.cursor()
    cursor.execute("""
        SELECT f.s3_key, f.analysis_type, f.ai_provider, f.created_at, f.score
        FROM feedback f
        INNER JOIN resumes r ON f.resume_id = r.id
        WHERE r.user_id = %s AND f.s3_key IS NOT NULL
        ORDER BY f.analysis_type, f.ai_provider, f.created_at DESC
    """, (user_id,))
    
    results = cursor.fetchall()
    cursor.close()
    connection.close()
    
    analyses = {}
    
    for result in results:
        analysis_type = result['analysis_type']
        ai_provider = result['ai_provider']
        
        # Crear estructura anidada si no existe
        if analysis_type not in analyses:
            analyses[analysis_type] = {}
        
        # Solo mantener el más reciente por tipo y proveedor
        if ai_provider not in analyses[analysis_type]:
            # Recuperar análisis completo desde S3
            analysis_data = get_analysis_from_s3(result['s3_key'])
            if analysis_data:
                # Aplanar la estructura para que el template pueda acceder directamente
                flattened_analysis = {
                    'score': analysis_data.get('score', 0),
                    'strengths': analysis_data.get('strengths', []),
                    'weaknesses': analysis_data.get('weaknesses', []),
                    'recommendations': analysis_data.get('recommendations', []),
                    'keywords': analysis_data.get('keywords', []),
                    'created_at': result['created_at'],
                    's3_key': result['s3_key'],
                    'analysis': analysis_data  # Mantener el análisis completo también
                }
                analyses[analysis_type][ai_provider] = flattened_analysis
    
    return analyses

def generate_smart_search_terms(cv_analysis):
    """Generar términos de búsqueda inteligentes basados en el análisis de CV usando IA"""
    try:
        # Combinar información del CV para generar términos de búsqueda
        cv_info = {
            'strengths': cv_analysis.get('strengths', []),
            'keywords': cv_analysis.get('keywords', []),
            'content_preview': cv_analysis.get('content', '')[:1000]  # Primeros 1000 caracteres
        }
        
        prompt = f"""
        Basándote en el siguiente análisis de CV, genera términos de búsqueda específicos para encontrar empleos relevantes.
        
        Fortalezas del candidato: {cv_info['strengths']}
        Palabras clave del CV: {cv_info['keywords']}
        Contenido del CV (muestra): {cv_info['content_preview']}
        
        Genera 5 términos de búsqueda específicos y relevantes que ayuden a encontrar empleos compatibles.
        Los términos deben ser:
        1. Específicos para el perfil profesional
        2. Incluir tecnologías, habilidades o roles mencionados
        3. Ser términos que realmente se usen en ofertas de trabajo
        
        Responde SOLO con una lista de términos separados por comas, sin explicaciones adicionales.
        Ejemplo: "Desarrollador Python, Analista de datos, Machine Learning, Django, SQL"
        """
        
        response = OPENAI_CLIENT.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "Eres un experto en reclutamiento que ayuda a generar términos de búsqueda efectivos para encontrar empleos relevantes."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=200,
            temperature=0.7
        )
        
        terms_text = response.choices[0].message.content.strip()
        search_terms = [term.strip() for term in terms_text.split(',')]
        
        # Filtrar términos vacíos y limitar a 5
        search_terms = [term for term in search_terms if term][:5]
        
        return search_terms
        
    except Exception as e:
        print(f"Error generando términos de búsqueda: {e}")
        # Fallback: usar palabras clave del CV
        keywords = cv_analysis.get('keywords', [])
        return keywords[:5] if keywords else ['desarrollador', 'analista', 'programador']

def remove_duplicate_jobs(jobs):
    """Eliminar trabajos duplicados basándose en título y empresa"""
    seen = set()
    unique_jobs = []
    
    for job in jobs:
        # Crear una clave única basada en título y empresa
        key = f"{job.get('title', '').lower().strip()}_{job.get('company', '').lower().strip()}"
        
        if key not in seen:
            seen.add(key)
            unique_jobs.append(job)
    
    return unique_jobs

def calculate_ai_job_compatibility(job, cv_analysis):
    """Calcular compatibilidad entre un trabajo y el CV usando IA con ponderación por área de experiencia"""
    try:
        # Preparar información del trabajo
        job_info = {
            'title': job.get('title', ''),
            'description': job.get('description', ''),
            'company': job.get('company', ''),
            'location': job.get('location', '')
        }
        
        # Preparar información del CV
        cv_info = {
            'strengths': cv_analysis.get('strengths', []),
            'keywords': cv_analysis.get('keywords', []),
            'score': cv_analysis.get('score', 0),
            'experience_areas': cv_analysis.get('experience_areas', []),
            'skill_level': cv_analysis.get('skill_level', 'intermedio')
        }
        
        prompt = f"""
        Analiza la compatibilidad entre este trabajo y el perfil del candidato, aplicando ponderación según área de experiencia.
        
        TRABAJO:
        Título: {job_info['title']}
        Empresa: {job_info['company']}
        Descripción: {job_info['description'][:500]}...
        
        PERFIL DEL CANDIDATO:
        Fortalezas principales: {cv_info['strengths']}
        Palabras clave del CV: {cv_info['keywords']}
        Áreas de experiencia: {cv_info['experience_areas']}
        Nivel de habilidad: {cv_info['skill_level']}
        Puntuación ATS del CV: {cv_info['score']}/100
        
        INSTRUCCIONES DE PONDERACIÓN:
        - Si el trabajo está en un área donde el candidato NO tiene experiencia: reducir compatibilidad en 20-40%
        - Si el trabajo requiere habilidades que el candidato no domina: reducir compatibilidad en 15-30%
        - Si el nivel del puesto es muy superior a la experiencia del candidato: reducir compatibilidad en 10-25%
        - Si hay coincidencia perfecta de área y habilidades: mantener o aumentar compatibilidad
        
        Calcula un porcentaje de compatibilidad del 0 al 100 considerando:
        1. Coincidencia de área de experiencia (peso: 35%)
        2. Coincidencia de habilidades técnicas (peso: 30%)
        3. Nivel del puesto vs experiencia (peso: 20%)
        4. Palabras clave coincidentes (peso: 15%)
        
        IMPORTANTE: No todos los trabajos deben tener alta compatibilidad. Sé realista con las puntuaciones.
        Responde SOLO con el número del porcentaje (ejemplo: 65)
        """
        
        response = OPENAI_CLIENT.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "Eres un experto en recursos humanos que evalúa la compatibilidad entre candidatos y ofertas de trabajo. Eres crítico y realista con las puntuaciones, no das puntuaciones altas a menos que haya una excelente coincidencia."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=50,
            temperature=0.2
        )
        
        compatibility_text = response.choices[0].message.content.strip()
        
        # Extraer el número del texto
        import re
        numbers = re.findall(r'\d+', compatibility_text)
        if numbers:
            compatibility = min(int(numbers[0]), 100)  # Limitar a 100
            return max(compatibility, 0)  # Asegurar que no sea negativo
        
        return 50  # Valor por defecto si no se puede extraer
        
    except Exception as e:
        print(f"Error calculando compatibilidad IA: {e}")
        # Fallback: usar método básico mejorado
        return calculate_basic_compatibility(job, cv_analysis)

def calculate_basic_compatibility(job, cv_analysis):
    """Método básico de compatibilidad sin IA como fallback con ponderación mejorada"""
    try:
        compatibility_score = 0
        
        # Texto del trabajo en minúsculas
        job_text = f"{job.get('title', '')} {job.get('description', '')}".lower()
        
        # Verificar coincidencias con fortalezas (peso mayor)
        strengths = cv_analysis.get('strengths', [])
        strength_matches = 0
        for strength in strengths:
            if isinstance(strength, str) and strength.lower() in job_text:
                strength_matches += 1
                compatibility_score += 12
        
        # Verificar coincidencias con palabras clave
        keywords = cv_analysis.get('keywords', [])
        keyword_matches = 0
        for keyword in keywords:
            if isinstance(keyword, str) and keyword.lower() in job_text:
                keyword_matches += 1
                compatibility_score += 8
        
        # Verificar áreas de experiencia
        experience_areas = cv_analysis.get('experience_areas', [])
        area_match = False
        for area in experience_areas:
            if isinstance(area, str) and area.lower() in job_text:
                area_match = True
                compatibility_score += 20
                break
        
        # Puntuación base del CV (reducida)
        base_score = cv_analysis.get('score', 50)
        compatibility_score += base_score * 0.2
        
        # Penalización si no hay coincidencias importantes
        if strength_matches == 0:
            compatibility_score *= 0.7  # Reducir 30%
        
        if not area_match and len(experience_areas) > 0:
            compatibility_score *= 0.8  # Reducir 20% si no coincide área
        
        # Asegurar que no todos los trabajos tengan alta compatibilidad
        if compatibility_score > 85:
            compatibility_score = 85  # Máximo realista para método básico
        
        # Limitar entre 15 y 85 para ser más realista
        return max(15, min(int(compatibility_score), 85))
        
    except Exception as e:
        print(f"Error en compatibilidad básica: {e}")
        return 45  # Valor más realista por defecto

@app.route('/cv_builder')
def cv_builder():
    """Alias para create_cv - Constructor de CV"""
    return create_cv()

@app.route('/create_cv')
def create_cv():
    """Constructor de CV estilo Harvard"""
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    # Verificar restricciones de suscripción
    user_id = session.get('user_id')
    can_create, message = check_user_limits(user_id, 'cv_creation')
    
    if not can_create:
        flash(f'Restricción de plan: {message}', 'error')
        return redirect(url_for('dashboard'))
    
    return render_template('create_cv.html')

@app.route('/get_user_cv_data', methods=['GET'])
def get_user_cv_data():
    """Obtener datos guardados del CV del usuario"""
    if not session.get('user_id'):
        return jsonify({'error': 'No autorizado'}), 401
    
    cv_id = request.args.get('cv_id')
    
    try:
        connection = get_db_connection()
        cursor = connection.cursor()
        
        if cv_id:
            # Obtener CV específico
            cursor.execute(
                "SELECT cv_name, personal_info, professional_summary, education, experience, skills, languages, certificates, format_options, ai_methodologies FROM user_cvs WHERE user_id = %s AND id = %s AND is_active = TRUE",
                (session['user_id'], cv_id)
            )
        else:
            # Obtener el CV más reciente
            cursor.execute(
                "SELECT cv_name, personal_info, professional_summary, education, experience, skills, languages, certificates, format_options, ai_methodologies FROM user_cvs WHERE user_id = %s AND is_active = TRUE ORDER BY updated_at DESC LIMIT 1",
                (session['user_id'],)
            )
        
        result = cursor.fetchone()
        cursor.close()
        connection.close()
        
        if result:
            # Función para deserializar JSON de forma segura
            def safe_json_loads(data, default):
                if data is None:
                    return default
                if isinstance(data, str):
                    try:
                        return json.loads(data)
                    except (json.JSONDecodeError, TypeError):
                        return default
                return data if isinstance(data, (dict, list)) else default
            
            # Asegurar que todos los campos JSON sean válidos
            try:
                personal_info = safe_json_loads(result['personal_info'], {})
                education = safe_json_loads(result['education'], [])
                experience = safe_json_loads(result['experience'], [])
                skills = safe_json_loads(result['skills'], [])
                languages = safe_json_loads(result['languages'], [])
                certificates = safe_json_loads(result['certificates'], [])
                format_options = safe_json_loads(result['format_options'], {})
                ai_methodologies = safe_json_loads(result['ai_methodologies'], {})
                
                return jsonify({
                    'success': True,
                    'data': {
                        'cv_name': result['cv_name'] if result['cv_name'] else 'Mi CV',
                        'personal_info': personal_info,
                        'professional_summary': result['professional_summary'] if result['professional_summary'] else '',
                        'education': education,
                        'experience': experience,
                        'skills': skills,
                        'languages': languages,
                        'certificates': certificates,
                        'format_options': format_options,
                        'ai_methodologies': ai_methodologies
                    }
                })
            except Exception as json_error:
                print(f"Error procesando datos JSON: {json_error}")
                return jsonify({
                    'success': True,
                    'data': {
                        'cv_name': result['cv_name'] if result['cv_name'] else 'Mi CV',
                        'personal_info': {},
                        'professional_summary': result['professional_summary'] if result['professional_summary'] else '',
                        'education': [],
                        'experience': [],
                        'skills': [],
                        'languages': [],
                        'certificates': [],
                        'format_options': {},
                        'ai_methodologies': {}
                    }
                })
        else:
            return jsonify({'success': True, 'data': None})
    except Exception as e:
        print(f"Error en get_user_cv_data: {e}")
        return jsonify({'error': str(e)}), 500

def improve_cv_with_ai(cv_data, target_language='es'):
    """Mejorar el CV usando OpenAI en dos etapas: 1) Aplicar metodologías, 2) Traducir todo el CV"""
    print(f"[DEBUG] === INICIANDO PROCESO DE IA COMPLETO ===")
    print(f"[DEBUG] Idioma objetivo: {target_language}")
    print(f"[DEBUG] OPENAI_CLIENT configurado: {OPENAI_CLIENT is not None}")
    print(f"[DEBUG] OPENAI_API_KEY presente: {bool(os.getenv('OPENAI_API_KEY'))}")
    
    if not OPENAI_CLIENT:
        print("[DEBUG] ERROR: OPENAI_CLIENT no está configurado")
        return cv_data
    
    # Mapeo de códigos de idioma a nombres completos
    language_names = {
        'es': 'español',
        'en': 'inglés', 
        'pt': 'portugués',
        'de': 'alemán',
        'fr': 'francés'
    }
    
    target_language_name = language_names.get(target_language, 'español')
    
    # Títulos de secciones por defecto en español
    section_titles = {
        'professional_summary': 'RESUMEN PROFESIONAL',
        'education': 'EDUCACIÓN',
        'experience': 'EXPERIENCIA PROFESIONAL',
        'skills': 'HABILIDADES',
        'languages': 'IDIOMAS',
        'certificates': 'CERTIFICACIONES'
    }
    
    try:
        # ETAPA 1: APLICAR METODOLOGÍAS (en español)
        print(f"[DEBUG] === ETAPA 1: APLICANDO METODOLOGÍAS ===")
        
        format_options = cv_data.get('format_options', {})
        tech_xyz = format_options.get('tech_xyz', False)
        tech_start = format_options.get('tech_start', False)
        selected_skills = cv_data.get('skills', [])
        
        print(f"[DEBUG] Metodologías seleccionadas - XYZ: {tech_xyz}, STAR: {tech_start}")
        print(f"[DEBUG] Habilidades seleccionadas: {selected_skills}")
        
        # Crear una copia de los datos para procesar
        improved_data = cv_data.copy()
        
        professional_summary = cv_data.get('professional_summary', '')
        experience = cv_data.get('experience', [])
        
        # Mejorar resumen profesional con metodología XYZ (en español)
        if professional_summary and tech_xyz:
            print(f"[DEBUG] Aplicando metodología XYZ al resumen profesional")
            print(f"[DEBUG] Resumen original: {professional_summary[:100]}...")
            
            prompt = f"""
            Mejora el siguiente resumen profesional aplicando la metodología XYZ (eXplica lo que hiciste, tus logros cuantificados Y el impacto/resultado obtenido).
            
            Resumen original: {professional_summary}
            
            Habilidades relevantes: {', '.join(selected_skills) if selected_skills else 'No especificadas'}
            
            Instrucciones:
            1. Mantén el tono profesional
            2. Incluye logros cuantificados cuando sea posible
            3. Enfócate en el impacto y resultados
            4. Máximo 150 palabras
            5. Responde SOLO con el resumen mejorado, sin explicaciones adicionales
            6. Mantén el texto en español
            """
            
            try:
                response = OPENAI_CLIENT.chat.completions.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": "Eres un experto en recursos humanos especializado en optimización de CVs. Siempre respondes en español."},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=300,
                    temperature=0.7
                )
                
                improved_summary = response.choices[0].message.content.strip()
                improved_data['professional_summary'] = improved_summary
                print(f"[DEBUG] Resumen mejorado con XYZ: {improved_summary[:100]}...")
            except Exception as openai_error:
                print(f"[DEBUG] Error con OpenAI API en resumen: {str(openai_error)}")
        else:
            print(f"[DEBUG] No se aplica metodología XYZ al resumen (vacío: {not professional_summary}, no seleccionado: {not tech_xyz})")
            

        
        # Mejorar experiencias laborales con metodología STAR (en español)
        print(f"[DEBUG] Procesando {len(experience)} experiencias laborales")
        
        if experience and tech_start:
            print(f"[DEBUG] Aplicando metodología STAR a experiencias laborales")
            
            for i, exp in enumerate(experience):
                if exp.get('description'):
                    print(f"[DEBUG] Procesando experiencia {i+1}: {exp.get('position', 'Sin puesto')} en {exp.get('company', 'Sin empresa')}")
                    
                    prompt = f"""
                    Mejora la siguiente descripción de experiencia laboral aplicando la metodología STAR (Situación, Tarea, Acción, Resultado).
                    
                    Puesto: {exp.get('position', 'No especificado')}
                    Empresa: {exp.get('company', 'No especificada')}
                    Descripción original: {exp['description']}
                    
                    Habilidades relevantes: {', '.join(selected_skills) if selected_skills else 'No especificadas'}
                    
                    Instrucciones:
                    1. Estructura la descripción usando STAR
                    2. Incluye logros cuantificados
                    3. Enfócate en resultados medibles
                    4. Máximo 200 palabras
                    5. Responde SOLO con la descripción mejorada, sin explicaciones adicionales
                    6. Mantén el texto en español
                    """
                    
                    try:
                        response = OPENAI_CLIENT.chat.completions.create(
                            model="gpt-3.5-turbo",
                            messages=[
                                {"role": "system", "content": "Eres un experto en recursos humanos especializado en optimización de CVs. Siempre respondes en español."},
                                {"role": "user", "content": prompt}
                            ],
                            max_tokens=400,
                            temperature=0.7
                        )
                        
                        improved_description = response.choices[0].message.content.strip()
                        improved_data['experience'][i]['description'] = improved_description
                        print(f"[DEBUG] Experiencia {i+1} mejorada con STAR: {improved_description[:100]}...")
                    except Exception as openai_error:
                        print(f"[DEBUG] Error con OpenAI API en experiencia {i+1}: {str(openai_error)}")
                else:
                    print(f"[DEBUG] Experiencia {i+1} sin descripción, saltando")
        else:
            print(f"[DEBUG] No se aplica metodología STAR (sin experiencias: {not experience}, no seleccionado: {not tech_start})")
        
        # ETAPA 2: TRADUCIR TODO EL CV AL IDIOMA SELECCIONADO
        print(f"[DEBUG] === ETAPA 2: TRADUCIENDO TODO EL CV ===")
        print(f"[DEBUG] Idioma objetivo código: {target_language}")
        print(f"[DEBUG] Idioma objetivo nombre: {target_language_name}")
        print(f"[DEBUG] Condición target_language != 'es': {target_language != 'es'}")
        
        if target_language != 'es':
            print(f"[DEBUG] ✅ Iniciando traducción completa al {target_language_name}")
            
            # Traducir información personal (excluyendo nombres)
            if improved_data.get('personal_info'):
                print(f"[DEBUG] Traduciendo información personal (excluyendo nombres)")
                personal_info = improved_data['personal_info']
                
                # Traducir solo el cargo/posición, no el nombre
                if personal_info.get('position'):
                    translate_prompt = f"""
                    Traduce este cargo o profesión al {target_language_name}: {personal_info['position']}
                    
                    Instrucciones:
                    1. Traduce solo el cargo profesional
                    2. Mantén el tono profesional
                    3. Responde SOLO con la traducción del cargo, sin explicaciones adicionales
                    """
                    
                    try:
                        response = OPENAI_CLIENT.chat.completions.create(
                            model="gpt-3.5-turbo",
                            messages=[
                                {"role": "system", "content": f"Eres un traductor profesional especializado en cargos y profesiones. Traduce siempre al {target_language_name}."},
                                {"role": "user", "content": translate_prompt}
                            ],
                            max_tokens=100,
                            temperature=0.3
                        )
                        
                        translated_position = response.choices[0].message.content.strip()
                        improved_data['personal_info']['position'] = translated_position
                        print(f"[DEBUG] Cargo traducido: {translated_position}")
                    except Exception as openai_error:
                        print(f"[DEBUG] Error traduciendo cargo: {str(openai_error)}")
            
            # Traducir resumen profesional
            if improved_data.get('professional_summary'):
                print(f"[DEBUG] Traduciendo resumen profesional")
                translate_prompt = f"""
                Traduce el siguiente resumen profesional al {target_language_name}, manteniendo el tono profesional y la estructura:
                
                {improved_data['professional_summary']}
                
                Instrucciones:
                1. Mantén el tono profesional
                2. Conserva la estructura y metodología aplicada
                3. NO traduzcas nombres propios de personas
                4. Responde SOLO con la traducción, sin explicaciones adicionales
                """
                
                try:
                    response = OPENAI_CLIENT.chat.completions.create(
                        model="gpt-3.5-turbo",
                        messages=[
                            {"role": "system", "content": f"Eres un traductor profesional especializado en CVs y documentos profesionales. Traduce siempre al {target_language_name}."},
                            {"role": "user", "content": translate_prompt}
                        ],
                        max_tokens=400,
                        temperature=0.3
                    )
                    
                    translated_summary = response.choices[0].message.content.strip()
                    improved_data['professional_summary'] = translated_summary
                    print(f"[DEBUG] Resumen traducido: {translated_summary[:100]}...")
                except Exception as openai_error:
                    print(f"[DEBUG] Error traduciendo resumen: {str(openai_error)}")
            
            # Traducir experiencias laborales
            if improved_data.get('experience'):
                print(f"[DEBUG] Traduciendo {len(improved_data['experience'])} experiencias laborales")
                
                for i, exp in enumerate(improved_data['experience']):
                    if exp.get('description'):
                        print(f"[DEBUG] Traduciendo experiencia {i+1}")
                        
                        translate_prompt = f"""
                        Traduce la siguiente descripción de experiencia laboral al {target_language_name}, manteniendo la estructura STAR y el contenido profesional:
                        
                        Puesto: {exp.get('position', '')}
                        Empresa: {exp.get('company', '')}
                        Descripción: {exp['description']}
                        
                        Instrucciones:
                        1. Traduce el puesto de trabajo
                        2. NO traduzcas nombres propios de empresas (mantén el nombre original de la empresa)
                        3. NO traduzcas nombres propios de personas
                        4. Mantén la estructura y metodología aplicada
                        5. Conserva el tono profesional
                        6. Responde en formato JSON con las claves: position, company, description
                        """
                        
                        try:
                            response = OPENAI_CLIENT.chat.completions.create(
                                model="gpt-3.5-turbo",
                                messages=[
                                    {"role": "system", "content": f"Eres un traductor profesional especializado en CVs. Traduce siempre al {target_language_name} y responde en formato JSON."},
                                    {"role": "user", "content": translate_prompt}
                                ],
                                max_tokens=500,
                                temperature=0.3
                            )
                            
                            translated_content = response.choices[0].message.content.strip()
                            
                            # Intentar parsear JSON, si falla usar solo la descripción
                            try:
                                import json
                                translated_exp = json.loads(translated_content)
                                if 'position' in translated_exp:
                                    improved_data['experience'][i]['position'] = translated_exp['position']
                                if 'company' in translated_exp:
                                    improved_data['experience'][i]['company'] = translated_exp['company']
                                if 'description' in translated_exp:
                                    improved_data['experience'][i]['description'] = translated_exp['description']
                            except:
                                # Si no es JSON válido, usar como descripción directa
                                improved_data['experience'][i]['description'] = translated_content
                            
                            print(f"[DEBUG] Experiencia {i+1} traducida")
                        except Exception as openai_error:
                            print(f"[DEBUG] Error traduciendo experiencia {i+1}: {str(openai_error)}")
            
            # Traducir otros campos usando la función existente
            improved_data = translate_cv_fields(improved_data, target_language_name)
            
            # Traducir títulos de secciones
            print(f"[DEBUG] Traduciendo títulos de secciones al {target_language_name}")
            for section_key, spanish_title in section_titles.items():
                translate_prompt = f"""
                Traduce este título de sección de CV al {target_language_name}: {spanish_title}
                
                Instrucciones:
                1. Traduce el título manteniendo el formato profesional
                2. Usa mayúsculas como en el original
                3. Responde SOLO con la traducción del título, sin explicaciones adicionales
                """
                
                try:
                    response = OPENAI_CLIENT.chat.completions.create(
                        model="gpt-3.5-turbo",
                        messages=[
                            {"role": "system", "content": f"Eres un traductor profesional especializado en CVs. Traduce siempre al {target_language_name}."},
                            {"role": "user", "content": translate_prompt}
                        ],
                        max_tokens=50,
                        temperature=0.3
                    )
                    
                    translated_title = response.choices[0].message.content.strip()
                    section_titles[section_key] = translated_title
                    print(f"[DEBUG] Título traducido: {spanish_title} -> {translated_title}")
                except Exception as openai_error:
                    print(f"[DEBUG] Error traduciendo título {section_key}: {str(openai_error)}")
            
            # Agregar títulos traducidos a los datos mejorados
            improved_data['section_titles'] = section_titles
            
        else:
            print(f"[DEBUG] ❌ No se requiere traducción (idioma español o código: {target_language})")
            # Agregar títulos en español por defecto
            improved_data['section_titles'] = section_titles
        
        print(f"[DEBUG] === PROCESO COMPLETO FINALIZADO ===")
        return improved_data
        
    except Exception as e:
        print(f"Error mejorando CV con IA: {str(e)}")
        return cv_data

def translate_cv_fields(cv_data, target_language_name):
    """Traducir campos adicionales del CV que no se procesan con IA"""
    if not OPENAI_CLIENT or target_language_name == 'español':
        return cv_data
    
    try:
        # Traducir educación (excluyendo nombres de instituciones)
        education = cv_data.get('education', [])
        if education:
            for edu in education:
                # Traducir título académico (degree)
                if edu.get('degree'):
                    translate_prompt = f"""
                    Traduce este título académico al {target_language_name}: {edu['degree']}
                    
                    Instrucciones:
                    1. Traduce solo el título académico o grado
                    2. NO traduzcas nombres propios de instituciones educativas
                    3. NO traduzcas nombres propios de personas
                    4. Responde solo con la traducción del título académico
                    """
                    try:
                        response = OPENAI_CLIENT.chat.completions.create(
                            model="gpt-3.5-turbo",
                            messages=[
                                {"role": "system", "content": f"Eres un traductor profesional especializado en CVs y documentos profesionales. Traduce al {target_language_name} pero NO traduzcas nombres propios."},
                                {"role": "user", "content": translate_prompt}
                            ],
                            max_tokens=100,
                            temperature=0.3
                        )
                        edu['degree'] = response.choices[0].message.content.strip()
                    except Exception:
                        pass  # Mantener original si hay error
                
                # Traducir carrera (career)
                if edu.get('career'):
                    translate_prompt = f"""
                    Traduce este nombre de carrera al {target_language_name}: {edu['career']}
                    
                    Instrucciones:
                    1. Traduce solo el nombre de la carrera o programa académico
                    2. NO traduzcas nombres propios de instituciones educativas
                    3. NO traduzcas nombres propios de personas
                    4. Responde solo con la traducción del nombre de la carrera
                    """
                    try:
                        response = OPENAI_CLIENT.chat.completions.create(
                            model="gpt-3.5-turbo",
                            messages=[
                                {"role": "system", "content": f"Eres un traductor profesional especializado en CVs y documentos profesionales. Traduce al {target_language_name} pero NO traduzcas nombres propios."},
                                {"role": "user", "content": translate_prompt}
                            ],
                            max_tokens=100,
                            temperature=0.3
                        )
                        edu['career'] = response.choices[0].message.content.strip()
                        print(f"[DEBUG] Carrera traducida: {edu['career']}")
                    except Exception:
                        pass  # Mantener original si hay error
        
        # Traducir certificados (incluyendo instituciones)
        certificates = cv_data.get('certificates', [])
        if certificates:
            for cert in certificates:
                # Traducir nombre del certificado
                if cert.get('title'):
                    translate_prompt = f"""
                    Traduce este nombre de certificado al {target_language_name}: {cert['title']}
                    
                    Instrucciones:
                    1. Detecta palabras en español que necesiten traducción (como 'Avanzado', 'Intermedio', 'Básico', etc.)
                    2. Traduce las palabras en español al {target_language_name} de forma profesional para un CV
                    3. Mantén nombres propios de software, marcas o tecnologías (Excel, Power BI, SAP, etc.)
                    4. Mantén el formato profesional apropiado para un CV
                    5. Responde SOLO con la traducción completa del nombre del certificado
                    """
                    try:
                        response = OPENAI_CLIENT.chat.completions.create(
                            model="gpt-3.5-turbo",
                            messages=[
                                {"role": "system", "content": f"Eres un traductor profesional especializado en CVs y documentos profesionales. Detecta palabras en español y tradúcelas al {target_language_name} manteniendo nombres propios de tecnologías."},
                                {"role": "user", "content": translate_prompt}
                            ],
                            max_tokens=100,
                            temperature=0.3
                        )
                        translated_name = response.choices[0].message.content.strip()
                        
                        # Limpiar formato no deseado (guiones, flechas, texto duplicado)
                        if ' - ' in translated_name:
                            # Si contiene guión, tomar solo la parte después del guión
                            translated_name = translated_name.split(' - ')[-1].strip()
                        if ' -> ' in translated_name:
                            # Si contiene flecha, tomar solo la parte después de la flecha
                            translated_name = translated_name.split(' -> ')[-1].strip()
                        
                        cert['title'] = translated_name
                        print(f"[DEBUG] Certificado traducido: {cert['title']}")
                    except Exception:
                        pass  # Mantener original si hay error
                
                # Traducir institución del certificado
                if cert.get('institution'):
                    translate_prompt = f"""
                    Traduce este nombre de institución al {target_language_name}: {cert['institution']}
                    
                    Instrucciones:
                    1. Traduce el nombre de la institución si es un nombre genérico
                    2. Si es un nombre propio específico de una empresa/universidad, manténlo igual
                    3. Responde solo con la traducción o el nombre original
                    """
                    try:
                        response = OPENAI_CLIENT.chat.completions.create(
                            model="gpt-3.5-turbo",
                            messages=[
                                {"role": "system", "content": f"Eres un traductor profesional especializado en CVs y documentos profesionales. Traduce al {target_language_name}."},
                                {"role": "user", "content": translate_prompt}
                            ],
                            max_tokens=100,
                            temperature=0.3
                        )
                        translated_institution = response.choices[0].message.content.strip()
                        
                        # Limpiar formato no deseado (guiones, flechas, texto duplicado)
                        if ' - ' in translated_institution:
                            # Si contiene guión, tomar solo la parte después del guión
                            translated_institution = translated_institution.split(' - ')[-1].strip()
                        if ' -> ' in translated_institution:
                            # Si contiene flecha, tomar solo la parte después de la flecha
                            translated_institution = translated_institution.split(' -> ')[-1].strip()
                        
                        cert['institution'] = translated_institution
                        print(f"[DEBUG] Institución del certificado traducida: {cert['institution']}")
                    except Exception:
                        pass  # Mantener original si hay error
        
        # Traducir idiomas y niveles de idioma
        languages = cv_data.get('languages', [])
        print(f"[DEBUG] Datos de idiomas antes de traducir: {languages}")
        if languages:
            for lang in languages:
                print(f"[DEBUG] Procesando idioma: {lang}")
                # Traducir nombre del idioma
                if lang.get('language'):
                    translate_prompt = f"""
                    Traduce SOLO el nombre de este idioma al {target_language_name}: {lang['language']}
                    
                    Instrucciones:
                    1. Traduce únicamente el nombre del idioma (Inglés, Francés, Alemán, Español, etc.)
                    2. NO incluyas niveles de competencia en la respuesta
                    3. Responde SOLO con el nombre del idioma traducido
                    
                    Ejemplos:
                    - Si el idioma es "Inglés" y traduces al francés, responde: "Anglais"
                    - Si el idioma es "Francés" y traduces al inglés, responde: "French"
                    """
                    try:
                        response = OPENAI_CLIENT.chat.completions.create(
                            model="gpt-3.5-turbo",
                            messages=[
                                {"role": "system", "content": f"Eres un traductor profesional especializado en CVs. Traduce ÚNICAMENTE nombres de idiomas al {target_language_name}, sin incluir niveles."},
                                {"role": "user", "content": translate_prompt}
                            ],
                            max_tokens=30,
                            temperature=0.1
                        )
                        translated_language = response.choices[0].message.content.strip()
                        
                        # Limpiar formato no deseado (guiones, flechas, texto duplicado)
                        if ' - ' in translated_language:
                            # Si contiene guión, tomar solo la parte después del guión
                            translated_language = translated_language.split(' - ')[-1].strip()
                        if ' -> ' in translated_language:
                            # Si contiene flecha, tomar solo la parte después de la flecha
                            translated_language = translated_language.split(' -> ')[-1].strip()
                        
                        lang['language'] = translated_language
                        print(f"[DEBUG] Idioma traducido: {lang['language']}")
                    except Exception:
                        pass  # Mantener original si hay error
                
                # Traducir nivel del idioma
                if lang.get('level'):
                    translate_prompt = f"""
                    Traduce SOLO este nivel de competencia de idioma al {target_language_name}: {lang['level']}
                    
                    Instrucciones:
                    1. Traduce únicamente el nivel de competencia (Básico, Intermedio, Avanzado, Nativo, etc.)
                    2. NO incluyas nombres de idiomas en la respuesta
                    3. Responde SOLO con el nivel traducido
                    
                    Ejemplos:
                    - Si el nivel es "Intermedio" y traduces al francés, responde: "Intermédiaire"
                    - Si el nivel es "Avanzado" y traduces al inglés, responde: "Advanced"
                    """
                    try:
                        response = OPENAI_CLIENT.chat.completions.create(
                            model="gpt-3.5-turbo",
                            messages=[
                                {"role": "system", "content": f"Eres un traductor profesional especializado en CVs. Traduce ÚNICAMENTE niveles de competencia al {target_language_name}, sin incluir nombres de idiomas."},
                                {"role": "user", "content": translate_prompt}
                            ],
                            max_tokens=30,
                            temperature=0.1
                        )
                        translated_level = response.choices[0].message.content.strip()
                        
                        # Limpiar formato no deseado (guiones, flechas, texto duplicado)
                        if ' - ' in translated_level:
                            # Si contiene guión, tomar solo la parte después del guión
                            translated_level = translated_level.split(' - ')[-1].strip()
                        if ' -> ' in translated_level:
                            # Si contiene flecha, tomar solo la parte después de la flecha
                            translated_level = translated_level.split(' -> ')[-1].strip()
                        
                        lang['level'] = translated_level
                        print(f"[DEBUG] Nivel de idioma traducido: {lang['level']}")
                    except Exception:
                        pass  # Mantener original si hay error
        
        print(f"[DEBUG] Datos de idiomas después de traducir: {languages}")
        
        # Traducir competencias/habilidades
        skills = cv_data.get('skills', [])
        if skills:
            translated_skills = []
            for skill in skills:
                if isinstance(skill, str) and skill.strip():
                    translate_prompt = f"""
                    Traduce esta competencia o habilidad profesional al {target_language_name}: {skill}
                    
                    Instrucciones CRÍTICAS:
                    1. Detecta si la palabra está en español y necesita traducción
                    2. Si está en español, tradúcela al {target_language_name} de forma profesional para un CV
                    3. Si ya está en el idioma objetivo o es un término técnico universal, manténla igual
                    4. Mantén el formato profesional apropiado para un CV
                    5. Responde ÚNICAMENTE con la traducción final, SIN incluir el texto original
                    6. NO uses guiones, flechas o separadores entre el original y la traducción
                    7. SOLO devuelve el resultado final traducido
                    
                    EJEMPLO CORRECTO:
                    - Input: "gestión de proyectos"
                    - Output: "project management" (NO "gestión de proyectos - project management")
                    """
                    try:
                        response = OPENAI_CLIENT.chat.completions.create(
                            model="gpt-3.5-turbo",
                            messages=[
                                {"role": "system", "content": f"Eres un traductor profesional. SOLO devuelve la traducción final al {target_language_name}, sin incluir el texto original, guiones, flechas o explicaciones."},
                                {"role": "user", "content": translate_prompt}
                            ],
                            max_tokens=50,
                            temperature=0.1
                        )
                        translated_skill = response.choices[0].message.content.strip()
                        
                        # Limpiar formato no deseado (guiones, flechas, texto duplicado)
                        if ' - ' in translated_skill:
                            # Si contiene guión, tomar solo la parte después del guión
                            translated_skill = translated_skill.split(' - ')[-1].strip()
                        if ' -> ' in translated_skill:
                            # Si contiene flecha, tomar solo la parte después de la flecha
                            translated_skill = translated_skill.split(' -> ')[-1].strip()
                        
                        translated_skills.append(translated_skill)
                        print(f"[DEBUG] Competencia traducida: {skill} -> {translated_skill}")
                    except Exception as e:
                        print(f"[DEBUG] Error traduciendo competencia '{skill}': {str(e)}")
                        translated_skills.append(skill)  # Mantener original si hay error
                else:
                    translated_skills.append(skill)  # Mantener si no es string válido
            
            cv_data['skills'] = translated_skills
        
        # Traducir fechas 'Presente' en experiencia laboral
        experience = cv_data.get('experience', [])
        if experience:
            # Obtener traducción de 'Presente'
            present_translation = 'Presente'  # Por defecto
            try:
                translate_prompt = f"""
                Traduce la palabra 'Presente' (como se usa en fechas de trabajo actual) al {target_language_name}.
                
                Instrucciones:
                1. Traduce solo la palabra que indica trabajo actual
                2. Responde SOLO con la traducción, sin explicaciones
                """
                response = OPENAI_CLIENT.chat.completions.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": f"Eres un traductor profesional especializado en CVs y documentos profesionales. Traduce al {target_language_name}."},
                        {"role": "user", "content": translate_prompt}
                    ],
                    max_tokens=20,
                    temperature=0.3
                )
                present_translation = response.choices[0].message.content.strip()
                print(f"[DEBUG] 'Presente' traducido a: {present_translation}")
            except Exception as e:
                print(f"[DEBUG] Error traduciendo 'Presente': {str(e)}")
            
            # Aplicar traducción a las fechas de experiencia
            for exp in experience:
                if exp.get('end_date') == 'Presente':
                    exp['end_date'] = present_translation
                    print(f"[DEBUG] Fecha actualizada: Presente -> {present_translation}")
        
        return cv_data
        
    except Exception as e:
        print(f"Error traduciendo campos del CV: {str(e)}")
        return cv_data

@app.route('/save_cv_draft', methods=['POST'])
def save_cv_draft():
    """Guardar borrador de CV sin validación estricta"""
    try:
        print(f"[DEBUG] save_cv_draft iniciado - user_id: {session.get('user_id')}")
        
        if not session.get('user_id'):
            print("[DEBUG] Error: No autorizado")
            return jsonify({'error': 'No autorizado'}), 401
        
        data = request.get_json()
        print(f"[DEBUG] Datos recibidos: {data is not None}")
        
        if not data:
            print("[DEBUG] Error: No se recibieron datos")
            return jsonify({'error': 'No se recibieron datos'}), 400
    except Exception as e:
        print(f"[DEBUG] Error en validación inicial: {str(e)}")
        return jsonify({'error': f'Error en validación inicial: {str(e)}'}), 500
    
    # Validar que las opciones de formato estén presentes
    if 'format_options' not in data:
        data['format_options'] = {'format': 'hardware', 'tech_xyz': False, 'tech_start': False}
    
    cv_id = data.get('cv_id')
    cv_name = data.get('cv_name', 'Mi CV')
    print(f"[DEBUG] cv_id: {cv_id}, cv_name: {cv_name}")
    
    # Guardar en la base de datos sin validación estricta
    try:
        print("[DEBUG] Intentando conectar a la base de datos")
        connection = get_db_connection()
        cursor = connection.cursor()
        print("[DEBUG] Conexión a base de datos exitosa")
        
        if cv_id:
            # Actualizar CV existente
            print(f"[DEBUG] Actualizando CV existente con ID: {cv_id}")
            cursor.execute(
                """
                UPDATE user_cvs SET 
                    cv_name = %s,
                    personal_info = %s,
                    professional_summary = %s,
                    education = %s,
                    experience = %s,
                    skills = %s,
                    languages = %s,
                    certificates = %s,
                    format_options = %s,
                    ai_methodologies = %s,
                    updated_at = CURRENT_TIMESTAMP
                WHERE id = %s AND user_id = %s
                """,
                (
                    cv_name,
                    json.dumps(data.get('personal_info', {})),
                    data.get('professional_summary', ''),
                    json.dumps(data.get('education', [])),
                    json.dumps(data.get('experience', [])),
                    json.dumps(data.get('skills', [])),
                    json.dumps(data.get('languages', [])),
                    json.dumps(data.get('certificates', [])),
                    json.dumps(data.get('format_options', {})),
                    json.dumps(data.get('ai_methodologies', {})),
                    cv_id,
                    session['user_id']
                )
            )
            print("[DEBUG] UPDATE ejecutado exitosamente")
        else:
            # Verificar límite de 10 CVs
            print("[DEBUG] Creando nuevo CV - verificando límite")
            cursor.execute(
                "SELECT COUNT(*) FROM user_cvs WHERE user_id = %s AND is_active = TRUE",
                (session['user_id'],)
            )
            cv_count = cursor.fetchone()['count']
            print(f"[DEBUG] CVs actuales del usuario: {cv_count}")
            
            if cv_count >= 10:
                print("[DEBUG] Límite de CVs alcanzado")
                return jsonify({'error': 'Has alcanzado el límite máximo de 10 CVs. Elimina uno para crear otro.'}), 400
            
            # Verificar si ya existe un CV con el mismo nombre
            print("[DEBUG] Verificando si existe CV con el mismo nombre")
            cursor.execute(
                "SELECT id FROM user_cvs WHERE user_id = %s AND cv_name = %s AND is_active = TRUE",
                (session['user_id'], cv_name)
            )
            existing_cv = cursor.fetchone()
            
            if existing_cv:
                # Actualizar CV existente
                existing_cv_id = existing_cv['id']
                print(f"[DEBUG] CV existente encontrado con ID: {existing_cv_id}, actualizando...")
                cursor.execute(
                    """
                    UPDATE user_cvs SET 
                        personal_info = %s,
                        professional_summary = %s,
                        education = %s,
                        experience = %s,
                        skills = %s,
                        languages = %s,
                        certificates = %s,
                        format_options = %s,
                        ai_methodologies = %s,
                        updated_at = CURRENT_TIMESTAMP
                    WHERE id = %s AND user_id = %s
                    """,
                    (
                        json.dumps(data.get('personal_info', {})),
                        data.get('professional_summary', ''),
                        json.dumps(data.get('education', [])),
                        json.dumps(data.get('experience', [])),
                        json.dumps(data.get('skills', [])),
                        json.dumps(data.get('languages', [])),
                        json.dumps(data.get('certificates', [])),
                        json.dumps(data.get('format_options', {})),
                        json.dumps(data.get('ai_methodologies', {})),
                        existing_cv_id,
                        session['user_id']
                    )
                )
                new_cv_id = existing_cv_id
                print(f"[DEBUG] CV existente actualizado con ID: {new_cv_id}")
            else:
                # Crear nuevo CV
                print("[DEBUG] Insertando nuevo CV")
                cursor.execute(
                    """
                    INSERT INTO user_cvs (user_id, cv_name, personal_info, professional_summary, education, experience, skills, languages, certificates, format_options, ai_methodologies)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                    RETURNING id
                    """,
                    (
                        session['user_id'],
                        cv_name,
                        json.dumps(data.get('personal_info', {})),
                        data.get('professional_summary', ''),
                        json.dumps(data.get('education', [])),
                        json.dumps(data.get('experience', [])),
                        json.dumps(data.get('skills', [])),
                        json.dumps(data.get('languages', [])),
                        json.dumps(data.get('certificates', [])),
                        json.dumps(data.get('format_options', {})),
                        json.dumps(data.get('ai_methodologies', {}))
                    )
                )
                new_cv_id = cursor.fetchone()['id']
                print(f"[DEBUG] Nuevo CV creado con ID: {new_cv_id}")
        
        print("[DEBUG] Haciendo commit a la base de datos")
        connection.commit()
        cursor.close()
        connection.close()
        print("[DEBUG] Conexión cerrada exitosamente")
        
        response_data = {
            'success': True, 
            'message': 'Borrador guardado exitosamente'
        }
        
        if not cv_id:  # Si es un nuevo CV, devolver el ID
            response_data['cv_id'] = new_cv_id
            
        print(f"[DEBUG] Respuesta exitosa: {response_data}")
        return jsonify(response_data)
    except Exception as e:
        print(f"[DEBUG] Error en save_cv_draft: {str(e)}")
        print(f"[DEBUG] Tipo de error: {type(e).__name__}")
        import traceback
        print(f"[DEBUG] Traceback: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500

@app.route('/my_cvs')
def my_cvs():
    """Página para gestionar múltiples CVs del usuario"""
    if 'user_id' not in session:
        return redirect(url_for('login'))
    return render_template('my_cvs.html')

@app.route('/get_user_cvs', methods=['GET'])
def get_user_cvs():
    """Obtener lista de CVs del usuario"""
    if not session.get('user_id'):
        return jsonify({'error': 'No autorizado'}), 401
    
    try:
        connection = get_db_connection()
        cursor = connection.cursor()
        cursor.execute(
            "SELECT id, cv_name, created_at, updated_at FROM user_cvs WHERE user_id = %s AND is_active = TRUE ORDER BY updated_at DESC",
            (session['user_id'],)
        )
        cvs = cursor.fetchall()
        cursor.close()
        connection.close()
        
        cv_list = []
        for cv in cvs:
            cv_list.append({
                'id': cv['id'],
                'name': cv['cv_name'],
                'created_at': cv['created_at'].strftime('%d/%m/%Y %H:%M') if cv['created_at'] else '',
                'updated_at': cv['updated_at'].strftime('%d/%m/%Y %H:%M') if cv['updated_at'] else ''
            })
        
        return jsonify({'cvs': cv_list})
    except Exception as e:
        print(f"[ERROR] Error en save_cv: {str(e)}")
        print(f"[ERROR] Tipo de error: {type(e).__name__}")
        import traceback
        print(f"[ERROR] Traceback: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500

@app.route('/delete_cv/<int:cv_id>', methods=['DELETE'])
def delete_cv(cv_id):
    """Eliminar un CV específico"""
    if not session.get('user_id'):
        return jsonify({'error': 'No autorizado'}), 401
    
    try:
        connection = get_db_connection()
        cursor = connection.cursor()
        
        # Verificar que el CV pertenece al usuario
        cursor.execute(
            "SELECT id FROM user_cvs WHERE id = %s AND user_id = %s",
            (cv_id, session['user_id'])
        )
        
        if not cursor.fetchone():
            return jsonify({'error': 'CV no encontrado'}), 404
        
        # Marcar como inactivo en lugar de eliminar
        cursor.execute(
            "UPDATE user_cvs SET is_active = FALSE WHERE id = %s AND user_id = %s",
            (cv_id, session['user_id'])
        )
        
        connection.commit()
        cursor.close()
        connection.close()
        
        return jsonify({'success': True, 'message': 'CV eliminado exitosamente'})
    except Exception as e:
        print(f"[ERROR] Error en save_cv: {str(e)}")
        print(f"[ERROR] Tipo de error: {type(e).__name__}")
        import traceback
        print(f"[ERROR] Traceback: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500

@app.route('/duplicate_cv/<int:cv_id>', methods=['POST'])
def duplicate_cv(cv_id):
    """Duplicar un CV existente"""
    if not session.get('user_id'):
        return jsonify({'error': 'No autorizado'}), 401
    
    try:
        connection = get_db_connection()
        cursor = connection.cursor()
        
        # Verificar límite de 10 CVs
        cursor.execute(
            "SELECT COUNT(*) FROM user_cvs WHERE user_id = %s AND is_active = TRUE",
            (session['user_id'],)
        )
        cv_count = cursor.fetchone()['count']
        
        if cv_count >= 10:
            return jsonify({'error': 'Has alcanzado el límite máximo de 10 CVs. Elimina uno para crear otro.'}), 400
        
        # Obtener datos del CV original
        cursor.execute(
            "SELECT cv_name, personal_info, professional_summary, education, experience, skills, languages, certificates, format_options FROM user_cvs WHERE id = %s AND user_id = %s AND is_active = TRUE",
            (cv_id, session['user_id'])
        )
        
        original_cv = cursor.fetchone()
        if not original_cv:
            return jsonify({'error': 'CV no encontrado'}), 404
        
        # Crear copia con nombre único
        new_name = f"{original_cv['cv_name']} - Copia"
        counter = 1
        while True:
            cursor.execute(
                "SELECT id FROM user_cvs WHERE user_id = %s AND cv_name = %s AND is_active = TRUE",
                (session['user_id'], new_name)
            )
            if not cursor.fetchone():
                break
            counter += 1
            new_name = f"{original_cv['cv_name']} - Copia {counter}"
        
        # Insertar CV duplicado
        cursor.execute(
            """
            INSERT INTO user_cvs (user_id, cv_name, personal_info, professional_summary, education, experience, skills, languages, certificates, format_options)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            RETURNING id
            """,
            (
                session['user_id'],
                new_name,
                original_cv['personal_info'],
                original_cv['professional_summary'],
                original_cv['education'],
                original_cv['experience'],
                original_cv['skills'],
                original_cv['languages'],
                original_cv['certificates'],
                original_cv['format_options']
            )
        )
        
        new_cv_id = cursor.fetchone()['id']
        
        connection.commit()
        cursor.close()
        connection.close()
        
        return jsonify({
            'success': True, 
            'message': 'CV duplicado exitosamente',
            'new_cv_id': new_cv_id,
            'new_cv_name': new_name
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/rename_cv/<int:cv_id>', methods=['PUT'])
def rename_cv(cv_id):
    """Renombrar un CV específico"""
    if not session.get('user_id'):
        return jsonify({'error': 'No autorizado'}), 401
    
    try:
        data = request.get_json()
        new_name = data.get('name', '').strip()
        
        if not new_name:
            return jsonify({'error': 'El nombre no puede estar vacío'}), 400
        
        if len(new_name) > 100:
            return jsonify({'error': 'El nombre no puede exceder 100 caracteres'}), 400
        
        connection = get_db_connection()
        cursor = connection.cursor()
        
        # Verificar que el CV pertenece al usuario
        cursor.execute(
            "SELECT id FROM user_cvs WHERE id = %s AND user_id = %s AND is_active = TRUE",
            (cv_id, session['user_id'])
        )
        
        if not cursor.fetchone():
            cursor.close()
            connection.close()
            return jsonify({'error': 'CV no encontrado'}), 404
        
        # Verificar que no existe otro CV con el mismo nombre
        cursor.execute(
            "SELECT id FROM user_cvs WHERE user_id = %s AND cv_name = %s AND id != %s AND is_active = TRUE",
            (session['user_id'], new_name, cv_id)
        )
        
        if cursor.fetchone():
            cursor.close()
            connection.close()
            return jsonify({'error': 'Ya existe un CV con ese nombre'}), 400
        
        # Actualizar el nombre
        cursor.execute(
            "UPDATE user_cvs SET cv_name = %s, updated_at = CURRENT_TIMESTAMP WHERE id = %s AND user_id = %s",
            (new_name, cv_id, session['user_id'])
        )
        
        connection.commit()
        cursor.close()
        connection.close()
        
        return jsonify({'success': True, 'message': 'CV renombrado exitosamente'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/preview_cv/<int:cv_id>')
def preview_cv(cv_id):
    """Generar preview HTML del CV para mostrar en modal"""
    if not session.get('user_id'):
        return redirect(url_for('login'))
    
    try:
        connection = get_db_connection()
        cursor = connection.cursor()
        
        # Obtener los datos del CV
        cursor.execute(
            "SELECT cv_name, personal_info, professional_summary, education, experience, skills, languages, certificates, format_options, ai_methodologies FROM user_cvs WHERE id = %s AND user_id = %s AND is_active = TRUE",
            (cv_id, session['user_id'])
        )
        result = cursor.fetchone()
        
        if not result:
            return '<div class="alert alert-danger m-3">CV no encontrado</div>', 404
        
        # Reconstruir los datos del CV - deserializar JSON strings
        def safe_json_loads(data, default):
            if data is None:
                return default
            if isinstance(data, str):
                try:
                    return json.loads(data)
                except (json.JSONDecodeError, TypeError):
                    return default
            return data if isinstance(data, (dict, list)) else default
        
        cv_data = {
            'cv_name': result['cv_name'] if result['cv_name'] else 'Mi CV',
            'personal_info': safe_json_loads(result['personal_info'], {}),
            'professional_summary': result['professional_summary'] if result['professional_summary'] else '',
            'education': safe_json_loads(result['education'], []),
            'experience': safe_json_loads(result['experience'], []),
            'skills': safe_json_loads(result['skills'], []),
            'languages': safe_json_loads(result['languages'], []),
            'certificates': safe_json_loads(result['certificates'], []),
            'format_options': safe_json_loads(result['format_options'], {'format': 'hardware', 'tech_xyz': False, 'tech_start': False}),
            'ai_methodologies': safe_json_loads(result['ai_methodologies'], {})
        }
        
        # Obtener idioma objetivo desde ai_methodologies (por defecto español)
        ai_methodologies = cv_data.get('ai_methodologies', {})
        target_language = ai_methodologies.get('target_language', 'es')
        
        # Aplicar mejoras de IA antes de generar HTML (igual que en export_cv)
        improved_cv_data = improve_cv_with_ai(cv_data, target_language)
        
        # Generar el HTML del CV con los datos mejorados por IA
        cv_html = generate_cv_html(improved_cv_data)
        
        # Crear HTML optimizado para preview con estilos responsivos
        preview_html = f"""
        <div style="background: white; padding: 20px; font-family: 'Times New Roman', serif; line-height: 1.4; color: #000; max-width: 100%; overflow-x: auto;">
            {cv_html}
        </div>
        """
        
        return preview_html
        
    except Exception as e:
        print(f"Error en preview_cv: {e}")
        return '<div class="alert alert-danger m-3">Error al generar el preview</div>', 500
    finally:
        if 'cursor' in locals():
            cursor.close()
        if 'connection' in locals():
            connection.close()

@app.route('/generate_pdf/<int:cv_id>')
def generate_pdf_route(cv_id):
    """Generar y descargar PDF del CV"""
    if not session.get('user_id'):
        return redirect(url_for('login'))
    
    # Redirigir a la función export_cv existente
    return export_cv(cv_id)

@app.route('/save_cv', methods=['POST'])
def save_cv():
    if not session.get('user_id'):
        return redirect(url_for('login'))
    
    data = request.get_json()
    
    # Validar datos requeridos
    if not data or not all(key in data for key in ['personal_info', 'education', 'experience']):
        return jsonify({'error': 'Faltan datos requeridos'}), 400
    
    # Validar que las opciones de formato estén presentes
    if 'format_options' not in data:
        data['format_options'] = {'format': 'hardware', 'tech_xyz': False, 'tech_start': False}
    
    # Obtener idioma objetivo (por defecto español)
    target_language = data.get('target_language', 'es')
    
    # Mejorar CV con IA antes de generar HTML
    print(f"[DEBUG] Datos originales antes de IA - resumen: {data.get('professional_summary', 'VACIO')[:50]}...")
    print(f"[DEBUG] Idioma objetivo: {target_language}")
    print(f"[DEBUG] Opciones de formato: {data.get('format_options', {})}")
    
    improved_data = improve_cv_with_ai(data, target_language)
    
    print(f"[DEBUG] Datos mejorados después de IA - resumen: {improved_data.get('professional_summary', 'VACIO')[:50]}...")
    print(f"[DEBUG] ¿Los datos cambiaron? {data.get('professional_summary') != improved_data.get('professional_summary')}")
    
    # Generar HTML del CV con los datos mejorados
    cv_html = generate_cv_html(improved_data)
    
    # Guardar en la base de datos
    try:
        connection = get_db_connection()
        cursor = connection.cursor()
        
        # Guardar una referencia mínima del CV (sin el HTML completo para optimizar almacenamiento)
        cursor.execute(
            "INSERT INTO resumes (user_id, filename, content) VALUES (%s, %s, %s) RETURNING id",
            (session['user_id'], data['personal_info'].get('name', 'Mi CV'), "CV creado por el usuario - datos estructurados en user_cv_data")
        )
        cv_id = cursor.fetchone()['id']
        
        # Verificar si hay un cv_id en los datos (para actualizar CV existente)
        cv_data_id = data.get('cv_id')
        
        if cv_data_id:
            # Actualizar CV existente
            cursor.execute(
                """
                UPDATE user_cvs SET
                    cv_name = %s,
                    personal_info = %s,
                    professional_summary = %s,
                    education = %s,
                    experience = %s,
                    skills = %s,
                    languages = %s,
                    certificates = %s,
                    format_options = %s,
                    ai_methodologies = %s,
                    updated_at = CURRENT_TIMESTAMP
                WHERE id = %s AND user_id = %s AND is_active = TRUE
                """,
                (
                    improved_data.get('cv_name', 'Mi CV'),
                    json.dumps(improved_data.get('personal_info', {})),
                    improved_data.get('professional_summary', ''),
                    json.dumps(improved_data.get('education', [])),
                    json.dumps(improved_data.get('experience', [])),
                    json.dumps(improved_data.get('skills', [])),
                    json.dumps(improved_data.get('languages', [])),
                    json.dumps(improved_data.get('certificates', [])),
                    json.dumps(improved_data.get('format_options', {})),
                    json.dumps({**improved_data.get('ai_methodologies', {}), 'target_language': target_language}),
                    cv_data_id,
                    session['user_id']
                )
            )
            new_cv_id = cv_data_id
        else:
            # Verificar límite de CVs (máximo 10)
            cursor.execute(
                "SELECT COUNT(*) FROM user_cvs WHERE user_id = %s AND is_active = TRUE",
                (session['user_id'],)
            )
            cv_count = cursor.fetchone()['count']
            
            if cv_count >= 10:
                return jsonify({'error': 'Has alcanzado el límite máximo de 10 CVs. Elimina alguno para crear uno nuevo.'}), 400
            
            # Crear nuevo CV
            # Generar nombre único para el CV
            base_cv_name = improved_data.get('cv_name', 'Mi CV')
            cv_name = base_cv_name
            counter = 1
            
            # Verificar si ya existe un CV con ese nombre
            while True:
                cursor.execute(
                    "SELECT id FROM user_cvs WHERE user_id = %s AND cv_name = %s AND is_active = TRUE",
                    (session['user_id'], cv_name)
                )
                if not cursor.fetchone():
                    break
                cv_name = f"{base_cv_name} ({counter})"
                counter += 1
            
            cursor.execute(
                """
                INSERT INTO user_cvs (user_id, cv_name, personal_info, professional_summary, education, experience, skills, languages, certificates, format_options, ai_methodologies)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s) RETURNING id
                """,
                (
                    session['user_id'],
                    cv_name,
                    json.dumps(improved_data.get('personal_info', {})),
                    improved_data.get('professional_summary', ''),
                    json.dumps(improved_data.get('education', [])),
                    json.dumps(improved_data.get('experience', [])),
                    json.dumps(improved_data.get('skills', [])),
                    json.dumps(improved_data.get('languages', [])),
                    json.dumps(improved_data.get('certificates', [])),
                    json.dumps(improved_data.get('format_options', {})),
                    json.dumps({**improved_data.get('ai_methodologies', {}), 'target_language': target_language})
                )
            )
            new_cv_id = cursor.fetchone()['id']
        
        connection.commit()
        cursor.close()
        connection.close()
        
        # Incrementar contador de uso para creación de CV
        increment_usage(session.get('user_id'), 'cv_creation')
        
        # Crear mensaje de éxito con idioma
        language_names = {
            'es': 'español',
            'en': 'inglés', 
            'pt': 'portugués',
            'de': 'alemán',
            'fr': 'francés'
        }
        language_name = language_names.get(target_language, 'español')
        success_message = f'CV guardado y mejorado con IA en {language_name}. Datos del usuario actualizados.'
        
        return jsonify({
            'success': True, 
            'cv_id': new_cv_id, 
            'message': success_message,
            'improved_data': improved_data,
            'cv_html': cv_html,
            'target_language': target_language
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/export_cv/<int:cv_id>')
def export_cv(cv_id):
    if not session.get('user_id'):
        return redirect(url_for('login'))
    
    # Obtener los datos del CV de la base de datos
    try:
        connection = get_db_connection()
        cursor = connection.cursor()
        
        # Obtener los datos estructurados del CV directamente desde user_cvs
        cursor.execute(
            "SELECT cv_name, personal_info, professional_summary, education, experience, skills, languages, certificates, format_options, ai_methodologies FROM user_cvs WHERE id = %s AND user_id = %s AND is_active = TRUE",
            (cv_id, session['user_id'])
        )
        result = cursor.fetchone()
        
        if not result:
            cursor.close()
            connection.close()
            return jsonify({'error': 'Datos del CV no encontrados'}), 404
        
        cursor.close()
        connection.close()
        
        # Función para deserializar JSON de forma segura
        def safe_json_loads(data, default):
            if data is None:
                return default
            if isinstance(data, str):
                try:
                    return json.loads(data)
                except (json.JSONDecodeError, TypeError):
                    return default
            return data if isinstance(data, (dict, list)) else default
        
        # Reconstruir los datos del CV - deserializar JSON strings
        cv_data = {
            'cv_name': result['cv_name'] if result['cv_name'] else 'Mi CV',
            'personal_info': safe_json_loads(result['personal_info'], {}),
            'professional_summary': result['professional_summary'] if result['professional_summary'] else '',
            'education': safe_json_loads(result['education'], []),
            'experience': safe_json_loads(result['experience'], []),
            'skills': safe_json_loads(result['skills'], []),
            'languages': safe_json_loads(result['languages'], []),
            'certificates': safe_json_loads(result['certificates'], []),
            'format_options': safe_json_loads(result['format_options'], {'format': 'hardware', 'tech_xyz': False, 'tech_start': False}),
            'ai_methodologies': safe_json_loads(result['ai_methodologies'], {})
        }
        
        # Obtener idioma objetivo desde ai_methodologies (por defecto español)
        ai_methodologies = cv_data.get('ai_methodologies', {})
        target_language = ai_methodologies.get('target_language', 'es')
        
        # Aplicar mejoras de IA antes de generar HTML (igual que en save_cv)
        improved_cv_data = improve_cv_with_ai(cv_data, target_language)
        
        # Generar el HTML con los datos mejorados por IA
        cv_html = generate_cv_html(improved_cv_data)
        
        # Crear HTML optimizado para PDF con estilos mejorados para impresión
        pdf_optimized_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>{improved_cv_data['cv_name']}</title>
            <style>
                @page {{
                    margin: 0;
                    size: A4;
                    /* Eliminar headers y footers del navegador */
                    @top-left {{ content: ""; }}
                    @top-center {{ content: ""; }}
                    @top-right {{ content: ""; }}
                    @bottom-left {{ content: ""; }}
                    @bottom-center {{ content: ""; }}
                    @bottom-right {{ content: ""; }}
                }}
                
                @media print {{
                    html, body {{
                        margin: 0 !important;
                        padding: 0 !important;
                        width: 210mm !important;
                        height: 297mm !important;
                        font-size: 11pt !important;
                        line-height: 1.3 !important;
                        -webkit-print-color-adjust: exact !important;
                        color-adjust: exact !important;
                    }}
                    
                    body {{
                        padding: 20mm 15mm 15mm 15mm !important;
                    }}
                    
                    .no-print {{ display: none !important; }}
                    .page-break {{ page-break-before: always; }}
                    .section-title {{
                        border-bottom: 1px solid #000 !important;
                        page-break-after: avoid;
                    }}
                    .item {{
                        page-break-inside: avoid;
                    }}
                    
                    /* Eliminar cualquier contenido generado automáticamente */
                    *::before, *::after {{
                        content: none !important;
                    }}
                }}
                
                @media screen {{
                    html {{
                        margin: 0;
                        padding: 0;
                    }}
                    body {{
                        max-width: 210mm;
                        margin: 0 auto;
                        padding: 20px;
                        box-shadow: 0 0 10px rgba(0,0,0,0.1);
                        background: white;
                        min-height: 100vh;
                    }}
                    .print-instructions {{
                        position: fixed;
                        top: 10px;
                        right: 10px;
                        background: #007bff;
                        color: white;
                        padding: 10px 15px;
                        border-radius: 5px;
                        font-family: Arial, sans-serif;
                        font-size: 14px;
                        z-index: 1000;
                        box-shadow: 0 2px 10px rgba(0,0,0,0.2);
                    }}
                    .print-instructions button {{
                        background: white;
                        color: #007bff;
                        border: none;
                        padding: 5px 10px;
                        border-radius: 3px;
                        cursor: pointer;
                        margin-left: 10px;
                        font-weight: bold;
                    }}
                }}
                
                /* Estilos base del CV */
                body {{
                    font-family: 'Times New Roman', serif;
                    margin: 0;
                    padding: 20px;
                    line-height: 1.4;
                    color: #000;
                }}
                
                .header {{
                    text-align: center;
                    padding-bottom: 10px;
                    margin-bottom: 20px;
                }}
                
                .name {{
                    font-size: 24px;
                    font-weight: bold;
                    margin-bottom: 5px;
                }}
                
                .contact {{
                    font-size: 12px;
                }}
                
                .section {{
                    margin-bottom: 20px;
                }}
                
                .section-title {{
                    font-size: 14px;
                    font-weight: bold;
                    border-bottom: 1px solid #000;
                    margin-bottom: 10px;
                    padding-bottom: 2px;
                }}
                
                .item {{
                    margin-bottom: 10px;
                }}
                
                .item-title {{
                    font-weight: bold;
                }}
                
                .item-subtitle {{
                    font-style: italic;
                }}
                
                .item-date {{
                    float: right;
                }}
                
                .skills-list {{
                    display: flex;
                    flex-wrap: wrap;
                    gap: 10px;
                }}
                
                .skill {{
                    background: #f0f0f0;
                    padding: 2px 8px;
                    border-radius: 3px;
                    font-size: 12px;
                }}
                
                .skill-xyz {{
                    background: #d1ecf1;
                    color: #0c5460;
                    font-weight: bold;
                }}
                
                .skill-start {{
                    background: #d4edda;
                    color: #155724;
                    font-weight: bold;
                }}
            </style>
            <script>
                function printPDF() {{
                    window.print();
                }}
                
                // Auto-abrir diálogo de impresión después de 1 segundo
                setTimeout(function() {{
                    if (confirm('¿Deseas abrir el diálogo de impresión para guardar como PDF?')) {{
                        window.print();
                    }}
                }}, 1000);
            </script>
        </head>
        <body>
            <div class="print-instructions no-print">
                📄 Presiona Ctrl+P para guardar como PDF
                <button onclick="printPDF()">Imprimir/PDF</button>
            </div>
            
            {cv_html[cv_html.find('<body>') + 6:cv_html.find('</body>')].strip() if '<body>' in cv_html else cv_html}
        </body>
        </html>
        """
        
        # Devolver el HTML optimizado para PDF
        response = make_response(pdf_optimized_html)
        response.headers['Content-Type'] = 'text/html; charset=utf-8'
        
        return response
        
    except Exception as e:
        # Mejorar el manejo de errores para evitar mensajes confusos
        error_message = str(e) if str(e) and str(e) != '0' else 'Error desconocido en la exportación del CV'
        
        # Log del error para debugging
        import traceback
        print(f"Error en export_cv: {error_message}")
        print(f"Tipo de excepción: {type(e)}")
        print(f"Traceback: {traceback.format_exc()}")
        
        return jsonify({'error': f'Error al obtener CV: {error_message}'}), 500

@app.route('/job_search')
def job_search():
    """Motor de búsqueda de empleos"""
    if 'user_id' not in session:
        return redirect(url_for('login'))
    return render_template('job_search.html')

@app.route('/search_jobs', methods=['POST'])
def search_jobs():
    """Buscar empleos en diferentes portales"""
    if 'user_id' not in session:
        return jsonify({'error': 'No autorizado'}), 401
    
    query = request.json.get('query', '')
    location = request.json.get('location', '')
    source = request.json.get('source', 'all')
    
    # Nota: Las funciones de scraping han sido deshabilitadas temporalmente
    # TODO: Implementar nuevas funciones de búsqueda de empleos
    jobs = []
    
    # Eliminar duplicados
    unique_jobs = remove_duplicate_jobs(jobs)
    
    # Obtener análisis de CV para calcular compatibilidad
    cv_analysis = get_latest_cv_analysis(session['user_id'])
    
    # Calcular compatibilidad con IA si hay análisis de CV
    if cv_analysis and unique_jobs:
        for job in unique_jobs:
            compatibility = calculate_ai_job_compatibility(job, cv_analysis)
            job['compatibility_score'] = compatibility
        
        # Ordenar por compatibilidad (mayor a menor)
        unique_jobs.sort(key=lambda x: x.get('compatibility_score', 0), reverse=True)
    
    # Guardar empleos en la base de datos
    save_jobs_to_db(unique_jobs)
    
    return jsonify({
        'jobs': unique_jobs,
        'total_found': len(unique_jobs),
        'has_ai_scoring': bool(cv_analysis)
    })

@app.route('/ai_job_search', methods=['POST'])
def ai_job_search():
    """Búsqueda inteligente de empleos con IA basada en el CV del usuario"""
    if 'user_id' not in session:
        return jsonify({'error': 'No autorizado'}), 401
    
    try:
        # Obtener el análisis de CV más reciente del usuario
        cv_analysis = get_latest_cv_analysis(session['user_id'])
        
        if not cv_analysis:
            return jsonify({
                'error': 'No se encontró un análisis de CV. Por favor, sube y analiza tu CV primero.'
            }), 400
        
        # Generar términos de búsqueda inteligentes basados en el CV
        search_terms = generate_smart_search_terms(cv_analysis)
        
        # Buscar empleos en múltiples portales
        all_jobs = []
        
        # Nota: Las funciones de scraping han sido deshabilitadas temporalmente
        # TODO: Implementar nuevas funciones de búsqueda de empleos
        all_jobs = []
        
        # Eliminar duplicados
        unique_jobs = remove_duplicate_jobs(all_jobs)
        
        # Calcular compatibilidad con IA para cada trabajo
        jobs_with_compatibility = []
        for job in unique_jobs:
            compatibility = calculate_ai_job_compatibility(job, cv_analysis)
            job['compatibility_score'] = compatibility
            jobs_with_compatibility.append(job)
        
        # Ordenar por compatibilidad (mayor a menor)
        jobs_with_compatibility.sort(key=lambda x: x['compatibility_score'], reverse=True)
        
        # Tomar los mejores 100 empleos
        top_jobs = jobs_with_compatibility[:100]
        
        # Guardar en base de datos
        if top_jobs:
            save_jobs_to_db(top_jobs)
        
        return jsonify({
            'jobs': top_jobs,
            'total_found': len(unique_jobs),
            'search_terms_used': search_terms[:3]
        })
        
    except Exception as e:
        print(f"Error en búsqueda IA: {e}")
        return jsonify({'error': f'Error interno del servidor: {str(e)}'}), 500

@app.route('/my_analyses')
def my_analyses():
    """Ver análisis previos del usuario organizados por tipo y proveedor de IA"""
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    # Obtener análisis organizados desde S3
    user_analyses = get_user_cv_analyses(session['user_id'])
    
    # Fallback: obtener análisis legacy de la base de datos
    connection = get_db_connection()
    legacy_analyses = []
    
    if connection:
        cursor = connection.cursor()
        cursor.execute("""
            SELECT f.id, r.filename, r.created_at, f.score, f.strengths, f.weaknesses, f.recommendations, f.keywords, f.analysis_type, f.ai_provider
            FROM resumes r
            INNER JOIN feedback f ON r.id = f.resume_id
            WHERE r.user_id = %s AND f.s3_key IS NULL
            ORDER BY r.created_at DESC
        """, (session['user_id'],))
        
        results = cursor.fetchall()
        for result in results:
            analysis = {
                'id': result['id'],
                'filename': result['filename'],
                'created_at': result['created_at'],
                'score': result['score'] if result['score'] else 0,
                'strengths': json.loads(result['strengths']) if result['strengths'] else [],
                'weaknesses': json.loads(result['weaknesses']) if result['weaknesses'] else [],
                'recommendations': json.loads(result['recommendations']) if result['recommendations'] else [],
                'keywords': json.loads(result['keywords']) if result['keywords'] else [],
                'analysis_type': result['analysis_type'] or 'general_health_check',
                'ai_provider': result['ai_provider'] or 'openai'
            }
            legacy_analyses.append(analysis)
        
        cursor.close()
        connection.close()
    
    # Definir nombres amigables para tipos de análisis
    analysis_type_names = {
        'general_health_check': 'Revisión General',
        'content_quality_analysis': 'Análisis de Calidad',
        'job_tailoring_optimization': 'Optimización para Empleos',
        'ats_compatibility_verification': 'Compatibilidad ATS',
        'tone_style_evaluation': 'Evaluación de Tono y Estilo',
        'industry_role_feedback': 'Feedback por Industria',
        'benchmarking_comparison': 'Comparación Benchmarking',
        'ai_improvement_suggestions': 'Sugerencias de IA',
        'visual_design_assessment': 'Evaluación de Diseño',
        'comprehensive_score': 'Puntuación Integral'
    }
    
    # Nombres amigables para proveedores de IA
    ai_provider_names = {
        'openai': 'OpenAI GPT',
        'anthropic': 'Anthropic Claude',
        'gemini': 'Google Gemini'
    }
    
    # Calcular el número total de análisis en S3
    s3_analyses_count = sum(len(providers) for providers in user_analyses.values())
    
    # Crear una lista plana de todos los análisis para JavaScript
    all_analyses = []
    
    # Agregar análisis de S3 (convertir estructura anidada a lista plana)
    for analysis_type, providers in user_analyses.items():
        for ai_provider, analysis_data in providers.items():
            all_analyses.append({
                'id': f"s3_{analysis_type}_{ai_provider}",
                'filename': f"Análisis {analysis_type_names.get(analysis_type, analysis_type)}",
                'created_at': analysis_data['created_at'],
                'score': analysis_data.get('score', 0),
                'strengths': analysis_data.get('strengths', []),
                'weaknesses': analysis_data.get('weaknesses', []),
                'recommendations': analysis_data.get('recommendations', []),
                'keywords': analysis_data.get('keywords', []),
                'analysis_type': analysis_type,
                'ai_provider': ai_provider,
                'source': 's3'
            })
    
    # Agregar análisis legacy
    all_analyses.extend(legacy_analyses)
    
    # Calcular puntuación promedio
    scores = []
    for analysis in all_analyses:
        score = analysis.get('score', 0)
        if score and score != 0:
            scores.append(score)
    
    avg_score = sum(scores) / len(scores) if scores else 0
    
    return render_template('my_analyses.html', 
                         s3_analyses=user_analyses,
                         legacy_analyses=legacy_analyses,
                         analyses=all_analyses,
                         type_names=analysis_type_names,
                         provider_names=ai_provider_names,
                         s3_analyses_count=s3_analyses_count,
                         avg_score=avg_score)

def generate_professional_summary_section(professional_summary, use_xyz, use_start, section_title='RESUMEN PROFESIONAL'):
    """Generar sección de resumen profesional con metodologías aplicadas"""
    enhanced_summary = professional_summary
    
    if use_xyz or use_start:
        # Palabras clave y frases para metodología XYZ (enfoque en tecnologías emergentes e innovación)
        xyz_enhancements = [
            "con enfoque en tecnologías emergentes",
            "especializado en soluciones innovadoras",
            "con experiencia en transformación digital",
            "orientado a la implementación de nuevas tecnologías",
            "con visión estratégica en innovación tecnológica"
        ]
        
        # Palabras clave y frases para metodología Start (enfoque en fundamentos sólidos y crecimiento)
        start_enhancements = [
            "con sólidos fundamentos técnicos",
            "enfocado en el crecimiento profesional continuo",
            "con base sólida en principios fundamentales",
            "orientado al desarrollo progresivo de competencias",
            "con enfoque metodológico y estructurado"
        ]
        
        # Aplicar mejoras según las metodologías seleccionadas
        if use_xyz and use_start:
            # Combinar ambas metodologías
            if xyz_enhancements and start_enhancements:
                enhanced_summary += f" {xyz_enhancements[0]} y {start_enhancements[0]}."
        elif use_xyz:
            # Solo metodología XYZ
            if xyz_enhancements:
                enhanced_summary += f" {xyz_enhancements[0]}."
        elif use_start:
            # Solo metodología Start
            if start_enhancements:
                enhanced_summary += f" {start_enhancements[0]}."
    
    return f'<div class="section"><div class="section-title">{section_title}</div><div style="text-align: justify; line-height: 1.4;">{enhanced_summary}</div></div>'

def enhance_experience_description(description, use_xyz, use_start):
    """Mejorar descripción de experiencia con metodologías aplicadas"""
    if not description or not description.strip():
        return description
    
    enhanced_description = description
    
    if use_xyz or use_start:
        # Frases para metodología XYZ (enfoque en innovación y tecnologías emergentes)
        xyz_phrases = [
            "implementando soluciones innovadoras",
            "utilizando tecnologías de vanguardia",
            "desarrollando estrategias disruptivas",
            "aplicando metodologías ágiles y modernas",
            "liderando iniciativas de transformación digital"
        ]
        
        # Frases para metodología Start (enfoque en fundamentos y crecimiento estructurado)
        start_phrases = [
            "aplicando metodologías estructuradas",
            "siguiendo mejores prácticas establecidas",
            "implementando procesos sistemáticos",
            "desarrollando competencias fundamentales",
            "estableciendo bases sólidas para el crecimiento"
        ]
        
        # Aplicar mejoras según las metodologías seleccionadas
        if use_xyz and use_start:
            # Combinar ambas metodologías
            if xyz_phrases and start_phrases:
                enhanced_description += f" Destacando por {xyz_phrases[0]} y {start_phrases[0]}."
        elif use_xyz:
            # Solo metodología XYZ
            if xyz_phrases:
                enhanced_description += f" Destacando por {xyz_phrases[0]}."
        elif use_start:
            # Solo metodología Start
            if start_phrases:
                enhanced_description += f" Destacando por {start_phrases[0]}."
    
    return enhanced_description

def generate_cv_html(cv_data):
    """Generar HTML del CV según el formato seleccionado"""
    personal = cv_data.get('personal_info', {})
    professional_summary = cv_data.get('professional_summary', '')
    education = cv_data.get('education', [])
    experience = cv_data.get('experience', [])
    skills = cv_data.get('skills', [])
    languages = cv_data.get('languages', [])
    certificates = cv_data.get('certificates', [])
    format_options = cv_data.get('format_options', {'format': 'hardware', 'summary_tech_xyz': False, 'summary_tech_start': False, 'experience_tech_xyz': False, 'experience_tech_start': False})
    
    # Obtener títulos de secciones (traducidos o por defecto en español)
    section_titles = cv_data.get('section_titles', {
        'professional_summary': 'RESUMEN PROFESIONAL',
        'education': 'EDUCACIÓN',
        'experience': 'EXPERIENCIA PROFESIONAL',
        'skills': 'HABILIDADES',
        'languages': 'IDIOMAS',
        'certificates': 'CERTIFICACIONES'
    })
    
    # Determinar el formato seleccionado
    is_ats_format = format_options.get('format') == 'ats'
    
    # Metodologías específicas por sección
    summary_tech_xyz = format_options.get('summary_tech_xyz', False)
    summary_tech_start = format_options.get('summary_tech_start', False)
    experience_tech_xyz = format_options.get('experience_tech_xyz', False)
    experience_tech_start = format_options.get('experience_tech_start', False)
    
    # Estilos base según el formato
    if is_ats_format:
        # Estilo ATS: Simple, sin diseño complejo, optimizado para sistemas de seguimiento
        styles = """
            body { font-family: Arial, sans-serif; margin: 0; padding: 20px; line-height: 1.4; }
            .header { text-align: left; border-bottom: 1px solid #000; padding-bottom: 10px; margin-bottom: 20px; }
            .name { font-size: 18px; font-weight: bold; margin-bottom: 5px; }
            .contact { font-size: 12px; }
            .section { margin-bottom: 20px; }
            .section-title { font-size: 14px; font-weight: bold; text-transform: uppercase; margin-bottom: 10px; }
            .item { margin-bottom: 10px; }
            .item-title { font-weight: bold; }
            .item-subtitle { font-style: italic; }
            .item-date { float: right; }
            .skills-list { margin-bottom: 10px; }
            .skill { margin-right: 5px; }
        """
    else:
        # Estilo Hardware: Formato profesional basado en la imagen de referencia
        styles = """
            body { 
                font-family: 'Times New Roman', serif; 
                margin: 0; 
                padding: 20px; 
                line-height: 1.4; 
                font-size: 11pt;
                color: #000;
            }
            .header { 
                text-align: left; 
                padding-bottom: 15px; 
                margin-bottom: 25px; 
            }
            .name { 
                font-size: 18pt; 
                font-weight: bold; 
                margin-bottom: 8px; 
                text-align: center;
            }
            .contact { 
                font-size: 10pt; 
                text-align: center;
                margin-bottom: 5px;
            }
            .section { 
                margin-bottom: 20px; 
            }
            .section-title { 
                font-size: 12pt; 
                font-weight: bold; 
                text-transform: uppercase;
                border-bottom: 1px solid #000; 
                margin-bottom: 12px;
                padding-bottom: 3px;
            }
            .item { 
                margin-bottom: 15px; 
                position: relative;
            }
            .item-title { 
                font-weight: bold; 
                font-size: 11pt;
            }
            .item-subtitle { 
                font-style: italic; 
                font-size: 10pt;
                margin-bottom: 3px;
            }
            .item-date { 
                float: right; 
                font-size: 10pt;
                font-weight: normal;
            }
            .item-description {
                text-align: justify;
                margin-top: 5px;
                font-size: 10pt;
                line-height: 1.3;
            }
            .skills-list { 
                text-align: justify;
                font-size: 10pt;
            }
            .skill { 
                display: inline;
                margin-right: 8px;
            }
            .languages-section {
                font-size: 10pt;
            }
            .language-item {
                margin-bottom: 5px;
            }
        """
    
    # Construir información de contacto
    contact_info = [personal.get('email', ''), personal.get('phone', ''), personal.get('address', '')]
    contact_info = [info for info in contact_info if info]  # Filtrar campos vacíos
    contact_line = ' | '.join(contact_info)
    
    # Construir redes sociales
    social_links = []
    if personal.get('linkedin'):
        social_links.append(f"LinkedIn: {personal.get('linkedin')}")
    if personal.get('github'):
        social_links.append(f"GitHub: {personal.get('github')}")
    if personal.get('website'):
        social_links.append(f"Web: {personal.get('website')}")
    
    social_line = ' | '.join(social_links) if social_links else ''
    
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>CV - {personal.get('name', '')}</title>
        <style>
            {styles}
        </style>
    </head>
    <body>
        <div class="header">
            <div class="name">{personal.get('name', '')}</div>
            <div style="border-bottom: 2px solid #000; margin: 5px 0;"></div>
            {f'<div style="font-size: 16pt; font-weight: 600; color: #000; margin-bottom: 8px; text-transform: uppercase; text-align: center;">{personal.get("position", "")}</div>' if personal.get('position') else ''}
            <div class="contact">
                {contact_line}
            </div>
            {f'<div class="contact" style="font-size: 11px; margin-top: 5px;">{social_line}</div>' if social_line else ''}
        </div>
        
        {generate_professional_summary_section(professional_summary, summary_tech_xyz, summary_tech_start, section_titles['professional_summary']) if professional_summary and professional_summary.strip() else ''}
        
        <div class="section">
            <div class="section-title">{section_titles['education']}</div>
    """
    
    for edu in education:
        html += f"""
            <div class="item">
                <div class="item-date">{edu.get('end_date', '')}</div>
                <div class="item-title">{edu.get('career', '')}</div>
                <div class="item-subtitle">{edu.get('institution', '')}</div>
                {f'<div class="item-description">{edu.get("description", "")}</div>' if edu.get('description', '').strip() else ''}
            </div>
        """
    
    html += f"""
        </div>
        
        <div class="section">
            <div class="section-title">{section_titles['experience']}</div>
    """
    
    for exp in experience:
        enhanced_description = enhance_experience_description(exp.get('description', ''), experience_tech_xyz, experience_tech_start)
        html += f"""
            <div class="item">
                <div class="item-date">{exp.get('start_date', '')} - {exp.get('end_date', '')}</div>
                <div class="item-title">{exp.get('position', '')}</div>
                <div class="item-subtitle">{exp.get('company', '')}</div>
                {f'<div class="item-description">{enhanced_description}</div>' if enhanced_description.strip() else ''}
            </div>
        """
    
    if skills:
        html += f"""
        </div>
        
        <div class="section">
            <div class="section-title">{section_titles['skills']}</div>
        """
        
        # Palabras clave para tecnologías XYZ (emergentes)
        xyz_keywords = ['AI', 'Machine Learning', 'Blockchain', 'IoT', 'Cloud', 'DevOps', 'React', 'Vue', 'Angular', 'Node.js']
        
        # Palabras clave para tecnologías Start (fundamentales)
        start_keywords = ['Java', 'Python', 'C++', 'SQL', 'HTML', 'CSS', 'JavaScript', 'Git', 'Linux', 'Windows']
        
        if is_ats_format:
            # Formato ATS: habilidades con estilos individuales
            html += '<div class="skills-list">'
            for skill in skills:
                skill_class = "skill"
                if any(keyword.lower() in skill.lower() for keyword in xyz_keywords):
                    skill_class = "skill skill-xyz"
                elif any(keyword.lower() in skill.lower() for keyword in start_keywords):
                    skill_class = "skill skill-start"
                html += f'<span class="{skill_class}">{skill}</span>'
            html += '</div>'
        else:
            # Formato Hardware: habilidades con estilos individuales
            html += '<div class="skills-list">'
            for skill in skills:
                skill_class = "skill"
                if any(keyword.lower() in skill.lower() for keyword in xyz_keywords):
                    skill_class = "skill skill-xyz"
                elif any(keyword.lower() in skill.lower() for keyword in start_keywords):
                    skill_class = "skill skill-start"
                html += f'<span class="{skill_class}">{skill}</span>'
            html += '</div>'
    
    if languages:
        html += f"""
        </div>
        
        <div class="section">
            <div class="section-title">{section_titles['languages']}</div>
        """
        
        for lang in languages:
            html += f'<div class="language-item">{lang.get("language", "")} - {lang.get("level", "")}</div>'
    
    if certificates:
        html += f"""
        </div>
        
        <div class="section">
            <div class="section-title">{section_titles['certificates']}</div>
        """
        
        for cert in certificates:
            html += f"""
            <div class="item">
                <div class="item-date">{cert.get('date', '')}</div>
                <div class="item-title">{cert.get('title', '')}</div>
                <div class="item-subtitle">{cert.get('institution', '')}</div>
            </div>
        """
    
    html += """
        </div>
    </body>
    </html>
    """
    
    return html

def scrape_computrabajo(query, location):
    """Scraping de empleos de CompuTrabajo"""
    jobs = []
    
    try:
        # Construir URL de búsqueda
        base_url = "https://www.computrabajo.com"
        search_url = f"{base_url}/empleos-publicados-en-{location.lower().replace(' ', '-')}?q={query.replace(' ', '+')}"
        
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        response = requests.get(search_url, headers=headers, timeout=10)
        # soup = BeautifulSoup(response.content, 'html.parser')  # Temporarily disabled
        
        # Buscar ofertas de empleo
        # job_listings = soup.find_all('div', class_='box_offer')  # Temporarily disabled
        job_listings = []  # Empty list when scraping is disabled
        
        for listing in job_listings[:10]:  # Limitar a 10 resultados
            try:
                title_elem = listing.find('h2', class_='fs18')
                company_elem = listing.find('p', class_='fs16')
                location_elem = listing.find('p', class_='fs13')
                
                if title_elem and company_elem:
                    title = title_elem.get_text(strip=True)
                    company = company_elem.get_text(strip=True)
                    job_location = location_elem.get_text(strip=True) if location_elem else location
                    
                    # Obtener enlace
                    link_elem = title_elem.find('a')
                    job_url = base_url + link_elem['href'] if link_elem else ''
                    
                    jobs.append({
                        'title': title,
                        'company': company,
                        'location': job_location,
                        'url': job_url,
                        'source': 'CompuTrabajo',
                        'description': ''
                    })
            except Exception as e:
                continue
                
    except Exception as e:
        print(f"Error scraping CompuTrabajo: {e}")
    
    return jobs

def scrape_indeed(query, location):
    """Scraping de empleos de Indeed"""
    jobs = []
    
    try:
        # Construir URL de búsqueda
        base_url = "https://www.indeed.com"
        search_url = f"{base_url}/jobs?q={query.replace(' ', '+')}&l={location.replace(' ', '+')}"
        
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        response = requests.get(search_url, headers=headers, timeout=10)
        # soup = BeautifulSoup(response.content, 'html.parser')  # Temporarily disabled
        
        # Buscar ofertas de empleo
        # job_listings = soup.find_all('div', class_='job_seen_beacon')  # Temporarily disabled
        job_listings = []  # Empty list when scraping is disabled
        
        for listing in job_listings[:10]:  # Limitar a 10 resultados
            try:
                title_elem = listing.find('h2', class_='jobTitle')
                company_elem = listing.find('span', class_='companyName')
                location_elem = listing.find('div', class_='companyLocation')
                
                if title_elem and company_elem:
                    title_link = title_elem.find('a')
                    title = title_link.get_text(strip=True) if title_link else title_elem.get_text(strip=True)
                    company = company_elem.get_text(strip=True)
                    job_location = location_elem.get_text(strip=True) if location_elem else location
                    
                    # Obtener enlace
                    job_url = base_url + title_link['href'] if title_link and title_link.get('href') else ''
                    
                    jobs.append({
                        'title': title,
                        'company': company,
                        'location': job_location,
                        'url': job_url,
                        'source': 'Indeed',
                        'description': ''
                    })
            except Exception as e:
                continue
                
    except Exception as e:
        print(f"Error scraping Indeed: {e}")
    
    return jobs

def save_jobs_to_db(jobs):
    """Guardar empleos en la base de datos"""
    connection = get_db_connection()
    if connection:
        cursor = connection.cursor()
        
        for job in jobs:
            try:
                # Verificar si el empleo ya existe
                cursor.execute(
                    "SELECT id FROM jobs WHERE title = %s AND company = %s AND url = %s",
                    (job['title'], job['company'], job['url'])
                )
                
                if not cursor.fetchone():
                    cursor.execute(
                        "INSERT INTO jobs (title, company, location, description, url, source) VALUES (%s, %s, %s, %s, %s, %s)",
                        (job['title'], job['company'], job['location'], job['description'], job['url'], job['source'])
                    )
            except Exception as e:
                print(f"Error guardando empleo: {e}")
                continue
        
        connection.commit()
        cursor.close()
        connection.close()

@app.route('/profile')
def profile():
    """Perfil del usuario"""
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    connection = get_db_connection()
    user = None
    analyses_count = 0
    last_analysis = None
    
    if connection:
        cursor = connection.cursor()
        cursor.execute(
            "SELECT id, username, first_name, last_name, phone, email, created_at, email_verified FROM users WHERE id = %s",
            (session['user_id'],)
        )
        user = cursor.fetchone()
        
        # Obtener estadísticas de análisis
        cursor.execute(
            "SELECT COUNT(*) FROM resumes WHERE user_id = %s",
            (session['user_id'],)
        )
        analyses_count = cursor.fetchone()['count']
        
        cursor.execute(
            "SELECT created_at FROM resumes WHERE user_id = %s ORDER BY created_at DESC LIMIT 1",
            (session['user_id'],)
        )
        result = cursor.fetchone()
        last_analysis = result['created_at'] if result else None
        
        cursor.close()
        connection.close()
    
    return render_template('profile.html', user=user, analyses_count=analyses_count, last_analysis=last_analysis)

def generate_professional_summary_section(professional_summary, use_xyz, use_start, section_title='RESUMEN PROFESIONAL'):
    """Generar sección de resumen profesional con metodologías aplicadas"""
    enhanced_summary = professional_summary
    
    if use_xyz or use_start:
        # Palabras clave y frases para metodología XYZ (enfoque en tecnologías emergentes e innovación)
        xyz_enhancements = [
            "con enfoque en tecnologías emergentes",
            "especializado en soluciones innovadoras",
            "con experiencia en transformación digital",
            "orientado a la implementación de nuevas tecnologías",
            "con visión estratégica en innovación tecnológica"
        ]
        
        # Palabras clave y frases para metodología Start (enfoque en fundamentos sólidos y crecimiento)
        start_enhancements = [
            "con sólidos fundamentos técnicos",
            "enfocado en el crecimiento profesional continuo",
            "con base sólida en principios fundamentales",
            "orientado al desarrollo progresivo de competencias",
            "con enfoque metodológico y estructurado"
        ]
        
        # Aplicar mejoras según las metodologías seleccionadas
        if use_xyz and use_start:
            # Combinar ambas metodologías
            if xyz_enhancements and start_enhancements:
                enhanced_summary += f" {xyz_enhancements[0]} y {start_enhancements[0]}."
        elif use_xyz:
            # Solo metodología XYZ
            if xyz_enhancements:
                enhanced_summary += f" {xyz_enhancements[0]}."
        elif use_start:
            # Solo metodología Start
            if start_enhancements:
                enhanced_summary += f" {start_enhancements[0]}."
    
    return f'<div class="section"><div class="section-title">{section_title}</div><div style="text-align: justify; line-height: 1.4;">{enhanced_summary}</div></div>'

def enhance_experience_description(description, use_xyz, use_start):
    """Mejorar descripción de experiencia con metodologías aplicadas"""
    if not description or not description.strip():
        return description
    
    enhanced_description = description
    
    if use_xyz or use_start:
        # Frases para metodología XYZ (enfoque en innovación y tecnologías emergentes)
        xyz_phrases = [
            "implementando soluciones innovadoras",
            "utilizando tecnologías de vanguardia",
            "desarrollando estrategias disruptivas",
            "aplicando metodologías ágiles y modernas",
            "liderando iniciativas de transformación digital"
        ]
        
        # Frases para metodología Start (enfoque en fundamentos y crecimiento estructurado)
        start_phrases = [
            "aplicando metodologías estructuradas",
            "siguiendo mejores prácticas establecidas",
            "implementando procesos sistemáticos",
            "desarrollando competencias fundamentales",
            "estableciendo bases sólidas para el crecimiento"
        ]
        
        # Aplicar mejoras según las metodologías seleccionadas
        if use_xyz and use_start:
            # Combinar ambas metodologías
            if xyz_phrases and start_phrases:
                enhanced_description += f" Destacando por {xyz_phrases[0]} y {start_phrases[0]}."
        elif use_xyz:
            # Solo metodología XYZ
            if xyz_phrases:
                enhanced_description += f" Destacando por {xyz_phrases[0]}."
        elif use_start:
            # Solo metodología Start
            if start_phrases:
                enhanced_description += f" Destacando por {start_phrases[0]}."
    
    return enhanced_description

def generate_cv_html(cv_data):
    """Generar HTML del CV según el formato seleccionado"""
    personal = cv_data.get('personal_info', {})
    professional_summary = cv_data.get('professional_summary', '')
    education = cv_data.get('education', [])
    experience = cv_data.get('experience', [])
    skills = cv_data.get('skills', [])
    languages = cv_data.get('languages', [])
    certificates = cv_data.get('certificates', [])
    format_options = cv_data.get('format_options', {'format': 'hardware', 'summary_tech_xyz': False, 'summary_tech_start': False, 'experience_tech_xyz': False, 'experience_tech_start': False})
    
    # Obtener títulos de secciones (traducidos o por defecto en español)
    section_titles = cv_data.get('section_titles', {
        'professional_summary': 'RESUMEN PROFESIONAL',
        'education': 'EDUCACIÓN',
        'experience': 'EXPERIENCIA PROFESIONAL',
        'skills': 'HABILIDADES',
        'languages': 'IDIOMAS',
        'certificates': 'CERTIFICACIONES'
    })
    
    # Determinar el formato seleccionado
    is_ats_format = format_options.get('format') == 'ats'
    
    # Metodologías específicas por sección
    summary_tech_xyz = format_options.get('summary_tech_xyz', False)
    summary_tech_start = format_options.get('summary_tech_start', False)
    experience_tech_xyz = format_options.get('experience_tech_xyz', False)
    experience_tech_start = format_options.get('experience_tech_start', False)
    
    # Estilos base según el formato
    if is_ats_format:
        # Estilo ATS: Simple, sin diseño complejo, optimizado para sistemas de seguimiento
        styles = """
            body { font-family: Arial, sans-serif; margin: 0; padding: 20px; line-height: 1.4; }
            .header { text-align: left; border-bottom: 1px solid #000; padding-bottom: 10px; margin-bottom: 20px; }
            .name { font-size: 18px; font-weight: bold; margin-bottom: 5px; }
            .contact { font-size: 12px; }
            .section { margin-bottom: 20px; }
            .section-title { font-size: 14px; font-weight: bold; text-transform: uppercase; margin-bottom: 10px; }
            .item { margin-bottom: 10px; }
            .item-title { font-weight: bold; }
            .item-subtitle { font-style: italic; }
            .item-date { float: right; }
            .skills-list { margin-bottom: 10px; }
            .skill { margin-right: 5px; }
        """
    else:
        # Estilo Hardware: Formato profesional basado en la imagen de referencia
        styles = """
            body { 
                font-family: 'Times New Roman', serif; 
                margin: 0; 
                padding: 20px; 
                line-height: 1.4; 
                font-size: 11pt;
                color: #000;
            }
            .header { 
                text-align: left; 
                padding-bottom: 15px; 
                margin-bottom: 25px; 
            }
            .name { 
                font-size: 18pt; 
                font-weight: bold; 
                margin-bottom: 8px; 
                text-align: center;
            }
            .contact { 
                font-size: 10pt; 
                text-align: center;
                margin-bottom: 5px;
            }
            .section { 
                margin-bottom: 20px; 
            }
            .section-title { 
                font-size: 12pt; 
                font-weight: bold; 
                text-transform: uppercase;
                border-bottom: 1px solid #000; 
                margin-bottom: 12px;
                padding-bottom: 3px;
            }
            .item { 
                margin-bottom: 15px; 
                position: relative;
            }
            .item-title { 
                font-weight: bold; 
                font-size: 11pt;
            }
            .item-subtitle { 
                font-style: italic; 
                font-size: 10pt;
                margin-bottom: 3px;
            }
            .item-date { 
                float: right; 
                font-size: 10pt;
                font-weight: normal;
            }
            .item-description {
                text-align: justify;
                margin-top: 5px;
                font-size: 10pt;
                line-height: 1.3;
            }
            .skills-list { 
                text-align: justify;
                font-size: 10pt;
            }
            .skill { 
                display: inline;
                margin-right: 8px;
            }
            .languages-section {
                font-size: 10pt;
            }
            .language-item {
                margin-bottom: 5px;
            }
        """
    
    # Construir información de contacto
    contact_info = [personal.get('email', ''), personal.get('phone', ''), personal.get('address', '')]
    contact_info = [info for info in contact_info if info]  # Filtrar campos vacíos
    contact_line = ' | '.join(contact_info)
    
    # Construir redes sociales
    social_links = []
    if personal.get('linkedin'):
        social_links.append(f"LinkedIn: {personal.get('linkedin')}")
    if personal.get('github'):
        social_links.append(f"GitHub: {personal.get('github')}")
    if personal.get('website'):
        social_links.append(f"Web: {personal.get('website')}")
    
    social_line = ' | '.join(social_links) if social_links else ''
    
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>CV - {personal.get('name', '')}</title>
        <style>
            {styles}
        </style>
    </head>
    <body>
        <div class="header">
            <div class="name">{personal.get('name', '')}</div>
            <div style="border-bottom: 2px solid #000; margin: 5px 0;"></div>
            {f'<div style="font-size: 16pt; font-weight: 600; color: #000; margin-bottom: 8px; text-transform: uppercase; text-align: center;">{personal.get("position", "")}</div>' if personal.get('position') else ''}
            <div class="contact">
                {contact_line}
            </div>
            {f'<div class="contact" style="font-size: 11px; margin-top: 5px;">{social_line}</div>' if social_line else ''}
        </div>
        
        {generate_professional_summary_section(professional_summary, summary_tech_xyz, summary_tech_start, section_titles['professional_summary']) if professional_summary and professional_summary.strip() else ''}
        
        <div class="section">
            <div class="section-title">{section_titles['education']}</div>
    """
    
    for edu in education:
        html += f"""
            <div class="item">
                <div class="item-date">{edu.get('end_date', '')}</div>
                <div class="item-title">{edu.get('career', '')}</div>
                <div class="item-subtitle">{edu.get('institution', '')}</div>
                {f'<div class="item-description">{edu.get("description", "")}</div>' if edu.get('description', '').strip() else ''}
            </div>
        """
    
    html += f"""
        </div>
        
        <div class="section">
            <div class="section-title">{section_titles['experience']}</div>
    """
    
    for exp in experience:
        enhanced_description = enhance_experience_description(exp.get('description', ''), experience_tech_xyz, experience_tech_start)
        html += f"""
            <div class="item">
                <div class="item-date">{exp.get('start_date', '')} - {exp.get('end_date', '')}</div>
                <div class="item-title">{exp.get('position', '')}</div>
                <div class="item-subtitle">{exp.get('company', '')}</div>
                {f'<div class="item-description">{enhanced_description}</div>' if enhanced_description.strip() else ''}
            </div>
        """
    
    if skills:
        html += f"""
        </div>
        
        <div class="section">
            <div class="section-title">{section_titles['skills']}</div>
        """
        
        # Palabras clave para tecnologías XYZ (emergentes)
        xyz_keywords = ['AI', 'Machine Learning', 'Blockchain', 'IoT', 'Cloud', 'DevOps', 'React', 'Vue', 'Angular', 'Node.js']
        
        # Palabras clave para tecnologías Start (fundamentales)
        start_keywords = ['Java', 'Python', 'C++', 'SQL', 'HTML', 'CSS', 'JavaScript', 'Git', 'Linux', 'Windows']
        
        if is_ats_format:
            # Formato ATS: habilidades con estilos individuales
            html += '<div class="skills-list">'
            for skill in skills:
                skill_class = "skill"
                if any(keyword.lower() in skill.lower() for keyword in xyz_keywords):
                    skill_class = "skill skill-xyz"
                elif any(keyword.lower() in skill.lower() for keyword in start_keywords):
                    skill_class = "skill skill-start"
                html += f'<span class="{skill_class}">{skill}</span>'
            html += '</div>'
        else:
            # Formato Hardware: habilidades con estilos individuales
            html += '<div class="skills-list">'
            for skill in skills:
                skill_class = "skill"
                if any(keyword.lower() in skill.lower() for keyword in xyz_keywords):
                    skill_class = "skill skill-xyz"
                elif any(keyword.lower() in skill.lower() for keyword in start_keywords):
                    skill_class = "skill skill-start"
                html += f'<span class="{skill_class}">{skill}</span>'
            html += '</div>'
    
    if languages:
        html += f"""
        </div>
        
        <div class="section">
            <div class="section-title">{section_titles['languages']}</div>
        """
        
        for lang in languages:
            html += f'<div class="language-item">{lang.get("language", "")} - {lang.get("level", "")}</div>'
    
    if certificates:
        html += f"""
        </div>
        
        <div class="section">
            <div class="section-title">{section_titles['certificates']}</div>
        """
        
        for cert in certificates:
            html += f"""
            <div class="item">
                <div class="item-date">{cert.get('date', '')}</div>
                <div class="item-title">{cert.get('title', '')}</div>
                <div class="item-subtitle">{cert.get('institution', '')}</div>
            </div>
        """
    
    html += """
        </div>
    </body>
    </html>
    """
    
    return html

def scrape_computrabajo(query, location):
    """Scraping de empleos de CompuTrabajo"""
    jobs = []
    
    try:
        # Construir URL de búsqueda
        base_url = "https://www.computrabajo.com"
        search_url = f"{base_url}/empleos-publicados-en-{location.lower().replace(' ', '-')}?q={query.replace(' ', '+')}"
        
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        response = requests.get(search_url, headers=headers, timeout=10)
        # soup = BeautifulSoup(response.content, 'html.parser')  # Temporarily disabled
        
        # Buscar ofertas de empleo
        # job_listings = soup.find_all('div', class_='box_offer')  # Temporarily disabled
        job_listings = []  # Empty list when scraping is disabled
        
        for listing in job_listings[:10]:  # Limitar a 10 resultados
            try:
                title_elem = listing.find('h2', class_='fs18')
                company_elem = listing.find('p', class_='fs16')
                location_elem = listing.find('p', class_='fs13')
                
                if title_elem and company_elem:
                    title = title_elem.get_text(strip=True)
                    company = company_elem.get_text(strip=True)
                    job_location = location_elem.get_text(strip=True) if location_elem else location
                    
                    # Obtener enlace
                    link_elem = title_elem.find('a')
                    job_url = base_url + link_elem['href'] if link_elem else ''
                    
                    jobs.append({
                        'title': title,
                        'company': company,
                        'location': job_location,
                        'url': job_url,
                        'source': 'CompuTrabajo',
                        'description': ''
                    })
            except Exception as e:
                continue
                
    except Exception as e:
        print(f"Error scraping CompuTrabajo: {e}")
    
    return jobs

def scrape_indeed(query, location):
    """Scraping de empleos de Indeed"""
    jobs = []
    
    try:
        # Construir URL de búsqueda
        base_url = "https://www.indeed.com"
        search_url = f"{base_url}/jobs?q={query.replace(' ', '+')}&l={location.replace(' ', '+')}"
        
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        response = requests.get(search_url, headers=headers, timeout=10)
        # soup = BeautifulSoup(response.content, 'html.parser')  # Temporarily disabled
        
        # Buscar ofertas de empleo
        # job_listings = soup.find_all('div', class_='job_seen_beacon')  # Temporarily disabled
        job_listings = []  # Empty list when scraping is disabled
        
        for listing in job_listings[:10]:  # Limitar a 10 resultados
            try:
                title_elem = listing.find('h2', class_='jobTitle')
                company_elem = listing.find('span', class_='companyName')
                location_elem = listing.find('div', class_='companyLocation')
                
                if title_elem and company_elem:
                    title_link = title_elem.find('a')
                    title = title_link.get_text(strip=True) if title_link else title_elem.get_text(strip=True)
                    company = company_elem.get_text(strip=True)
                    job_location = location_elem.get_text(strip=True) if location_elem else location
                    
                    # Obtener enlace
                    job_url = base_url + title_link['href'] if title_link and title_link.get('href') else ''
                    
                    jobs.append({
                        'title': title,
                        'company': company,
                        'location': job_location,
                        'url': job_url,
                        'source': 'Indeed',
                        'description': ''
                    })
            except Exception as e:
                continue
                
    except Exception as e:
        print(f"Error scraping Indeed: {e}")
    
    return jobs

def save_jobs_to_db(jobs):
    """Guardar empleos en la base de datos"""
    connection = get_db_connection()
    if connection:
        cursor = connection.cursor()
        
        for job in jobs:
            try:
                # Verificar si el empleo ya existe
                cursor.execute(
                    "SELECT id FROM jobs WHERE title = %s AND company = %s AND url = %s",
                    (job['title'], job['company'], job['url'])
                )
                
                if not cursor.fetchone():
                    cursor.execute(
                        "INSERT INTO jobs (title, company, location, description, url, source) VALUES (%s, %s, %s, %s, %s, %s)",
                        (job['title'], job['company'], job['location'], job['description'], job['url'], job['source'])
                    )
            except Exception as e:
                print(f"Error guardando empleo: {e}")
                continue
        
        connection.commit()
        cursor.close()
        connection.close()

@app.route('/delete_analysis/<int:analysis_id>', methods=['DELETE'])
def delete_analysis(analysis_id):
    """Eliminar un análisis específico"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'No autorizado'}), 401
    
    connection = get_db_connection()
    if not connection:
        return jsonify({'success': False, 'message': 'Error de conexión a la base de datos'}), 500
    
    try:
        cursor = connection.cursor()
        
        # Verificar que el análisis pertenece al usuario a través de la tabla resumes
        cursor.execute(
            "SELECT f.id FROM feedback f JOIN resumes r ON f.resume_id = r.id WHERE f.id = %s AND r.user_id = %s",
            (analysis_id, session['user_id'])
        )
        
        if not cursor.fetchone():
            return jsonify({'success': False, 'message': 'Análisis no encontrado'}), 404
        
        # Eliminar el análisis de la tabla feedback
        cursor.execute(
            "DELETE FROM feedback WHERE id = %s AND resume_id IN (SELECT id FROM resumes WHERE user_id = %s)",
            (analysis_id, session['user_id'])
        )
        
        connection.commit()
        cursor.close()
        connection.close()
        
        return jsonify({'success': True, 'message': 'Análisis eliminado correctamente'})
        
    except Exception as e:
        print(f"Error eliminando análisis: {e}")
        return jsonify({'success': False, 'message': 'Error interno del servidor'}), 500

@app.route('/delete_account', methods=['DELETE'])
def delete_account():
    """Eliminar cuenta de usuario"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'No autorizado'}), 401
    
    connection = get_db_connection()
    if not connection:
        return jsonify({'success': False, 'message': 'Error de conexión a la base de datos'}), 500
    
    try:
        cursor = connection.cursor()
        user_id = session['user_id']
        
        # Eliminar todos los análisis del usuario
        cursor.execute("DELETE FROM feedback WHERE resume_id IN (SELECT id FROM resumes WHERE user_id = %s)", (user_id,))
        
        # Eliminar la cuenta del usuario
        cursor.execute("DELETE FROM users WHERE id = %s", (user_id,))
        
        connection.commit()
        cursor.close()
        connection.close()
        
        # Limpiar la sesión
        session.clear()
        
        return jsonify({'success': True, 'message': 'Cuenta eliminada correctamente'})
        
    except Exception as e:
        print(f"Error eliminando cuenta: {e}")
        return jsonify({'success': False, 'message': 'Error interno del servidor'}), 500

@app.route('/change_password', methods=['POST'])
def change_password():
    """Cambiar contraseña del usuario"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'No autorizado'}), 401
    
    data = request.get_json()
    current_password = data.get('current_password')
    new_password = data.get('new_password')
    
    if not current_password or not new_password:
        return jsonify({'success': False, 'message': 'Faltan datos requeridos'}), 400
    
    if len(new_password) < 6:
        return jsonify({'success': False, 'message': 'La nueva contraseña debe tener al menos 6 caracteres'}), 400
    
    connection = get_db_connection()
    if not connection:
        return jsonify({'success': False, 'message': 'Error de conexión a la base de datos'}), 500
    
    cursor = None
    try:
        cursor = connection.cursor()
        
        # Verificar contraseña actual
        cursor.execute(
            "SELECT password_hash FROM users WHERE id = %s",
            (session['user_id'],)
        )
        
        user_data = cursor.fetchone()
        print(f"Debug - user_data: {user_data}")
        print(f"Debug - session user_id: {session.get('user_id')}")
        
        if not user_data:
            cursor.close()
            connection.close()
            return jsonify({'success': False, 'message': 'Usuario no encontrado'}), 400
            
        if not user_data[0]:
            cursor.close()
            connection.close()
            return jsonify({'success': False, 'message': 'Password hash no encontrado'}), 400
            
        if not check_password_hash(user_data[0], current_password):
            cursor.close()
            connection.close()
            return jsonify({'success': False, 'message': 'Contraseña actual incorrecta'}), 400
        
        # Actualizar contraseña
        hashed_password = generate_password_hash(new_password)
        cursor.execute(
            "UPDATE users SET password_hash = %s WHERE id = %s",
            (hashed_password, session['user_id'])
        )
        
        # Verificar si se actualizó alguna fila
        if cursor.rowcount == 0:
            cursor.close()
            connection.close()
            return jsonify({'success': False, 'message': 'No se pudo actualizar la contraseña. Usuario no encontrado.'}), 400
        
        connection.commit()
        cursor.close()
        connection.close()
        
        return jsonify({'success': True, 'message': 'Contraseña cambiada correctamente'})
        
    except Exception as e:
        print(f"Error cambiando contraseña: {e}")
        if cursor:
            cursor.close()
        if connection:
            connection.close()
        return jsonify({'success': False, 'message': 'Error interno del servidor'}), 500

# ==================== RUTAS DE ADMINISTRACIÓN ====================

def admin_required(f):
    """Decorador para verificar que el usuario sea administrador"""
    from functools import wraps
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            flash('Debes iniciar sesión para acceder a esta página', 'error')
            return redirect(url_for('login'))
        if session.get('user_role') != 'admin':
            flash('No tienes permisos para acceder a esta página', 'error')
            return redirect(url_for('dashboard'))
        return f(*args, **kwargs)
    return decorated_function

@app.route('/admin')
@app.route('/admin/dashboard')
@admin_required
def admin_dashboard():
    """Panel principal de administración"""
    username = session.get('username', 'unknown')
    add_console_log('INFO', f'Admin accedió al panel principal: {username}', 'ADMIN')
    return render_template('admin/dashboard.html')

@app.route('/admin/console')
@admin_required
def admin_console():
    """Consola en tiempo real del servidor"""
    username = session.get('username', 'unknown')
    add_console_log('INFO', f'Admin accedió a la consola del servidor: {username}', 'ADMIN')
    return render_template('admin/console.html')

@app.route('/admin/stats')
@admin_required
def admin_stats():
    """Estadísticas de usuarios"""
    username = session.get('username', 'unknown')
    add_console_log('INFO', f'Admin accedió a estadísticas: {username}', 'ADMIN')
    
    connection = get_db_connection()
    if not connection:
        flash('Error de conexión a la base de datos', 'error')
        return redirect(url_for('admin_dashboard'))
    
    try:
        cursor = connection.cursor()
        
        # Usuarios en tiempo real (últimos 5 minutos)
        cursor.execute("""
            SELECT COUNT(*) as active_users 
            FROM users 
            WHERE last_login >= NOW() - INTERVAL '5 minutes'
        """)
        active_users = cursor.fetchone()['active_users']
        
        # Usuarios conectados hoy
        cursor.execute("""
            SELECT COUNT(*) as daily_users 
            FROM users 
            WHERE DATE(last_login) = CURRENT_DATE
        """)
        daily_users = cursor.fetchone()['daily_users']
        
        # Usuarios conectados esta semana
        cursor.execute("""
            SELECT COUNT(*) as weekly_users 
            FROM users 
            WHERE last_login >= DATE_TRUNC('week', NOW())
        """)
        weekly_users = cursor.fetchone()['weekly_users']
        
        # Usuarios conectados este mes
        cursor.execute("""
            SELECT COUNT(*) as monthly_users 
            FROM users 
            WHERE last_login >= DATE_TRUNC('month', NOW())
        """)
        monthly_users = cursor.fetchone()['monthly_users']
        
        # Total de usuarios registrados
        cursor.execute("SELECT COUNT(*) as total_users FROM users")
        total_users = cursor.fetchone()['total_users']
        
        cursor.close()
        connection.close()
        
        stats = {
            'active_users': active_users,
            'daily_users': daily_users,
            'weekly_users': weekly_users,
            'monthly_users': monthly_users,
            'total_users': total_users
        }
        
        return render_template('admin/stats.html', stats=stats)
        
    except Exception as e:
        print(f"Error obteniendo estadísticas: {e}")
        flash('Error obteniendo estadísticas', 'error')
        return redirect(url_for('admin_dashboard'))

@app.route('/admin/database')
@admin_required
def admin_database():
    """Interfaz de consultas a la base de datos"""
    return render_template('admin/database.html')

@app.route('/admin/search_users', methods=['POST'])
@admin_required
def admin_search_users():
    """Buscar usuarios por título profesional"""
    search_term = request.form.get('search_query', '').strip()
    
    if not search_term:
        flash('Por favor, introduce un término de búsqueda', 'warning')
        return redirect(url_for('admin_database'))
    
    connection = get_db_connection()
    if not connection:
        flash('Error de conexión a la base de datos', 'error')
        return redirect(url_for('admin_database'))
    
    try:
        cursor = connection.cursor()
        
        # Buscar en los datos de CV de usuarios
        cursor.execute("""
            SELECT u.id, u.username, u.email, u.created_at, u.last_login,
                   cv.personal_info, cv.professional_summary, cv.education, 
                   cv.experience, cv.skills
            FROM users u
            LEFT JOIN user_cv_data cv ON u.id = cv.user_id
            WHERE cv.personal_info ILIKE %s 
               OR cv.professional_summary ILIKE %s
               OR cv.education ILIKE %s
               OR cv.experience ILIKE %s
               OR cv.skills ILIKE %s
        """, (f'%{search_term}%', f'%{search_term}%', f'%{search_term}%', f'%{search_term}%', f'%{search_term}%'))
        
        results = cursor.fetchall()
        cursor.close()
        connection.close()
        
        return render_template('admin/search_results.html', results=results, search_term=search_term)
        
    except Exception as e:
        print(f"Error en búsqueda de usuarios: {e}")
        flash('Error realizando la búsqueda', 'error')
        return redirect(url_for('admin_database'))

@app.route('/admin/export_users', methods=['POST'])
@admin_required
def admin_export_users():
    """Exportar usuarios a Excel"""
    search_term = request.form.get('search_query', '')
    
    connection = get_db_connection()
    if not connection:
        flash('Error de conexión a la base de datos', 'error')
        return redirect(url_for('admin_database'))
    
    try:
        import pandas as pd
        from io import BytesIO
        
        cursor = connection.cursor()
        
        if search_term:
            cursor.execute("""
                SELECT u.id, u.username, u.email, u.created_at, u.last_login,
                       cv.personal_info, cv.professional_summary, cv.education, 
                       cv.experience, cv.skills
                FROM users u
                LEFT JOIN user_cv_data cv ON u.id = cv.user_id
                WHERE cv.personal_info ILIKE %s 
                   OR cv.professional_summary ILIKE %s
                   OR cv.education ILIKE %s
                   OR cv.experience ILIKE %s
                   OR cv.skills ILIKE %s
            """, (f'%{search_term}%', f'%{search_term}%', f'%{search_term}%', f'%{search_term}%', f'%{search_term}%'))
        else:
            cursor.execute("""
                SELECT u.id, u.username, u.email, u.created_at, u.last_login,
                       cv.personal_info, cv.professional_summary, cv.education, 
                       cv.experience, cv.skills
                FROM users u
                LEFT JOIN user_cv_data cv ON u.id = cv.user_id
            """)
        
        results = cursor.fetchall()
        cursor.close()
        connection.close()
        
        # Crear DataFrame
        df = pd.DataFrame(results, columns=['ID', 'Usuario', 'Email', 'Fecha Registro', 'Último Login', 
                                          'Info Personal', 'Resumen Profesional', 'Educación', 'Experiencia', 'Habilidades'])
        
        # Crear archivo Excel en memoria
        output = BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Usuarios', index=False)
        
        output.seek(0)
        
        filename = f"usuarios_{search_term}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx" if search_term else f"todos_usuarios_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=filename
        )
        
    except ImportError:
        flash('pandas no está instalado. No se puede exportar a Excel.', 'error')
        return redirect(url_for('admin_database'))
    except Exception as e:
        print(f"Error exportando usuarios: {e}")
        flash('Error exportando usuarios', 'error')
        return redirect(url_for('admin_database'))

@app.route('/admin/users')
@admin_required
def admin_users():
    """Gestión de usuarios"""
    search = request.args.get('search', '')
    page = int(request.args.get('page', 1))
    per_page = 20
    
    connection = get_db_connection()
    if not connection:
        flash('Error de conexión a la base de datos', 'error')
        return redirect(url_for('admin_dashboard'))
    
    try:
        cursor = connection.cursor()
        
        # Construir consulta con búsqueda
        base_query = """
            SELECT id, username, email, role, is_banned, ban_until, ban_reason, 
                   last_login, created_at, email_verified
            FROM users
        """
        
        if search:
            base_query += " WHERE username ILIKE %s OR email ILIKE %s"
            search_param = f"%{search}%"
            cursor.execute(base_query + " ORDER BY created_at DESC LIMIT %s OFFSET %s", 
                         (search_param, search_param, per_page, (page-1)*per_page))
        else:
            cursor.execute(base_query + " ORDER BY created_at DESC LIMIT %s OFFSET %s", 
                         (per_page, (page-1)*per_page))
        
        users = cursor.fetchall()
        
        # Contar total para paginación
        if search:
            cursor.execute("SELECT COUNT(*) FROM users WHERE username ILIKE %s OR email ILIKE %s", 
                         (search_param, search_param))
        else:
            cursor.execute("SELECT COUNT(*) FROM users")
        
        total_users = cursor.fetchone()['count']
        total_pages = (total_users + per_page - 1) // per_page
        
        cursor.close()
        connection.close()
        
        return render_template('admin/users.html', 
                             users=users, 
                             search=search, 
                             page=page, 
                             total_pages=total_pages)
        
    except Exception as e:
        print(f"Error obteniendo usuarios: {e}")
        flash('Error obteniendo usuarios', 'error')
        return redirect(url_for('admin_dashboard'))

@app.route('/admin/ban_user', methods=['POST'])
@admin_required
def admin_ban_user():
    """Banear usuario"""
    data = request.get_json()
    user_id = data.get('user_id')
    ban_type = data.get('ban_type')  # 'permanent' o 'temporary'
    ban_duration = data.get('ban_duration')  # en días para bans temporales
    ban_reason = data.get('ban_reason', '')
    
    if not user_id:
        return jsonify({'success': False, 'message': 'ID de usuario requerido'})
    
    connection = get_db_connection()
    if not connection:
        return jsonify({'success': False, 'message': 'Error de conexión a la base de datos'})
    
    try:
        cursor = connection.cursor()
        
        if ban_type == 'permanent':
            cursor.execute("""
                UPDATE users 
                SET is_banned = TRUE, ban_until = NULL, ban_reason = %s
                WHERE id = %s
            """, (ban_reason, user_id))
        else:  # temporary
            from datetime import datetime, timedelta
            ban_until = datetime.now() + timedelta(days=int(ban_duration))
            cursor.execute("""
                UPDATE users 
                SET is_banned = TRUE, ban_until = %s, ban_reason = %s
                WHERE id = %s
            """, (ban_until, ban_reason, user_id))
        
        connection.commit()
        cursor.close()
        connection.close()
        
        return jsonify({'success': True, 'message': 'Usuario baneado correctamente'})
        
    except Exception as e:
        print(f"Error baneando usuario: {e}")
        return jsonify({'success': False, 'message': 'Error interno del servidor'})

@app.route('/admin/unban_user', methods=['POST'])
@admin_required
def admin_unban_user():
    """Desbanear usuario"""
    data = request.get_json()
    user_id = data.get('user_id')
    
    if not user_id:
        return jsonify({'success': False, 'message': 'ID de usuario requerido'})
    
    connection = get_db_connection()
    if not connection:
        return jsonify({'success': False, 'message': 'Error de conexión a la base de datos'})
    
    try:
        cursor = connection.cursor()
        
        cursor.execute("""
            UPDATE users 
            SET is_banned = FALSE, ban_until = NULL, ban_reason = NULL
            WHERE id = %s
        """, (user_id,))
        
        connection.commit()
        cursor.close()
        connection.close()
        
        return jsonify({'success': True, 'message': 'Usuario desbaneado correctamente'})
        
    except Exception as e:
        print(f"Error desbaneando usuario: {e}")
        return jsonify({'success': False, 'message': 'Error interno del servidor'})

@app.route('/admin/delete_user', methods=['POST'])
@admin_required
def admin_delete_user():
    """Eliminar usuario"""
    data = request.get_json()
    user_id = data.get('user_id')
    
    if not user_id:
        return jsonify({'success': False, 'message': 'ID de usuario requerido'})
    
    # Prevenir que el admin se elimine a sí mismo
    if int(user_id) == session['user_id']:
        return jsonify({'success': False, 'message': 'No puedes eliminar tu propia cuenta'})
    
    connection = get_db_connection()
    if not connection:
        return jsonify({'success': False, 'message': 'Error de conexión a la base de datos'})
    
    try:
        cursor = connection.cursor()
        
        # Eliminar usuario (las relaciones se eliminan en cascada)
        cursor.execute("DELETE FROM users WHERE id = %s", (user_id,))
        
        connection.commit()
        cursor.close()
        connection.close()
        
        return jsonify({'success': True, 'message': 'Usuario eliminado correctamente'})
        
    except Exception as e:
        print(f"Error eliminando usuario: {e}")
        return jsonify({'success': False, 'message': 'Error interno del servidor'})

@app.route('/admin/get_user_subscription/<int:user_id>')
@admin_required
def admin_get_user_subscription(user_id):
    """Obtener suscripción actual del usuario"""
    try:
        from subscription_system import get_user_subscription
        user_subscription = get_user_subscription(user_id)
        current_plan = user_subscription.get('plan_type', 'Sin suscripción') if user_subscription else 'Sin suscripción'
        
        return jsonify({
            'success': True,
            'current_plan': current_plan
        })
    except Exception as e:
        print(f"Error obteniendo suscripción: {e}")
        return jsonify({'success': False, 'message': 'Error obteniendo suscripción'})

@app.route('/admin/update_subscription', methods=['POST'])
@admin_required
def admin_update_subscription():
    """Actualizar suscripción de usuario"""
    data = request.get_json()
    user_id = data.get('user_id')
    new_plan = data.get('new_plan')
    reason = data.get('reason', '')
    
    if not user_id or not new_plan:
        return jsonify({'success': False, 'message': 'Datos requeridos faltantes'})
    
    # Corregir nombres de planes válidos
    if new_plan not in ['free_trial', 'standard', 'pro']:
        return jsonify({'success': False, 'message': 'Plan de suscripción inválido'})
    
    connection = get_db_connection()
    if not connection:
        return jsonify({'success': False, 'message': 'Error de conexión a la base de datos'})
    
    try:
        cursor = connection.cursor()
        
        # Verificar si el usuario existe
        cursor.execute("SELECT username FROM users WHERE id = %s", (user_id,))
        user = cursor.fetchone()
        if not user:
            return jsonify({'success': False, 'message': 'Usuario no encontrado'})
        
        # Verificar si ya tiene una suscripción activa
        cursor.execute("SELECT id FROM subscriptions WHERE user_id = %s AND status = 'active'", (user_id,))
        existing_subscription = cursor.fetchone()
        
        if existing_subscription:
            # Desactivar suscripción existente
            cursor.execute("""
                UPDATE subscriptions 
                SET status = 'cancelled'
                WHERE user_id = %s AND status = 'active'
            """, (user_id,))
        
        # Crear nueva suscripción
        from datetime import datetime, timedelta
        start_date = datetime.now()
        
        # Calcular fecha de vencimiento y precio según el plan
        if new_plan == 'free_trial':
            end_date = start_date + timedelta(days=7)
            amount = 0
        elif new_plan == 'standard':
            end_date = start_date + timedelta(days=30)
            amount = 9990
        elif new_plan == 'pro':
            end_date = start_date + timedelta(days=30)
            amount = 19990
            
        # Insertar nueva suscripción
        cursor.execute("""
            INSERT INTO subscriptions (user_id, plan_type, status, start_date, end_date, amount)
            VALUES (%s, %s, 'active', %s, %s, %s)
        """, (user_id, new_plan, start_date, end_date, amount))
        
        # Actualizar información del usuario
        cursor.execute("""
            UPDATE users 
            SET current_plan = %s, subscription_status = 'active', subscription_end_date = %s
            WHERE id = %s
        """, (new_plan, end_date, user_id))
        
        # Registrar el cambio en logs de administrador
        admin_username = session.get('username', 'Admin')
        log_message = f"Admin {admin_username} asignó suscripción {new_plan} al usuario {user['username']}. Razón: {reason}"
        add_console_log('INFO', log_message, 'ADMIN')
        
        connection.commit()
        cursor.close()
        connection.close()
        
        return jsonify({
            'success': True, 
            'message': f'Suscripción {new_plan} asignada exitosamente',
            'new_plan': new_plan,
            'end_date': end_date.strftime('%Y-%m-%d')
        })
        
    except Exception as e:
        print(f"Error actualizando suscripción: {e}")
        add_console_log('ERROR', f'Error asignando suscripción: {str(e)}', 'ADMIN')
        if connection:
            connection.rollback()
        return jsonify({'success': False, 'message': f'Error interno del servidor: {str(e)}'})

# Sistema de logs global para la consola de administrador
import logging
from collections import deque
from datetime import datetime
import threading

# Cola de logs para la consola (máximo 1000 entradas)
console_logs = deque(maxlen=1000)
logs_lock = threading.Lock()

def add_console_log(level, message, source='SYSTEM'):
    """Agregar un log a la consola de administrador"""
    with logs_lock:
        log_entry = {
            'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'level': level,
            'source': source,
            'message': str(message)
        }
        console_logs.append(log_entry)

# Configurar logging personalizado
class ConsoleLogHandler(logging.Handler):
    def emit(self, record):
        try:
            msg = self.format(record)
            level = record.levelname
            add_console_log(level, msg, 'APP')
        except Exception:
            pass

# Configurar el handler personalizado
console_handler = ConsoleLogHandler()
console_handler.setLevel(logging.INFO)
formatter = logging.Formatter('%(name)s - %(message)s')
console_handler.setFormatter(formatter)

# Agregar el handler al logger de Flask
app.logger.addHandler(console_handler)
app.logger.setLevel(logging.INFO)

# Manejador de errores global para capturar errores sin que sean fatales
@app.errorhandler(Exception)
def handle_exception(e):
    """Manejador global de excepciones para evitar errores fatales"""
    import traceback
    
    # Log del error en la consola de administrador
    error_msg = f"Error en {request.endpoint or 'unknown'}: {str(e)}"
    add_console_log('ERROR', error_msg, 'APP')
    
    # Log detallado para debugging
    app.logger.error(f"Excepción no manejada: {str(e)}")
    app.logger.error(f"Traceback: {traceback.format_exc()}")
    
    # Si es una petición AJAX, devolver JSON
    if request.is_json or 'application/json' in request.headers.get('Accept', ''):
        return jsonify({
            'success': False,
            'error': 'Error interno del servidor',
            'message': str(e) if app.debug else 'Ha ocurrido un error inesperado'
        }), 500
    
    # Para peticiones normales, mostrar página de error
    flash(f'Ha ocurrido un error: {str(e)}', 'error')
    return redirect(request.referrer or url_for('index'))

# Manejador específico para errores 404
@app.errorhandler(404)
def not_found_error(error):
    add_console_log('WARNING', f'Página no encontrada: {request.url}', 'APP')
    flash('Página no encontrada', 'error')
    return redirect(url_for('index'))

# Manejador específico para errores 500
@app.errorhandler(500)
def internal_error(error):
    add_console_log('ERROR', f'Error interno del servidor: {str(error)}', 'APP')
    flash('Error interno del servidor', 'error')
    return redirect(url_for('index'))

@app.route('/admin/console_logs')
@admin_required
def admin_console_logs():
    """API para obtener logs del servidor en tiempo real"""
    try:
        import psutil
        current_time = datetime.now()
        
        # Información del sistema en tiempo real
        try:
            cpu_percent = psutil.cpu_percent(interval=0.1)
            memory = psutil.virtual_memory()
            disk = psutil.disk_usage('/')
            
            # Agregar logs de sistema si no hay logs recientes
            if not console_logs or (current_time - datetime.strptime(console_logs[-1]['timestamp'], '%Y-%m-%d %H:%M:%S')).seconds > 5:
                add_console_log('INFO', f'Sistema - CPU: {cpu_percent:.1f}%, RAM: {memory.percent:.1f}%, Disco: {disk.percent:.1f}%', 'SYSTEM')
        except Exception as e:
            add_console_log('WARNING', f'Error obteniendo métricas del sistema: {str(e)}', 'SYSTEM')
        
        # Logs de actividad de base de datos
        try:
            conn = get_db_connection()
            if conn:
                cursor = conn.cursor()
            
                # Contar usuarios activos (últimos 5 minutos)
                cursor.execute("""
                    SELECT COUNT(*) as active_count FROM users 
                    WHERE last_login >= NOW() - INTERVAL '5 minutes'
                """)
                result = cursor.fetchone()
                active_users = result['active_count'] if result else 0
                
                # Últimos logins
                cursor.execute("""
                    SELECT email, last_login 
                    FROM users 
                    WHERE last_login IS NOT NULL 
                    ORDER BY last_login DESC 
                    LIMIT 5
                """)
                recent_logins = cursor.fetchall()
                
                add_console_log('INFO', f'Base de datos conectada - {active_users} usuarios activos', 'DATABASE')
                
                for login in recent_logins:
                    if login['last_login']:
                        add_console_log('INFO', f'Login de usuario: {login["email"]}', 'AUTH')
                
                cursor.close()
                conn.close()
                
        except Exception as db_error:
            add_console_log('ERROR', f'Error de base de datos: {str(db_error)}', 'DATABASE')
        
        # Logs de advertencias del sistema
        try:
            if 'memory' in locals() and memory.percent > 80:
                add_console_log('WARNING', f'Uso alto de memoria detectado: {memory.percent:.1f}%', 'SYSTEM')
            
            if 'cpu_percent' in locals() and cpu_percent > 80:
                add_console_log('WARNING', f'Uso alto de CPU detectado: {cpu_percent:.1f}%', 'SYSTEM')
        except Exception:
            pass
        
        # Convertir deque a lista para JSON
        with logs_lock:
            logs_list = list(console_logs)
        
        return jsonify({
            'success': True,
            'logs': logs_list[-50:],  # Últimos 50 logs
            'timestamp': current_time.strftime('%Y-%m-%d %H:%M:%S')
        })
        
    except Exception as e:
        add_console_log('ERROR', f'Error en consola de administrador: {str(e)}', 'ADMIN')
        return jsonify({
             'success': False,
             'error': str(e),
             'logs': []
         })

@app.route('/admin/check_s3')
def check_s3_connection():
    """Check S3 connection and create bucket if needed"""
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    # Check if user is admin
    connection = get_db_connection()
    if connection:
        cursor = connection.cursor()
        cursor.execute("SELECT role FROM users WHERE id = %s", (session['user_id'],))
        user = cursor.fetchone()
        cursor.close()
        connection.close()
        
        if not user or user['role'] != 'admin':
            return jsonify({'success': False, 'error': 'Access denied'})
    else:
        return jsonify({'success': False, 'error': 'Database connection failed'})
    
    try:
        from s3_utils import check_s3_connection, create_bucket_if_not_exists
        
        # Check S3 connection
        connection_result = check_s3_connection()
        if not connection_result['success']:
            return jsonify({
                'success': False,
                'error': f"S3 connection failed: {connection_result['error']}"
            })
        
        # Create bucket if it doesn't exist
        bucket_result = create_bucket_if_not_exists()
        if not bucket_result['success']:
            return jsonify({
                'success': False,
                'error': f"Failed to create S3 bucket: {bucket_result['error']}"
            })
        
        return jsonify({
            'success': True,
            'message': 'S3 connection successful and bucket is ready',
            'bucket_name': bucket_result.get('bucket_name', 'Unknown')
        })
        
    except ImportError:
        return jsonify({
            'success': False,
            'error': 'S3 utilities not available. Please check s3_utils.py'
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Unexpected error: {str(e)}'
        })

@app.route('/admin/content')
@admin_required
def admin_content():
    """Panel de edición de contenido del sitio"""
    username = session.get('username', 'unknown')
    add_console_log('INFO', f'Admin accedió a edición de contenido: {username}', 'ADMIN')
    
    connection = get_db_connection()
    if not connection:
        flash('Error de conexión a la base de datos', 'error')
        return redirect(url_for('admin_dashboard'))
    
    try:
        cursor = connection.cursor()
        
        # Obtener contenido del sitio
        cursor.execute("""
            SELECT section, content_key, content_value, updated_at 
            FROM site_content 
            ORDER BY section, content_key
        """)
        site_content = cursor.fetchall()
        
        # Obtener consejos del día
        cursor.execute("""
            SELECT id, title, description, icon, color, is_active, updated_at 
            FROM daily_tips 
            ORDER BY id
        """)
        daily_tips = cursor.fetchall()
        
        cursor.close()
        connection.close()
        
        return render_template('admin/content.html', 
                             site_content=site_content, 
                             daily_tips=daily_tips)
        
    except Exception as e:
        app.logger.error(f"Error obteniendo contenido: {e}")
        flash('Error obteniendo contenido del sitio', 'error')
        return redirect(url_for('admin_dashboard'))

@app.route('/admin/content/update', methods=['POST'])
@admin_required
def admin_update_content():
    """Actualizar contenido del sitio"""
    username = session.get('username', 'unknown')
    user_id = session.get('user_id')
    
    section = request.form.get('section')
    content_key = request.form.get('content_key')
    content_value = request.form.get('content_value')
    
    if not all([section, content_key, content_value]):
        return jsonify({'success': False, 'message': 'Datos incompletos'})
    
    connection = get_db_connection()
    if not connection:
        return jsonify({'success': False, 'message': 'Error de conexión a la base de datos'})
    
    try:
        cursor = connection.cursor()
        
        cursor.execute("""
            UPDATE site_content 
            SET content_value = %s, updated_at = CURRENT_TIMESTAMP, updated_by = %s 
            WHERE section = %s AND content_key = %s
        """, (content_value, user_id, section, content_key))
        
        if cursor.rowcount == 0:
            cursor.execute("""
                INSERT INTO site_content (section, content_key, content_value, updated_by) 
                VALUES (%s, %s, %s, %s)
            """, (section, content_key, content_value, user_id))
        
        connection.commit()
        cursor.close()
        connection.close()
        
        add_console_log('INFO', f'Admin {username} actualizó contenido: {section}.{content_key}', 'CONTENT')
        return jsonify({'success': True, 'message': 'Contenido actualizado correctamente'})
        
    except Exception as e:
        app.logger.error(f"Error actualizando contenido: {e}")
        return jsonify({'success': False, 'message': 'Error actualizando contenido'})

@app.route('/admin/update_content', methods=['POST'])
@admin_required
def update_content_inline():
    """Actualizar contenido desde edición directa de cuadros completos en dashboard"""
    if 'user_id' not in session or session.get('user_role') != 'admin':
        return jsonify({'success': False, 'message': 'No autorizado'})
    
    try:
        # Obtener datos JSON del request
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'message': 'No se recibieron datos'})
        
        section = data.get('element_id')
        
        # Debug logging para identificar el problema
        app.logger.info(f"Datos recibidos: {data}")
        app.logger.info(f"Section/element_id: {section} (tipo: {type(section)})")
        
        # Validar que section sea un string válido
        if not section or not isinstance(section, str):
            app.logger.error(f"Section inválido: {section}")
            return jsonify({'success': False, 'message': f'Identificador de sección inválido: {section}'})
        
        # Obtener todos los campos del cuadro
        title = data.get('title')
        description = data.get('description')
        icon = data.get('icon')
        icon_color = data.get('icon_color')
        button_text = data.get('button_text')
        button_icon = data.get('button_icon')
        subtitle = data.get('subtitle')
        section_title = data.get('section_title')
        section_icon = data.get('section_icon')
        section_icon_color = data.get('section_icon_color')
        stats_label1 = data.get('stats_label1')
        stats_label2 = data.get('stats_label2')
        
        if not section:
            return jsonify({'success': False, 'message': 'Falta el identificador de sección'})
        
        connection = get_db_connection()
        if not connection:
            return jsonify({'success': False, 'message': 'Error de conexión a la base de datos'})
        
        cursor = connection.cursor()
        username = session.get('username', 'unknown')
        user_id = session.get('user_id')
        
        # Manejar diferentes tipos de secciones
        if section.startswith('daily_tip_'):
            # Actualizar consejo del día específico
            try:
                tip_id = int(section.split('_')[-1])
                app.logger.info(f"Actualizando tip con ID: {tip_id}")
                
                # Verificar que el tip existe
                cursor.execute("SELECT id FROM daily_tips WHERE id = %s", (tip_id,))
                result = cursor.fetchone()
                
                if not result:
                    app.logger.error(f"No se encontró tip con ID: {tip_id}")
                    return jsonify({'success': False, 'message': f'No se encontró el consejo con ID: {tip_id}'})
                
                # Actualizar campos
                if title:
                    cursor.execute("UPDATE daily_tips SET title = %s WHERE id = %s", (title, tip_id))
                    app.logger.info(f"Título actualizado: {title}")
                
                if description:
                    cursor.execute("UPDATE daily_tips SET description = %s WHERE id = %s", (description, tip_id))
                    app.logger.info(f"Descripción actualizada: {description}")
                
                if icon:
                    cursor.execute("UPDATE daily_tips SET icon = %s WHERE id = %s", (icon, tip_id))
                    app.logger.info(f"Icono actualizado: {icon}")
                
                if icon_color:
                    # Limpiar el prefijo 'text-' si existe
                    color_value = icon_color.replace('text-', '') if icon_color.startswith('text-') else icon_color
                    cursor.execute("UPDATE daily_tips SET color = %s WHERE id = %s", (color_value, tip_id))
                    app.logger.info(f"Color actualizado: {color_value}")
                    
            except ValueError as e:
                app.logger.error(f"Error convirtiendo tip_id: {e}")
                return jsonify({'success': False, 'message': 'ID de consejo inválido'})
            except Exception as e:
                app.logger.error(f"Error actualizando daily_tip: {e}")
                return jsonify({'success': False, 'message': f'Error actualizando consejo: {str(e)}'})
        
        elif section.startswith('quick_action_'):
            # Actualizar acciones rápidas
            action_type = section.replace('quick_action_', '')
            
            # Función helper para insertar o actualizar quick actions
            def upsert_quick_action(content_key, content_value):
                try:
                    cursor.execute("""
                        UPDATE site_content 
                        SET content_value = %s, updated_at = CURRENT_TIMESTAMP, updated_by = %s 
                        WHERE section = %s AND content_key = %s
                    """, (content_value, user_id, 'quick_actions', content_key))
                    
                    if cursor.rowcount == 0:
                        cursor.execute("""
                            INSERT INTO site_content (section, content_key, content_value, updated_by) 
                            VALUES (%s, %s, %s, %s)
                        """, ('quick_actions', content_key, content_value, user_id))
                except Exception as e:
                    app.logger.error(f"Error en upsert_quick_action para {content_key}: {e}")
                    raise
            
            if title:
                upsert_quick_action(f'{action_type}_title', title)
            
            if description:
                upsert_quick_action(f'{action_type}_description', description)
            
            if icon:
                upsert_quick_action(f'{action_type}_icon', icon)
            
            if button_text:
                upsert_quick_action(f'{action_type}_button_text', button_text)
        
        elif section == 'welcome_section':
            # Actualizar sección de bienvenida
            def upsert_welcome(content_key, content_value):
                try:
                    cursor.execute("""
                        UPDATE site_content 
                        SET content_value = %s, updated_at = CURRENT_TIMESTAMP, updated_by = %s 
                        WHERE section = %s AND content_key = %s
                    """, (content_value, user_id, 'welcome', content_key))
                    
                    if cursor.rowcount == 0:
                        cursor.execute("""
                            INSERT INTO site_content (section, content_key, content_value, updated_by) 
                            VALUES (%s, %s, %s, %s)
                        """, ('welcome', content_key, content_value, user_id))
                except Exception as e:
                    app.logger.error(f"Error en upsert_welcome para {content_key}: {e}")
                    raise
            
            if title:
                upsert_welcome('title', title)
            
            if description:
                upsert_welcome('subtitle', description)
        
        elif section == 'tips_section':
            # Actualizar configuración general de la sección de consejos
            if title:
                try:
                    cursor.execute("""
                        UPDATE site_content 
                        SET content_value = %s, updated_at = CURRENT_TIMESTAMP, updated_by = %s 
                        WHERE section = %s AND content_key = %s
                    """, (title, user_id, 'tips', 'section_title'))
                    
                    if cursor.rowcount == 0:
                        cursor.execute("""
                            INSERT INTO site_content (section, content_key, content_value, updated_by) 
                            VALUES (%s, %s, %s, %s)
                        """, ('tips', 'section_title', title, user_id))
                except Exception as e:
                    app.logger.error(f"Error actualizando tips_section: {e}")
                    raise
        
        elif section.startswith('tip_'):
            # Actualizar consejos de análisis individuales
            tip_type = section  # tip_format, tip_keywords, etc.
            
            # Función helper para insertar o actualizar
            def upsert_content(content_key, content_value):
                try:
                    # Intentar actualizar primero
                    cursor.execute("""
                        UPDATE site_content 
                        SET content_value = %s, updated_at = CURRENT_TIMESTAMP, updated_by = %s 
                        WHERE section = %s AND content_key = %s
                    """, (content_value, user_id, 'analysis_tips', content_key))
                    
                    # Si no se actualizó ninguna fila, insertar
                    if cursor.rowcount == 0:
                        cursor.execute("""
                            INSERT INTO site_content (section, content_key, content_value, updated_by) 
                            VALUES (%s, %s, %s, %s)
                        """, ('analysis_tips', content_key, content_value, user_id))
                except Exception as e:
                    app.logger.error(f"Error en upsert_content para {content_key}: {e}")
                    raise
            
            if title:
                upsert_content(f'{tip_type}_title', title)
            
            if description:
                upsert_content(f'{tip_type}_description', description)
            
            if icon:
                upsert_content(f'{tip_type}_icon', icon)
            
            if icon_color:
                upsert_content(f'{tip_type}_icon_color', icon_color)
        
        connection.commit()
        cursor.close()
        connection.close()
        
        add_console_log('INFO', f'Cuadro actualizado por edición directa - {username}: {section}', 'ADMIN')
        
        return jsonify({'success': True, 'message': 'Contenido actualizado exitosamente'})
        
    except Exception as e:
        app.logger.error(f"Error actualizando contenido inline: {e}")
        return jsonify({'success': False, 'message': str(e)})

@app.route('/admin/tips/add', methods=['POST'])
@admin_required
def admin_add_tip():
    """Agregar nuevo consejo del día"""
    username = session.get('username', 'unknown')
    user_id = session.get('user_id')
    
    data = request.get_json()
    if not data:
        return jsonify({'success': False, 'message': 'No se recibieron datos'})
    
    title = data.get('title', 'Nuevo Consejo')
    description = data.get('description', 'Descripción del consejo')
    icon = data.get('icon', 'fas fa-lightbulb')
    color = data.get('color', 'primary')
    
    if not all([title, description]):
        return jsonify({'success': False, 'message': 'Título y descripción son requeridos'})
    
    connection = get_db_connection()
    if not connection:
        return jsonify({'success': False, 'message': 'Error de conexión a la base de datos'})
    
    try:
        cursor = connection.cursor()
        
        cursor.execute("""
            INSERT INTO daily_tips (title, description, icon, color, updated_by) 
            VALUES (%s, %s, %s, %s, %s)
        """, (title, description, icon, color, user_id))
        
        connection.commit()
        cursor.close()
        connection.close()
        
        add_console_log('INFO', f'Admin {username} agregó nuevo consejo: {title}', 'CONTENT')
        return jsonify({'success': True, 'message': 'Consejo agregado correctamente'})
        
    except Exception as e:
        app.logger.error(f"Error agregando consejo: {e}")
        return jsonify({'success': False, 'message': 'Error agregando consejo'})

@app.route('/admin/tips/update', methods=['POST'])
@admin_required
def admin_update_tip():
    """Actualizar consejo del día"""
    username = session.get('username', 'unknown')
    user_id = session.get('user_id')
    
    tip_id = request.form.get('tip_id')
    title = request.form.get('title')
    description = request.form.get('description')
    icon = request.form.get('icon')
    color = request.form.get('color')
    is_active = request.form.get('is_active') == 'true'
    
    if not all([tip_id, title, description]):
        return jsonify({'success': False, 'message': 'Datos incompletos'})
    
    connection = get_db_connection()
    if not connection:
        return jsonify({'success': False, 'message': 'Error de conexión a la base de datos'})
    
    try:
        cursor = connection.cursor()
        
        cursor.execute("""
            UPDATE daily_tips 
            SET title = %s, description = %s, icon = %s, color = %s, 
                is_active = %s, updated_at = CURRENT_TIMESTAMP, updated_by = %s 
            WHERE id = %s
        """, (title, description, icon, color, is_active, user_id, tip_id))
        
        connection.commit()
        cursor.close()
        connection.close()
        
        add_console_log('INFO', f'Admin {username} actualizó consejo ID {tip_id}', 'CONTENT')
        return jsonify({'success': True, 'message': 'Consejo actualizado correctamente'})
        
    except Exception as e:
        app.logger.error(f"Error actualizando consejo: {e}")
        return jsonify({'success': False, 'message': 'Error actualizando consejo'})

@app.route('/admin/daily_tips/update', methods=['POST'])
@admin_required
def admin_update_daily_tip():
    """Actualizar consejo del día desde modal"""
    username = session.get('username', 'unknown')
    user_id = session.get('user_id')
    
    data = request.get_json()
    if not data:
        return jsonify({'success': False, 'message': 'No se recibieron datos'})
    
    section = data.get('section')
    title = data.get('title')
    description = data.get('description')
    icon = data.get('icon')
    icon_color = data.get('icon_color')
    
    # Extraer el ID del consejo del section (formato: daily_tip_X)
    if not section or not section.startswith('daily_tip_'):
        return jsonify({'success': False, 'message': 'Sección inválida'})
    
    try:
        tip_id = int(section.split('_')[2])
    except (IndexError, ValueError):
        return jsonify({'success': False, 'message': 'ID de consejo inválido'})
    
    if not all([title, description]):
        return jsonify({'success': False, 'message': 'Título y descripción son requeridos'})
    
    connection = get_db_connection()
    if not connection:
        return jsonify({'success': False, 'message': 'Error de conexión a la base de datos'})
    
    try:
        cursor = connection.cursor()
        
        cursor.execute("""
            UPDATE daily_tips 
            SET title = %s, description = %s, icon = %s, color = %s, 
                updated_at = CURRENT_TIMESTAMP, updated_by = %s 
            WHERE id = %s
        """, (title, description, icon, icon_color, user_id, tip_id))
        
        connection.commit()
        cursor.close()
        connection.close()
        
        add_console_log('INFO', f'Admin {username} actualizó consejo del día ID {tip_id}', 'CONTENT')
        return jsonify({'success': True, 'message': 'Consejo actualizado correctamente'})
        
    except Exception as e:
        app.logger.error(f"Error actualizando consejo del día: {e}")
        return jsonify({'success': False, 'message': 'Error actualizando consejo'})

@app.route('/admin/tips/update_inline', methods=['POST'])
@admin_required
def admin_update_tip_inline():
    """Actualizar consejo del día en línea"""
    username = session.get('username', 'unknown')
    user_id = session.get('user_id')
    
    data = request.get_json()
    if not data:
        return jsonify({'success': False, 'message': 'No se recibieron datos'})
    
    tip_id = data.get('tip_id')
    field = data.get('field')
    value = data.get('value')
    
    # Validación específica por campo
    if tip_id is None or not field:
        return jsonify({'success': False, 'message': 'Datos incompletos'})
    
    if field == 'icon':
        icon = data.get('icon')
        color = data.get('color')
        if not icon or not color:
            return jsonify({'success': False, 'message': 'Icono y color requeridos'})
    elif not value:
        return jsonify({'success': False, 'message': 'Valor requerido'})
    
    connection = get_db_connection()
    if not connection:
        return jsonify({'success': False, 'message': 'Error de conexión a la base de datos'})
    
    try:
        cursor = connection.cursor()
        
        # Usar directamente el ID real del consejo
        real_tip_id = int(tip_id)
        
        # Actualizar el campo específico
        if field == 'title':
            cursor.execute("UPDATE daily_tips SET title = %s, updated_at = CURRENT_TIMESTAMP, updated_by = %s WHERE id = %s", 
                         (value, user_id, real_tip_id))
        elif field == 'description':
            cursor.execute("UPDATE daily_tips SET description = %s, updated_at = CURRENT_TIMESTAMP, updated_by = %s WHERE id = %s", 
                         (value, user_id, real_tip_id))
        elif field == 'icon':
            icon = data.get('icon')
            color = data.get('color')
            cursor.execute("UPDATE daily_tips SET icon = %s, color = %s, updated_at = CURRENT_TIMESTAMP, updated_by = %s WHERE id = %s", 
                         (icon, color, user_id, real_tip_id))
        else:
            return jsonify({'success': False, 'message': 'Campo no válido'})
        
        connection.commit()
        cursor.close()
        connection.close()
        
        add_console_log('INFO', f'Admin {username} actualizó consejo ID {real_tip_id} campo {field}', 'CONTENT')
        return jsonify({'success': True, 'message': 'Consejo actualizado correctamente'})
        
    except Exception as e:
        app.logger.error(f"Error actualizando consejo en línea: {e}")
        return jsonify({'success': False, 'message': 'Error actualizando consejo'})

@app.route('/admin/tips/delete', methods=['POST'])
@admin_required
def admin_delete_tip():
    """Eliminar consejo del día"""
    username = session.get('username', 'unknown')
    
    data = request.get_json()
    if not data:
        return jsonify({'success': False, 'message': 'No se recibieron datos'})
    
    tip_id = data.get('tip_id')
    
    if not tip_id:
        return jsonify({'success': False, 'message': 'ID de consejo requerido'})
    
    connection = get_db_connection()
    if not connection:
        return jsonify({'success': False, 'message': 'Error de conexión a la base de datos'})
    
    try:
        cursor = connection.cursor()
        
        # Usar directamente el ID real del consejo
        real_tip_id = int(tip_id)
        
        cursor.execute("DELETE FROM daily_tips WHERE id = %s", (real_tip_id,))
        
        connection.commit()
        cursor.close()
        connection.close()
        
        add_console_log('INFO', f'Admin {username} eliminó consejo ID {tip_id}', 'CONTENT')
        return jsonify({'success': True, 'message': 'Consejo eliminado correctamente'})
        
    except Exception as e:
        app.logger.error(f"Error eliminando consejo: {e}")
        return jsonify({'success': False, 'message': 'Error eliminando consejo'})

# ==================== RUTAS DEL BLOG DE TIPS Y SUGERENCIAS ====================

@app.route('/tips-sugerencias')
def blog_tips():
    """Mostrar blog de tips y sugerencias para usuarios"""
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    connection = get_db_connection()
    if not connection:
        flash('Error de conexión a la base de datos', 'error')
        return redirect(url_for('dashboard'))
    
    try:
        cursor = connection.cursor(cursor_factory=RealDictCursor)
        
        # Obtener posts publicados con información del autor y conteo de reacciones
        cursor.execute("""
            SELECT bp.*, u.username as author_name,
                   COUNT(br.id) as total_reactions
            FROM blog_posts bp
            LEFT JOIN users u ON bp.author_id = u.id
            LEFT JOIN blog_reactions br ON bp.id = br.post_id
            WHERE bp.is_published = TRUE
            GROUP BY bp.id, u.username
            ORDER BY bp.created_at DESC
        """)
        posts = cursor.fetchall()
        
        # Obtener reacciones del usuario actual para cada post
        user_reactions = {}
        if posts:
            post_ids = [post['id'] for post in posts]
            cursor.execute("""
                SELECT post_id, emoji
                FROM blog_reactions
                WHERE user_id = %s AND post_id = ANY(%s)
            """, (session['user_id'], post_ids))
            user_reactions = {row['post_id']: row['emoji'] for row in cursor.fetchall()}
        
        # Obtener conteo de reacciones por emoji para cada post
        reactions_count = {}
        if posts:
            cursor.execute("""
                SELECT post_id, emoji, COUNT(*) as count
                FROM blog_reactions
                WHERE post_id = ANY(%s)
                GROUP BY post_id, emoji
            """, (post_ids,))
            for row in cursor.fetchall():
                if row['post_id'] not in reactions_count:
                    reactions_count[row['post_id']] = {}
                reactions_count[row['post_id']][row['emoji']] = row['count']
        
        cursor.close()
        connection.close()
        
        return render_template('blog_tips.html', 
                             posts=posts, 
                             user_reactions=user_reactions,
                             reactions_count=reactions_count)
        
    except Exception as e:
        app.logger.error(f"Error cargando blog: {e}")
        flash('Error cargando el blog', 'error')
        return redirect(url_for('dashboard'))

@app.route('/tips-sugerencias/react', methods=['POST'])
def blog_react():
    """Agregar o quitar reacción a un post del blog"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'No autorizado'})
    
    data = request.get_json()
    post_id = data.get('post_id')
    emoji = data.get('emoji')
    
    if not post_id or not emoji:
        return jsonify({'success': False, 'message': 'Datos incompletos'})
    
    connection = get_db_connection()
    if not connection:
        return jsonify({'success': False, 'message': 'Error de conexión'})
    
    try:
        cursor = connection.cursor()
        
        # Verificar si ya existe la reacción
        cursor.execute("""
            SELECT id FROM blog_reactions 
            WHERE post_id = %s AND user_id = %s AND emoji = %s
        """, (post_id, session['user_id'], emoji))
        
        existing = cursor.fetchone()
        
        if existing:
            # Quitar reacción
            cursor.execute("""
                DELETE FROM blog_reactions 
                WHERE post_id = %s AND user_id = %s AND emoji = %s
            """, (post_id, session['user_id'], emoji))
            action = 'removed'
        else:
            # Agregar reacción (primero quitar cualquier otra reacción del usuario en este post)
            cursor.execute("""
                DELETE FROM blog_reactions 
                WHERE post_id = %s AND user_id = %s
            """, (post_id, session['user_id']))
            
            cursor.execute("""
                INSERT INTO blog_reactions (post_id, user_id, emoji)
                VALUES (%s, %s, %s)
            """, (post_id, session['user_id'], emoji))
            action = 'added'
        
        # Obtener nuevo conteo de reacciones
        cursor.execute("""
            SELECT emoji, COUNT(*) as count
            FROM blog_reactions
            WHERE post_id = %s
            GROUP BY emoji
        """, (post_id,))
        reactions = {row[0]: row[1] for row in cursor.fetchall()}
        
        connection.commit()
        cursor.close()
        connection.close()
        
        return jsonify({
            'success': True, 
            'action': action,
            'reactions': reactions
        })
        
    except Exception as e:
        app.logger.error(f"Error procesando reacción: {e}")
        return jsonify({'success': False, 'message': 'Error procesando reacción'})

# ==================== RUTAS DE ADMINISTRACIÓN DEL BLOG ====================

@app.route('/admin/blog')
@admin_required
def admin_blog():
    """Panel de administración del blog"""
    connection = get_db_connection()
    if not connection:
        flash('Error de conexión a la base de datos', 'error')
        return redirect(url_for('admin_dashboard'))
    
    try:
        cursor = connection.cursor(cursor_factory=RealDictCursor)
        
        # Obtener todos los posts con información del autor
        cursor.execute("""
            SELECT bp.*, u.username as author_name
            FROM blog_posts bp
            LEFT JOIN users u ON bp.author_id = u.id
            ORDER BY bp.created_at DESC
        """)
        posts = cursor.fetchall()
        
        cursor.close()
        connection.close()
        
        return render_template('admin/blog.html', posts=posts)
        
    except Exception as e:
        app.logger.error(f"Error cargando admin blog: {e}")
        flash('Error cargando el panel de blog', 'error')
        return redirect(url_for('admin_dashboard'))

@app.route('/admin/blog/create', methods=['GET', 'POST'])
@admin_required
def admin_blog_create():
    """Crear nuevo post del blog"""
    if request.method == 'POST':
        title = request.form.get('title')
        content = request.form.get('content')
        image_option = request.form.get('image_option', 'url')
        image_url = request.form.get('image_url')
        is_published = 'is_published' in request.form
        
        if not title or not content:
            flash('Título y contenido son requeridos', 'error')
            return render_template('admin/blog_create.html')
        
        # Manejar imagen
        final_image_url = None
        if image_option == 'url' and image_url:
            final_image_url = image_url
        elif image_option == 'upload' and 'image_file' in request.files:
            image_file = request.files['image_file']
            if image_file and image_file.filename:
                try:
                    # Validar archivo
                    if not image_file.content_type.startswith('image/'):
                        flash('El archivo debe ser una imagen válida', 'error')
                        return render_template('admin/blog_create.html')
                    
                    # Validar tamaño (5MB máximo)
                    image_file.seek(0, 2)  # Ir al final del archivo
                    file_size = image_file.tell()
                    image_file.seek(0)  # Volver al inicio
                    
                    if file_size > 5 * 1024 * 1024:  # 5MB
                        flash('La imagen es demasiado grande. Máximo 5MB', 'error')
                        return render_template('admin/blog_create.html')
                    
                    # Guardar imagen en base de datos
                    final_image_url = save_image_to_database(image_file)
                    if not final_image_url:
                        flash('Error al guardar la imagen', 'error')
                        return render_template('admin/blog_create.html')
                        
                except Exception as e:
                    app.logger.error(f"Error procesando imagen: {e}")
                    flash('Error procesando la imagen', 'error')
                    return render_template('admin/blog_create.html')
        
        connection = get_db_connection()
        if not connection:
            flash('Error de conexión a la base de datos', 'error')
            return render_template('admin/blog_create.html')
        
        try:
            cursor = connection.cursor()
            
            cursor.execute("""
                INSERT INTO blog_posts (title, content, image_url, author_id, is_published)
                VALUES (%s, %s, %s, %s, %s)
            """, (title, content, final_image_url, session['user_id'], is_published))
            
            connection.commit()
            cursor.close()
            connection.close()
            
            username = session.get('username', 'unknown')
            add_console_log('INFO', f'Admin {username} creó nuevo post: {title}', 'CONTENT')
            flash('Post creado exitosamente', 'success')
            return redirect(url_for('admin_blog'))
            
        except Exception as e:
            app.logger.error(f"Error creando post: {e}")
            flash('Error creando el post', 'error')
    
    return render_template('admin/blog_create.html')

@app.route('/admin/blog/edit/<int:post_id>', methods=['GET', 'POST'])
@admin_required
def admin_blog_edit(post_id):
    """Editar post del blog"""
    connection = get_db_connection()
    if not connection:
        flash('Error de conexión a la base de datos', 'error')
        return redirect(url_for('admin_blog'))
    
    try:
        cursor = connection.cursor(cursor_factory=RealDictCursor)
        
        if request.method == 'POST':
            title = request.form.get('title')
            content = request.form.get('content')
            image_url = request.form.get('image_url')
            is_published = 'is_published' in request.form
            
            if not title or not content:
                flash('Título y contenido son requeridos', 'error')
                return redirect(url_for('admin_blog_edit', post_id=post_id))
            
            cursor.execute("""
                UPDATE blog_posts 
                SET title = %s, content = %s, image_url = %s, is_published = %s, updated_at = CURRENT_TIMESTAMP
                WHERE id = %s
            """, (title, content, image_url, is_published, post_id))
            
            connection.commit()
            cursor.close()
            connection.close()
            
            username = session.get('username', 'unknown')
            add_console_log('INFO', f'Admin {username} editó post ID {post_id}', 'CONTENT')
            flash('Post actualizado exitosamente', 'success')
            return redirect(url_for('admin_blog'))
        
        # GET request - mostrar formulario de edición
        cursor.execute("SELECT * FROM blog_posts WHERE id = %s", (post_id,))
        post = cursor.fetchone()
        
        if not post:
            flash('Post no encontrado', 'error')
            return redirect(url_for('admin_blog'))
        
        cursor.close()
        connection.close()
        
        return render_template('admin/blog_edit.html', post=post)
        
    except Exception as e:
        app.logger.error(f"Error editando post: {e}")
        flash('Error editando el post', 'error')
        return redirect(url_for('admin_blog'))

@app.route('/admin/blog/delete/<int:post_id>', methods=['POST'])
@admin_required
def admin_blog_delete(post_id):
    """Eliminar post del blog"""
    connection = get_db_connection()
    if not connection:
        return jsonify({'success': False, 'message': 'Error de conexión'})
    
    try:
        cursor = connection.cursor()
        
        # Obtener información del post antes de eliminarlo
        cursor.execute("SELECT title, image_url FROM blog_posts WHERE id = %s", (post_id,))
        post = cursor.fetchone()
        
        if not post:
            return jsonify({'success': False, 'message': 'Post no encontrado'})
        
        title, image_url = post
        
        # Si la imagen es una imagen subida (formato /image/ID), eliminarla de la BD
        if image_url and image_url.startswith('/image/'):
            try:
                image_id = int(image_url.split('/')[-1])
                cursor.execute("DELETE FROM uploaded_images WHERE id = %s", (image_id,))
            except (ValueError, IndexError):
                # Si no se puede extraer el ID, continuar sin eliminar la imagen
                pass
        
        # Eliminar el post
        cursor.execute("DELETE FROM blog_posts WHERE id = %s", (post_id,))
        
        connection.commit()
        cursor.close()
        connection.close()
        
        username = session.get('username', 'unknown')
        add_console_log('INFO', f'Admin {username} eliminó post: {title}', 'CONTENT')
        return jsonify({'success': True, 'message': 'Post eliminado exitosamente'})
        
    except Exception as e:
        app.logger.error(f"Error eliminando post: {e}")
        return jsonify({'success': False, 'message': 'Error eliminando el post'})

# ==================== GENERADOR DE CARTAS DE PRESENTACIÓN ====================

@app.route('/cover_letter_generator')
@login_required
def cover_letter_generator():
    """Generador de cartas de presentación con IA"""
    # Verificar restricciones de suscripción
    user_id = session.get('user_id')
    can_generate, message = check_user_limits(user_id, 'cover_letter_generation')
    
    if not can_generate:
        flash(f'Restricción de plan: {message}', 'error')
        return redirect(url_for('dashboard'))
    
    return render_template('cover_letter_generator.html')

@app.route('/generate_cover_letter', methods=['POST'])
@login_required
def generate_cover_letter():
    """Generar carta de presentación con IA"""
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    # Verificar restricciones de suscripción
    user_id = session.get('user_id')
    can_generate, message = check_user_limits(user_id, 'cover_letter_generation')
    
    if not can_generate:
        return jsonify({'success': False, 'message': f'Restricción de plan: {message}'})
    
    try:
        # Obtener datos del formulario
        cv_file = request.files.get('cv_file')
        job_title = request.form.get('job_title', '').strip()
        company_name = request.form.get('company_name', '').strip()
        job_description = request.form.get('job_description', '').strip()
        language = request.form.get('language', 'es')
        
        if not cv_file or cv_file.filename == '':
            return jsonify({'success': False, 'message': 'Debe subir un archivo de CV'})
        
        if not job_title:
            return jsonify({'success': False, 'message': 'El título del puesto es requerido'})
        
        # Validar archivo
        if not allowed_file(cv_file.filename):
            return jsonify({'success': False, 'message': 'Tipo de archivo no permitido. Solo se permiten PDF, DOC y DOCX.'})
        
        # Extraer texto del CV
        filename = secure_filename(cv_file.filename)
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            filepath = temp_file.name
            cv_file.save(filepath)
            
            original_extension = filename.rsplit('.', 1)[1].lower() if '.' in filename else None
            cv_text = extract_text_from_file(filepath, original_extension)
            
            # Limpiar archivo temporal
            try:
                os.unlink(filepath)
            except PermissionError:
                pass
        
        if not cv_text:
            return jsonify({'success': False, 'message': 'No se pudo extraer texto del CV'})
        
        # Generar carta de presentación con IA
        print(f"DEBUG: Generating cover letter for {job_title} at {company_name}")  # Debug
        cover_letter = generate_cover_letter_with_ai(
            cv_text=cv_text,
            job_title=job_title,
            company_name=company_name,
            job_description=job_description,
            language=language
        )
        
        print(f"DEBUG: Generated cover letter length: {len(cover_letter) if cover_letter else 0}")  # Debug
        print(f"DEBUG: Cover letter preview: {cover_letter[:200] if cover_letter else 'None'}...")  # Debug
        
        if not cover_letter:
            return jsonify({'success': False, 'message': 'Error generando la carta de presentación'})
        
        # Guardar en base de datos
        connection = get_db_connection()
        if connection:
            cursor = connection.cursor()
            cursor.execute("""
                INSERT INTO cover_letters (user_id, job_title, company_name, content, language, created_at)
                VALUES (%s, %s, %s, %s, %s, %s)
                RETURNING id
            """, (user_id, job_title, company_name, cover_letter, language, datetime.now()))
            
            cover_letter_id = cursor.fetchone()['id']
            connection.commit()
            cursor.close()
            connection.close()
            
            # Incrementar contador de uso
            increment_usage(user_id, 'cover_letter_generation')
            
            username = session.get('username', 'unknown')
            add_console_log('INFO', f'Usuario {username} generó carta para: {job_title} en {company_name}', 'COVER_LETTER')
            
            return jsonify({
                'success': True, 
                'cover_letter': cover_letter,
                'cover_letter_id': cover_letter_id
            })
        
        return jsonify({'success': False, 'message': 'Error guardando la carta de presentación'})
        
    except Exception as e:
        app.logger.error(f"Error generando carta de presentación: {e}")
        return jsonify({'success': False, 'message': 'Error interno del servidor'})

@app.route('/download_cover_letter/<int:cover_letter_id>')
@login_required
def download_cover_letter(cover_letter_id):
    """Descargar carta de presentación en PDF"""
    try:
        connection = get_db_connection()
        if not connection:
            flash('Error de conexión a la base de datos', 'error')
            return redirect(url_for('cover_letter_generator'))
        
        cursor = connection.cursor()
        cursor.execute("""
            SELECT job_title, company_name, content, language, created_at
            FROM cover_letters 
            WHERE id = %s AND user_id = %s
        """, (cover_letter_id, session['user_id']))
        
        cover_letter_data = cursor.fetchone()
        cursor.close()
        connection.close()
        
        if not cover_letter_data:
            flash('Carta de presentación no encontrada', 'error')
            return redirect(url_for('cover_letter_generator'))
        
        # Extraer datos del diccionario (RealDictCursor devuelve un diccionario)
        job_title = cover_letter_data['job_title']
        company_name = cover_letter_data['company_name']
        content = cover_letter_data['content']
        language = cover_letter_data['language']
        created_at = cover_letter_data['created_at']
        
        # Configuración de idiomas para el PDF
        language_config = {
            'es': {'title': 'Carta de Presentación', 'closing': 'Atentamente'},
            'en': {'title': 'Cover Letter', 'closing': 'Sincerely'},
            'pt': {'title': 'Carta de Apresentação', 'closing': 'Atenciosamente'},
            'de': {'title': 'Anschreiben', 'closing': 'Mit freundlichen Grüßen'},
            'fr': {'title': 'Lettre de Motivation', 'closing': 'Cordialement'}
        }
        
        lang_config = language_config.get(language, language_config['es'])
        
        # Generar PDF
        # Procesar el contenido para preservar los saltos de línea
        # Dividir el contenido en párrafos y formatear cada uno
        paragraphs = content.split('\n')
        formatted_paragraphs = []
        
        for para in paragraphs:
            para = para.strip()
            if para:  # Solo agregar párrafos no vacíos
                formatted_paragraphs.append(f'<p>{para}</p>')
        
        formatted_content = '\n'.join(formatted_paragraphs)
        
        print(f"DEBUG: Language: {language}")  # Debug
        print(f"DEBUG: Original content: {content[:200]}...")  # Debug
        print(f"DEBUG: Formatted content: {formatted_content[:200]}...")  # Debug
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                body {{ font-family: Arial, sans-serif; line-height: 1.6; margin: 40px; color: #333; }}
                .header {{ text-align: center; margin-bottom: 30px; border-bottom: 2px solid #333; padding-bottom: 20px; }}
                .content {{ text-align: justify; margin-bottom: 30px; }}
                .content p {{ margin-bottom: 15px; }}
                .signature {{ margin-top: 30px; }}
            </style>
        </head>
        <body>
            <div class="header">
                <h1>{lang_config['title']}</h1>
                <p><strong>{job_title}</strong> - {company_name}</p>
            </div>
            <div class="content">
                {formatted_content}
            </div>
            <div class="signature">
                <p>{lang_config['closing']},<br>
                {session.get('username', 'Candidato')}</p>
            </div>
        </body>
        </html>
        """
        
        # Debug: Imprimir el HTML completo
        print(f"DEBUG: Complete HTML: {html_content}")
        
        # Generar PDF usando ReportLab
        pdf_content = generate_pdf_with_reportlab(html_content, f"Carta_{job_title}_{company_name}")
        
        if pdf_content:
            filename = f"carta_presentacion_{job_title}_{company_name}.pdf".replace(' ', '_')
            
            response = make_response(pdf_content)
            response.headers['Content-Type'] = 'application/pdf'
            response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
            
            return response
        else:
            flash('Error generando el PDF', 'error')
            return redirect(url_for('cover_letter_generator'))
            
    except Exception as e:
        app.logger.error(f"Error descargando carta: {e}")
        flash('Error descargando la carta de presentación', 'error')
        return redirect(url_for('cover_letter_generator'))

def generate_cover_letter_with_ai(cv_text, job_title, company_name, job_description, language='es'):
    """Generar carta de presentación usando OpenAI siguiendo principios de Harvard"""
    try:
        # Mapeo de idiomas
        language_prompts = {
            'es': {
                'lang_name': 'español',
                'greeting': 'Estimado/a',
                'closing': 'Atentamente'
            },
            'en': {
                'lang_name': 'inglés',
                'greeting': 'Dear',
                'closing': 'Sincerely'
            },
            'pt': {
                'lang_name': 'portugués',
                'greeting': 'Prezado/a',
                'closing': 'Atenciosamente'
            },
            'de': {
                'lang_name': 'alemán',
                'greeting': 'Sehr geehrte/r',
                'closing': 'Mit freundlichen Grüßen'
            },
            'fr': {
                'lang_name': 'francés',
                'greeting': 'Cher/Chère',
                'closing': 'Cordialement'
            }
        }
        
        lang_config = language_prompts.get(language, language_prompts['es'])
        
        # Prompt siguiendo principios de Harvard
        prompt = f"""
Eres un experto en redacción de cartas de presentación siguiendo los principios de Harvard Business School. 

Genera una carta de presentación profesional en {lang_config['lang_name']} que siga estos principios clave:

1. DIRIGIDA Y PERSONALIZADA: Dirígela específicamente al puesto y empresa
2. CONECTA CON LA EMPRESA: Demuestra comprensión de la misión y valores
3. ENFOCADA EN EL VALOR: Explica qué puedes hacer por ellos, no solo lo que has hecho
4. MUESTRA ENTUSIASMO: Transmite interés genuino
5. BREVE Y CONCISA: Máximo una página
6. LENGUAJE DE ACCIÓN: Describe logros de manera impactante
7. NO REPITAS EL CV: Complementa, no dupliques

INFORMACIÓN:
- Puesto: {job_title}
- Empresa: {company_name}
- Descripción del puesto: {job_description if job_description else 'No proporcionada'}

CV DEL CANDIDATO:
{cv_text[:3000]}  # Limitar texto para evitar tokens excesivos

Genera una carta de presentación que:
- Use un saludo apropiado ({lang_config['greeting']})
- Tenga 3-4 párrafos bien estructurados
- Conecte específicamente con {company_name}
- Destaque 2-3 logros clave del CV que sean relevantes para {job_title}
- Explique el valor específico que aportaría a la empresa
- Termine con {lang_config['closing']}
- Sea profesional pero con personalidad
- No exceda 400 palabras

Formato: Solo el contenido de la carta, sin encabezados adicionales.
"""
        
        if OPENAI_CLIENT:
            print(f"DEBUG: Calling OpenAI API for cover letter generation")  # Debug
            response = OPENAI_CLIENT.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "Eres un experto en redacción de cartas de presentación profesionales siguiendo los estándares de Harvard Business School."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=800,
                temperature=0.7
            )
            
            generated_content = response.choices[0].message.content.strip()
            print(f"DEBUG: OpenAI response length: {len(generated_content)}")  # Debug
            print(f"DEBUG: OpenAI response preview: {generated_content[:200]}...")  # Debug
            return generated_content
        else:
            print(f"DEBUG: OpenAI client not available")  # Debug
            return None
            
    except Exception as e:
        app.logger.error(f"Error en generate_cover_letter_with_ai: {e}")
        return None

@app.route('/compare_cv_job', methods=['GET', 'POST'])
def compare_cv_job():
    """Comparador de CV con ofertas laborales usando IA"""
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    # Verificar restricciones de suscripción
    user_id = session.get('user_id')
    can_analyze, message = check_user_limits(user_id, 'cv_analysis')
    
    if not can_analyze:
        flash(f'Restricción de plan: {message}', 'error')
        return redirect(url_for('dashboard'))
    
    username = session.get('username', 'unknown')
    
    if request.method == 'POST':
        add_console_log('INFO', f'Usuario inició comparación CV-Oferta: {username}', 'CV_COMPARE')
        
        # Verificar que se subió un archivo
        if 'file' not in request.files:
            flash('No se seleccionó ningún archivo', 'error')
            return redirect(request.url)
        
        file = request.files['file']
        job_description = request.form.get('job_description', '').strip()
        
        # Validaciones
        if file.filename == '':
            flash('No se seleccionó ningún archivo', 'error')
            return redirect(request.url)
            
        if not job_description:
            flash('Debe ingresar la descripción de la oferta laboral', 'error')
            return redirect(request.url)
        
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            
            # Usar archivo temporal
            with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                filepath = temp_file.name
                file.save(filepath)
                
                # Extraer texto del CV
                original_extension = filename.rsplit('.', 1)[1].lower() if '.' in filename else None
                cv_text = extract_text_from_file(filepath, original_extension)
                
                # Eliminar archivo temporal
                try:
                    os.unlink(filepath)
                except PermissionError:
                    pass
                
                if cv_text:
                    # Realizar comparación con IA
                    comparison_result = compare_cv_with_job_ai(cv_text, job_description)
                    
                    # Asegurar que siempre tengamos un resultado válido
                    if not comparison_result:
                        print("WARNING: No se obtuvo resultado del análisis IA, usando valores por defecto")
                        add_console_log('WARNING', f'Análisis IA falló para usuario {username}, usando valores por defecto', 'CV_COMPARE')
                        comparison_result = {
                            'match_percentage': 0,
                            'strengths': ['No se pudieron identificar fortalezas específicas'],
                            'keywords_found': ['No se encontraron palabras clave'],
                            'improvements': ['Revisar el formato del CV y la descripción del trabajo'],
                            'summary': 'No se pudo completar el análisis automático. El CV y la descripción del trabajo han sido procesados, pero se recomienda una revisión manual.'
                        }
                    
                    # Validar y limpiar el resultado
                    try:
                        # Asegurar que match_percentage sea un entero válido
                        if 'match_percentage' not in comparison_result or not isinstance(comparison_result['match_percentage'], (int, float)):
                            comparison_result['match_percentage'] = 0
                        else:
                            comparison_result['match_percentage'] = max(0, min(100, int(comparison_result['match_percentage'])))
                        
                        # Asegurar que las listas sean válidas
                        for key in ['strengths', 'keywords_found', 'improvements']:
                            if key not in comparison_result or not isinstance(comparison_result[key], list):
                                comparison_result[key] = [f'No se pudieron identificar {key}']
                            elif not comparison_result[key]:  # Lista vacía
                                comparison_result[key] = [f'No se encontraron {key} específicos']
                        
                        # Asegurar que summary sea una cadena válida
                        if 'summary' not in comparison_result or not isinstance(comparison_result['summary'], str):
                            comparison_result['summary'] = 'Resumen no disponible'
                        
                        print(f"INFO: Resultado validado - Match: {comparison_result['match_percentage']}%")
                        
                    except Exception as e:
                        print(f"ERROR validando resultado: {e}")
                        add_console_log('ERROR', f'Error validando resultado para usuario {username}: {str(e)}', 'CV_COMPARE')
                        comparison_result = {
                            'match_percentage': 0,
                            'strengths': ['Error en el procesamiento'],
                            'keywords_found': ['Error en el procesamiento'],
                            'improvements': ['Intenta nuevamente con un formato de archivo diferente'],
                            'summary': 'Ocurrió un error durante la validación del análisis. Por favor, intenta nuevamente.'
                        }
                    
                    # Incrementar contador de uso
                    increment_usage(session.get('user_id'), 'cv_analysis')
                    
                    # Guardar resultado en base de datos
                    try:
                        connection = get_db_connection()
                        if connection:
                            cursor = connection.cursor()
                            
                            analysis_id = str(uuid.uuid4())
                            cursor.execute("""
                                INSERT INTO cv_analyses 
                                (id, user_id, filename, analysis_type, analysis_result, created_at)
                                VALUES (%s, %s, %s, %s, %s, %s)
                            """, (
                                analysis_id,
                                user_id,
                                filename,
                                'cv_job_comparison',
                                json.dumps(comparison_result),
                                datetime.now()
                            ))
                            
                            connection.commit()
                            cursor.close()
                            connection.close()
                            
                            add_console_log('INFO', f'Comparación CV-Oferta completada: {filename} por {username}', 'CV_COMPARE')
                            
                    except Exception as e:
                        add_console_log('ERROR', f'Error guardando comparación: {str(e)}', 'CV_COMPARE')
                        print(f"ERROR guardando en base de datos: {e}")
                        # Continuar sin fallar, ya que tenemos el resultado
                    
                    # Siempre retornar el resultado, incluso si hubo error en BD
                    return render_template('cv_job_comparison_result.html', 
                                         result=comparison_result,
                                         filename=filename,
                                         job_description=job_description[:200] + '...' if len(job_description) > 200 else job_description)
                else:
                    flash('No se pudo extraer texto del archivo CV', 'error')
        else:
            flash('Tipo de archivo no permitido. Solo se permiten archivos PDF, DOC y DOCX.', 'error')
    
    return render_template('compare_cv_job.html')

def compare_cv_with_job_ai(cv_text, job_description):
    """Comparar CV con oferta laboral usando IA de Gemini"""
    try:
        # Configurar Gemini
        import google.generativeai as genai
        
        gemini_api_key = os.getenv('GEMINI_API_KEY')
        if not gemini_api_key:
            print("ERROR: GEMINI_API_KEY no configurada")
            add_console_log('ERROR', 'GEMINI_API_KEY no configurada', 'CV_COMPARE')
            return None
            
        genai.configure(api_key=gemini_api_key)
        model = genai.GenerativeModel('gemini-1.5-flash')
        print(f"INFO: Configuración de Gemini exitosa, modelo: gemini-1.5-flash")
        
        # Prompt para la comparación
        prompt = f"""
Eres un experto en recursos humanos y análisis de currículums. Analiza el siguiente CV en relación con la oferta laboral proporcionada.

CV DEL CANDIDATO:
{cv_text}

OFERTA LABORAL:
{job_description}

Por favor, proporciona un análisis detallado en formato JSON con la siguiente estructura:

{{
    "match_percentage": [número del 0 al 100],
    "strengths": [
        "Fortaleza 1 específica encontrada",
        "Fortaleza 2 específica encontrada",
        "Fortaleza 3 específica encontrada"
    ],
    "keywords_found": [
        "palabra clave 1",
        "palabra clave 2",
        "palabra clave 3"
    ],
    "improvements": [
        "Mejora 1 específica recomendada",
        "Mejora 2 específica recomendada",
        "Mejora 3 específica recomendada"
    ],
    "summary": "Resumen ejecutivo de la compatibilidad entre el CV y la oferta laboral"
}}

Analiza específicamente:
1. Coincidencia de habilidades técnicas
2. Experiencia relevante
3. Nivel de educación requerido
4. Palabras clave importantes
5. Competencias blandas

Sé específico y constructivo en tus recomendaciones.
"""
        
        print(f"INFO: Enviando prompt a Gemini (longitud: {len(prompt)} caracteres)")
        response = model.generate_content(prompt)
        
        if response and response.text:
            print(f"INFO: Respuesta recibida de Gemini (longitud: {len(response.text)} caracteres)")
            # Intentar parsear la respuesta JSON
            try:
                # Limpiar la respuesta para extraer solo el JSON
                response_text = response.text.strip()
                
                # Buscar el JSON en la respuesta
                start_idx = response_text.find('{')
                end_idx = response_text.rfind('}') + 1
                
                if start_idx != -1 and end_idx != -1:
                    json_str = response_text[start_idx:end_idx]
                    print(f"INFO: JSON extraído (longitud: {len(json_str)} caracteres)")
                    result = json.loads(json_str)
                    
                    # Validar estructura del resultado
                    required_keys = ['match_percentage', 'strengths', 'keywords_found', 'improvements', 'summary']
                    if all(key in result for key in required_keys):
                        # Asegurar que match_percentage sea un número válido
                        try:
                            result['match_percentage'] = int(result['match_percentage'])
                            if result['match_percentage'] < 0 or result['match_percentage'] > 100:
                                result['match_percentage'] = 0
                        except (ValueError, TypeError):
                            result['match_percentage'] = 0
                        
                        # Asegurar que las listas sean listas válidas
                        for list_key in ['strengths', 'keywords_found', 'improvements']:
                            if not isinstance(result[list_key], list):
                                result[list_key] = []
                        
                        # Asegurar que summary sea una cadena
                        if not isinstance(result['summary'], str):
                            result['summary'] = 'Resumen no disponible'
                        
                        print(f"INFO: Análisis completado exitosamente")
                        return result
                    else:
                        print(f"ERROR: Estructura JSON incompleta: {result.keys()}")
                        add_console_log('ERROR', f'Estructura JSON incompleta: {result.keys()}', 'CV_COMPARE')
                        return None
                else:
                    print(f"ERROR: No se encontró JSON válido en la respuesta")
                    print(f"Respuesta completa: {response_text[:1000]}...")
                    add_console_log('ERROR', f'No se encontró JSON válido en respuesta de Gemini', 'CV_COMPARE')
                    return None
                    
            except json.JSONDecodeError as e:
                print(f"ERROR: Error parseando JSON de Gemini: {e}")
                print(f"Respuesta recibida: {response.text[:500]}...")
                add_console_log('ERROR', f'Error parseando JSON de Gemini: {str(e)}', 'CV_COMPARE')
                return None
        else:
            print("ERROR: No se recibió respuesta de Gemini")
            add_console_log('ERROR', 'No se recibió respuesta de Gemini', 'CV_COMPARE')
            return None
            
    except Exception as e:
        print(f"ERROR en compare_cv_with_job_ai: {e}")
        add_console_log('ERROR', f'Excepción en compare_cv_with_job_ai: {str(e)}', 'CV_COMPARE')
        import traceback
        print(f"Traceback completo: {traceback.format_exc()}")
        return None

# Manejadores de errores
@app.errorhandler(404)
def page_not_found(error):
    """Manejar errores 404"""
    return render_template('404.html'), 404

@app.errorhandler(500)
def internal_server_error(error):
    """Manejar errores 500"""
    return render_template('500.html'), 500

if __name__ == '__main__':
    init_database()
    app.run(debug=True)