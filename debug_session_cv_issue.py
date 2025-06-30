#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script de debug para investigar el problema de sesión con CVs
que se cargan exitosamente pero luego no pasan a la selección de IA
"""

import os
import sys
import tempfile
from werkzeug.utils import secure_filename
import PyPDF2
from docx import Document

def allowed_file(filename):
    """Verificar si el archivo tiene una extensión permitida"""
    ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_file_debug(filepath, file_extension=None):
    """Versión debug de extract_text_from_file para identificar problemas específicos"""
    print(f"\n=== INICIANDO EXTRACCIÓN DE TEXTO ===")
    print(f"Archivo: {filepath}")
    print(f"Extensión: {file_extension}")
    
    # Validaciones iniciales
    if not filepath or not os.path.exists(filepath):
        print(f"❌ ERROR: Archivo no existe: {filepath}")
        return None
    
    if not os.access(filepath, os.R_OK):
        print(f"❌ ERROR: No se puede leer el archivo: {filepath}")
        return None
    
    file_size = os.path.getsize(filepath)
    if file_size == 0:
        print(f"❌ ERROR: El archivo está vacío: {filepath}")
        return None
    
    if file_size > 50 * 1024 * 1024:  # 50MB límite
        print(f"❌ ERROR: Archivo demasiado grande: {file_size} bytes")
        return None
    
    print(f"✅ Archivo válido: {file_size} bytes")
    
    # Determinar extensión del archivo
    if file_extension is None:
        if '.' not in filepath:
            # Detectar tipo por header
            try:
                with open(filepath, 'rb') as f:
                    header = f.read(16)
                    print(f"Header del archivo: {header}")
                    
                    if header.startswith(b'%PDF'):
                        file_extension = 'pdf'
                        print("✅ Detectado como PDF por header")
                    elif header.startswith(b'PK\x03\x04'):
                        file_extension = 'docx'
                        print("✅ Detectado como DOCX por header")
                    elif header.startswith(b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'):
                        file_extension = 'doc'
                        print("✅ Detectado como DOC por header")
                    else:
                        print(f"❌ Header no reconocido: {header}")
                        return None
            except Exception as e:
                print(f"❌ Error leyendo header: {e}")
                return None
        else:
            file_extension = filepath.rsplit('.', 1)[1].lower()
    
    print(f"Extensión final: {file_extension}")
    
    text = ""
    
    try:
        if file_extension == 'pdf':
            print("📄 Procesando PDF...")
            text = _extract_from_pdf_debug(filepath)
        elif file_extension in ['doc', 'docx']:
            print(f"📝 Procesando {file_extension.upper()}...")
            text = _extract_from_word_debug(filepath, file_extension)
        else:
            print(f"❌ Extensión no soportada: {file_extension}")
            return None
    
    except Exception as e:
        print(f"❌ Error general: {e}")
        import traceback
        print(f"Traceback: {traceback.format_exc()}")
        return None
    
    if not text or len(text.strip()) == 0:
        print(f"❌ ERROR CRÍTICO: No se extrajo texto del archivo")
        print(f"Texto retornado: {repr(text)}")
        return None
    
    text = text.strip()
    print(f"\n=== RESULTADO FINAL ===")
    print(f"✅ Texto extraído: {len(text)} caracteres")
    print(f"Primeros 100 caracteres: {repr(text[:100])}")
    
    # Verificaciones adicionales que podrían causar problemas
    if len(text) < 50:
        print(f"⚠️  ADVERTENCIA: Texto muy corto (menos de 50 caracteres)")
        print(f"Esto podría causar problemas en el análisis")
    
    # Verificar caracteres problemáticos
    non_printable = sum(1 for c in text if not c.isprintable() and not c.isspace())
    if non_printable > len(text) * 0.1:
        print(f"⚠️  ADVERTENCIA: {non_printable} caracteres no imprimibles ({non_printable/len(text)*100:.1f}%)")
    
    # Verificar si el texto parece ser un CV
    cv_keywords = ['experiencia', 'educación', 'habilidades', 'trabajo', 'universidad', 'empresa', 'proyecto']
    found_keywords = sum(1 for keyword in cv_keywords if keyword.lower() in text.lower())
    
    if found_keywords < 2:
        print(f"⚠️  ADVERTENCIA: El texto no parece ser un CV (solo {found_keywords} palabras clave encontradas)")
        print(f"Palabras buscadas: {cv_keywords}")
    else:
        print(f"✅ El texto parece ser un CV ({found_keywords} palabras clave encontradas)")
    
    return text

def _extract_from_pdf_debug(filepath):
    """Extraer texto de PDF con debug detallado"""
    text = ""
    
    try:
        with open(filepath, 'rb') as file:
            # Verificar header
            file.seek(0)
            header = file.read(4)
            if not header.startswith(b'%PDF'):
                raise ValueError(f"No es un PDF válido. Header: {header}")
            
            file.seek(0)
            pdf_reader = PyPDF2.PdfReader(file)
            
            # Verificar encriptación
            if pdf_reader.is_encrypted:
                print("⚠️  PDF encriptado, intentando desencriptar...")
                try:
                    pdf_reader.decrypt('')
                    print("✅ PDF desencriptado")
                except:
                    print("❌ No se puede desencriptar el PDF")
                    return None
            
            num_pages = len(pdf_reader.pages)
            print(f"PDF tiene {num_pages} páginas")
            
            if num_pages == 0:
                print("❌ PDF sin páginas")
                return None
            
            # Extraer texto página por página
            pages_with_text = 0
            total_chars = 0
            
            for i, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text += page_text + "\n"
                        pages_with_text += 1
                        total_chars += len(page_text)
                        print(f"  Página {i+1}: {len(page_text)} caracteres")
                    else:
                        print(f"  Página {i+1}: Sin texto extraíble")
                except Exception as page_error:
                    print(f"  ❌ Error en página {i+1}: {page_error}")
                    continue
            
            print(f"Texto extraído de {pages_with_text}/{num_pages} páginas")
            print(f"Total de caracteres: {total_chars}")
            
            if pages_with_text == 0:
                print(f"❌ PROBLEMA: Ninguna página tiene texto extraíble")
                print(f"Esto podría ser un PDF escaneado que requiere OCR")
    
    except Exception as e:
        print(f"❌ Error procesando PDF: {e}")
        raise
    
    return text

def _extract_from_word_debug(filepath, file_extension):
    """Extraer texto de Word con debug detallado"""
    text = ""
    
    try:
        # Verificar header
        with open(filepath, 'rb') as f:
            header = f.read(8)
            if file_extension == 'docx' and not header.startswith(b'PK'):
                print(f"⚠️  Archivo no parece DOCX válido. Header: {header}")
            elif file_extension == 'doc' and not header.startswith(b'\xd0\xcf'):
                print(f"⚠️  Archivo no parece DOC válido. Header: {header}")
        
        # Abrir documento
        doc = Document(filepath)
        
        # Extraer párrafos
        paragraph_count = 0
        empty_paragraphs = 0
        
        for i, paragraph in enumerate(doc.paragraphs):
            para_text = paragraph.text.strip()
            if para_text:
                text += para_text + '\n'
                paragraph_count += 1
            else:
                empty_paragraphs += 1
        
        print(f"Párrafos con contenido: {paragraph_count}")
        print(f"Párrafos vacíos: {empty_paragraphs}")
        
        # Extraer tablas
        table_count = 0
        table_cells = 0
        
        if hasattr(doc, 'tables') and doc.tables:
            for table in doc.tables:
                table_count += 1
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            text += cell_text + ' '
                            table_cells += 1
                text += '\n'
        
        print(f"Tablas procesadas: {table_count}")
        print(f"Celdas con contenido: {table_cells}")
        
        if paragraph_count == 0 and table_cells == 0:
            print(f"❌ PROBLEMA: No se encontró contenido en párrafos ni tablas")
    
    except Exception as e:
        print(f"❌ Error procesando Word: {e}")
        raise
    
    return text

def simulate_cv_upload_process(test_file_path):
    """Simular el proceso completo de carga de CV"""
    print(f"\n{'='*80}")
    print(f"SIMULANDO PROCESO DE CARGA DE CV")
    print(f"{'='*80}")
    
    if not os.path.exists(test_file_path):
        print(f"❌ Archivo no encontrado: {test_file_path}")
        return False
    
    filename = os.path.basename(test_file_path)
    print(f"Archivo: {filename}")
    
    # Paso 1: Verificar si está permitido
    if not allowed_file(filename):
        print(f"❌ Archivo no permitido según allowed_file()")
        return False
    
    print(f"✅ Archivo permitido")
    
    # Paso 2: Simular guardado temporal
    try:
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            temp_path = temp_file.name
            
            # Copiar contenido del archivo original al temporal
            with open(test_file_path, 'rb') as original:
                temp_file.write(original.read())
        
        print(f"✅ Archivo guardado temporalmente: {temp_path}")
        
        # Paso 3: Extraer extensión
        original_extension = filename.rsplit('.', 1)[1].lower() if '.' in filename else None
        print(f"Extensión original: {original_extension}")
        
        # Paso 4: Extraer texto
        text_content = extract_text_from_file_debug(temp_path, original_extension)
        
        # Paso 5: Limpiar archivo temporal
        try:
            os.unlink(temp_path)
            print(f"✅ Archivo temporal eliminado")
        except:
            print(f"⚠️  No se pudo eliminar archivo temporal: {temp_path}")
        
        # Paso 6: Evaluar resultado
        if text_content:
            print(f"\n✅ PROCESO EXITOSO")
            print(f"El CV debería pasar a la selección de IA")
            print(f"Contenido de sesión simulado:")
            print(f"  - cv_content: {len(text_content)} caracteres")
            print(f"  - cv_filename: {filename}")
            return True
        else:
            print(f"\n❌ PROCESO FALLIDO")
            print(f"El CV NO pasará a la selección de IA")
            print(f"Razón: No se extrajo texto del archivo")
            return False
    
    except Exception as e:
        print(f"❌ Error durante el proceso: {e}")
        import traceback
        print(f"Traceback: {traceback.format_exc()}")
        return False

def main():
    """Función principal"""
    print("=== DEBUG DE PROBLEMA DE SESIÓN CON CVs ===")
    print("Este script simula el proceso completo de carga de CV")
    print("para identificar por qué algunos CVs no pasan a la selección de IA")
    
    # Verificar dependencias
    try:
        import PyPDF2
        print(f"✅ PyPDF2 versión: {PyPDF2.__version__}")
    except ImportError as e:
        print(f"❌ Error importando PyPDF2: {e}")
        return
    
    try:
        from docx import Document
        print(f"✅ python-docx disponible")
    except ImportError as e:
        print(f"❌ Error importando python-docx: {e}")
        return
    
    # Si se proporciona un archivo como argumento, probarlo
    if len(sys.argv) > 1:
        test_file_path = sys.argv[1]
        simulate_cv_upload_process(test_file_path)
    else:
        print("\nUso: python debug_session_cv_issue.py <ruta_del_cv>")
        print("Ejemplo: python debug_session_cv_issue.py mi_cv.pdf")
        
        # Buscar archivos de prueba en el directorio actual
        current_dir = os.getcwd()
        test_files = []
        
        for file in os.listdir(current_dir):
            if file.lower().endswith(('.pdf', '.doc', '.docx')):
                test_files.append(file)
        
        if test_files:
            print(f"\nArchivos encontrados en el directorio actual:")
            for i, file in enumerate(test_files, 1):
                print(f"{i}. {file}")
            
            try:
                choice = input("\nSelecciona un archivo para probar (número) o presiona Enter para salir: ")
                if choice.strip():
                    file_index = int(choice) - 1
                    if 0 <= file_index < len(test_files):
                        test_file_path = os.path.join(current_dir, test_files[file_index])
                        simulate_cv_upload_process(test_file_path)
            except (ValueError, IndexError, KeyboardInterrupt):
                print("Saliendo...")
        else:
            print("\nNo se encontraron archivos PDF, DOC o DOCX en el directorio actual")
    
    print("\n=== POSIBLES CAUSAS DEL PROBLEMA ===")
    print("1. Archivos PDF escaneados que no tienen texto extraíble")
    print("2. Archivos Word corruptos o con formato inválido")
    print("3. Archivos con contenido muy corto (menos de 50 caracteres)")
    print("4. Archivos con muchos caracteres no imprimibles")
    print("5. Problemas de permisos o acceso al archivo")
    print("6. Archivos que no contienen contenido típico de CV")
    
    print("\n=== RECOMENDACIONES ===")
    print("1. Implementar validación de contenido mínimo")
    print("2. Agregar OCR para PDFs escaneados")
    print("3. Mejorar mensajes de error específicos")
    print("4. Implementar logging detallado en producción")
    print("5. Agregar validación de palabras clave de CV")

if __name__ == "__main__":
    main()