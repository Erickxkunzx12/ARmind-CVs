#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script de debug para identificar problemas en la extracción de texto de archivos PDF y Word
"""

import os
import sys
import tempfile
from werkzeug.utils import secure_filename
import PyPDF2
from docx import Document

def allowed_file(filename):
    """Verificar si el archivo tiene una extensión permitida"""
    ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_file_debug(filepath, file_extension=None):
    """Versión debug de extract_text_from_file con logging detallado"""
    text = ""
    
    print(f"\n=== INICIANDO EXTRACCIÓN ===")
    print(f"Archivo: {filepath}")
    print(f"Extensión proporcionada: {file_extension}")
    
    # Verificar que el archivo existe
    if not os.path.exists(filepath):
        print(f"❌ ERROR: El archivo no existe: {filepath}")
        return None
    
    # Verificar permisos de lectura
    if not os.access(filepath, os.R_OK):
        print(f"❌ ERROR: No se puede leer el archivo: {filepath}")
        return None
    
    # Verificar tamaño del archivo
    file_size = os.path.getsize(filepath)
    print(f"Tamaño del archivo: {file_size} bytes")
    
    if file_size == 0:
        print(f"❌ ERROR: El archivo está vacío")
        return None
    
    # Si no se proporciona extensión, intentar obtenerla del filepath
    if file_extension is None:
        if '.' not in filepath:
            # Si no hay extensión, intentar detectar el tipo de archivo
            try:
                with open(filepath, 'rb') as f:
                    header = f.read(16)  # Leer más bytes para mejor detección
                    print(f"Header del archivo (16 bytes): {header}")
                    
                    if header.startswith(b'%PDF'):
                        file_extension = 'pdf'
                        print("✅ Detectado como PDF por header")
                    elif header.startswith(b'PK\x03\x04'):
                        file_extension = 'docx'
                        print("✅ Detectado como DOCX por header")
                    elif header.startswith(b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'):
                        file_extension = 'doc'
                        print("✅ Detectado como DOC por header")
                    else:
                        print(f"❌ Header no reconocido: {header}")
                        print("Intentando con extensión del nombre de archivo...")
                        return None
            except Exception as e:
                print(f"❌ Error leyendo header: {e}")
                return None
        else:
            file_extension = filepath.rsplit('.', 1)[1].lower()
            print(f"Extensión extraída del nombre: {file_extension}")
    
    print(f"Extensión final a usar: {file_extension}")
    
    try:
        if file_extension == 'pdf':
            print("📄 Procesando como PDF...")
            
            try:
                with open(filepath, 'rb') as file:
                    # Verificar que el archivo se puede abrir
                    file.seek(0, 2)  # Ir al final
                    file_length = file.tell()
                    file.seek(0)  # Volver al inicio
                    print(f"Longitud del archivo PDF: {file_length} bytes")
                    
                    # Intentar crear el lector PDF
                    try:
                        pdf_reader = PyPDF2.PdfReader(file)
                        print(f"✅ PDF Reader creado exitosamente")
                        
                        # Verificar si el PDF está encriptado
                        if pdf_reader.is_encrypted:
                            print(f"⚠️  ADVERTENCIA: El PDF está encriptado")
                            try:
                                pdf_reader.decrypt('')  # Intentar con contraseña vacía
                                print(f"✅ PDF desencriptado con contraseña vacía")
                            except:
                                print(f"❌ ERROR: No se puede desencriptar el PDF")
                                return None
                        
                        num_pages = len(pdf_reader.pages)
                        print(f"Número de páginas: {num_pages}")
                        
                        if num_pages == 0:
                            print(f"❌ ERROR: El PDF no tiene páginas")
                            return None
                        
                        total_chars = 0
                        for i, page in enumerate(pdf_reader.pages):
                            print(f"Procesando página {i+1}/{num_pages}...")
                            try:
                                page_text = page.extract_text()
                                page_chars = len(page_text) if page_text else 0
                                total_chars += page_chars
                                print(f"  Página {i+1}: {page_chars} caracteres")
                                
                                if page_text:
                                    text += page_text + "\n"
                                else:
                                    print(f"  ⚠️  Página {i+1} no tiene texto extraíble")
                                    
                            except Exception as page_error:
                                print(f"  ❌ Error en página {i+1}: {page_error}")
                                continue
                        
                        print(f"Total de caracteres extraídos: {total_chars}")
                        
                    except Exception as reader_error:
                        print(f"❌ Error creando PDF Reader: {reader_error}")
                        print(f"Tipo de error: {type(reader_error).__name__}")
                        
                        # Intentar diagnóstico adicional
                        file.seek(0)
                        first_line = file.readline()
                        print(f"Primera línea del archivo: {first_line}")
                        
                        return None
                        
            except Exception as file_error:
                print(f"❌ Error abriendo archivo PDF: {file_error}")
                return None
        
        elif file_extension in ['doc', 'docx']:
            print(f"📝 Procesando como {file_extension.upper()}...")
            
            try:
                # Verificar si es realmente un archivo de Word
                with open(filepath, 'rb') as f:
                    header = f.read(8)
                    if file_extension == 'docx' and not header.startswith(b'PK'):
                        print(f"⚠️  ADVERTENCIA: El archivo no parece ser un DOCX válido")
                    elif file_extension == 'doc' and not header.startswith(b'\xd0\xcf'):
                        print(f"⚠️  ADVERTENCIA: El archivo no parece ser un DOC válido")
                
                doc = Document(filepath)
                print(f"✅ Documento Word abierto exitosamente")
                
                num_paragraphs = len(doc.paragraphs)
                print(f"Número de párrafos: {num_paragraphs}")
                
                if num_paragraphs == 0:
                    print(f"❌ ERROR: El documento no tiene párrafos")
                    return None
                
                total_chars = 0
                non_empty_paragraphs = 0
                
                for i, paragraph in enumerate(doc.paragraphs):
                    para_text = paragraph.text
                    para_chars = len(para_text)
                    total_chars += para_chars
                    
                    if para_text.strip():  # Solo contar párrafos con contenido
                        non_empty_paragraphs += 1
                        text += para_text + '\n'
                
                print(f"Párrafos con contenido: {non_empty_paragraphs}/{num_paragraphs}")
                print(f"Total de caracteres extraídos: {total_chars}")
                
                # Intentar extraer texto de tablas también
                if hasattr(doc, 'tables') and doc.tables:
                    print(f"Encontradas {len(doc.tables)} tablas")
                    for i, table in enumerate(doc.tables):
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    text += cell.text + ' '
                        text += '\n'
                    
            except Exception as docx_error:
                print(f"❌ Error procesando documento Word: {docx_error}")
                print(f"Tipo de error: {type(docx_error).__name__}")
                return None
        else:
            print(f"❌ Extensión no soportada: {file_extension}")
            return None
    
    except Exception as e:
        print(f"❌ Error general al extraer texto: {e}")
        print(f"Tipo de error: {type(e).__name__}")
        import traceback
        print(f"Traceback: {traceback.format_exc()}")
        return None
    
    final_text = text.strip()
    print(f"\n=== RESULTADO FINAL ===")
    print(f"Texto extraído: {len(final_text)} caracteres")
    
    if len(final_text) > 0:
        print(f"✅ ÉXITO: Se extrajo texto del archivo")
        print(f"Primeros 100 caracteres: {repr(final_text[:100])}")
        if len(final_text) > 100:
            print(f"Últimos 100 caracteres: {repr(final_text[-100:])}")
        
        # Verificar calidad del texto
        if len(final_text.strip()) < 50:
            print(f"⚠️  ADVERTENCIA: El texto extraído es muy corto (menos de 50 caracteres)")
        
        # Verificar si hay caracteres extraños
        printable_chars = sum(1 for c in final_text if c.isprintable() or c.isspace())
        if printable_chars < len(final_text) * 0.8:
            print(f"⚠️  ADVERTENCIA: El texto contiene muchos caracteres no imprimibles")
        
        return final_text
    else:
        print(f"❌ ERROR: No se extrajo ningún texto")
        return None

def simulate_upload_process():
    """Simular el proceso completo de carga como en la aplicación"""
    print("\n" + "="*60)
    print("SIMULANDO PROCESO COMPLETO DE CARGA")
    print("="*60)
    
    # Crear un archivo de prueba simple
    test_content = "Este es un CV de prueba.\n\nNombre: Juan Pérez\nExperiencia: 5 años en desarrollo\nEducación: Ingeniería en Sistemas"
    
    # Crear archivo temporal
    with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as temp_file:
        temp_file.write(test_content)
        temp_path = temp_file.name
    
    print(f"Archivo temporal creado: {temp_path}")
    
    try:
        # Simular el proceso de la aplicación
        filename = "test_cv.txt"
        print(f"Nombre de archivo simulado: {filename}")
        
        # Verificar si está permitido
        if not allowed_file(filename):
            print(f"❌ Archivo no permitido según allowed_file()")
            return
        
        print(f"✅ Archivo permitido")
        
        # Extraer extensión
        original_extension = filename.rsplit('.', 1)[1].lower() if '.' in filename else None
        print(f"Extensión original: {original_extension}")
        
        # Intentar extraer texto
        text_content = extract_text_from_file_debug(temp_path, original_extension)
        
        if text_content:
            print(f"\n✅ PROCESO EXITOSO")
            print(f"Texto extraído: {len(text_content)} caracteres")
        else:
            print(f"\n❌ PROCESO FALLIDO")
    
    finally:
        # Limpiar archivo temporal
        try:
            os.unlink(temp_path)
            print(f"Archivo temporal eliminado")
        except:
            print(f"No se pudo eliminar archivo temporal: {temp_path}")

def main():
    """Función principal"""
    print("=== SCRIPT DE DEBUG PARA EXTRACCIÓN DE ARCHIVOS ===")
    print("Este script identificará problemas en la extracción de texto")
    
    # Verificar dependencias
    try:
        import PyPDF2
        print(f"✅ PyPDF2 versión: {PyPDF2.__version__}")
    except ImportError as e:
        print(f"❌ Error importando PyPDF2: {e}")
        return
    
    try:
        from docx import Document
        print(f"✅ python-docx disponible")
    except ImportError as e:
        print(f"❌ Error importando python-docx: {e}")
        return
    
    # Simular proceso de carga
    simulate_upload_process()
    
    # Si se proporciona un archivo como argumento, probarlo
    if len(sys.argv) > 1:
        test_file_path = sys.argv[1]
        if os.path.exists(test_file_path):
            filename = os.path.basename(test_file_path)
            original_extension = filename.rsplit('.', 1)[1].lower() if '.' in filename else None
            extract_text_from_file_debug(test_file_path, original_extension)
        else:
            print(f"❌ Archivo no encontrado: {test_file_path}")
    
    print("\n=== RECOMENDACIONES ===")
    print("1. Verificar que los archivos PDF no estén corruptos o encriptados")
    print("2. Asegurarse de que los archivos Word no estén dañados")
    print("3. Verificar que los archivos tengan contenido de texto extraíble")
    print("4. Considerar agregar más validaciones en la aplicación")
    print("5. Implementar logging detallado para identificar problemas específicos")

if __name__ == "__main__":
    main()