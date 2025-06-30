#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Función mejorada para extraer texto de archivos PDF y Word
con manejo robusto de errores y validaciones
"""

import os
import PyPDF2
from docx import Document
import logging

# Configurar logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def extract_text_from_file_improved(filepath, file_extension=None):
    """
    Versión mejorada de extract_text_from_file con manejo robusto de errores
    
    Args:
        filepath (str): Ruta al archivo
        file_extension (str): Extensión del archivo (opcional)
    
    Returns:
        str: Texto extraído del archivo, o None si hay error
    """
    
    # Validaciones iniciales
    if not filepath or not os.path.exists(filepath):
        logger.error(f"Archivo no existe: {filepath}")
        return None
    
    if not os.access(filepath, os.R_OK):
        logger.error(f"No se puede leer el archivo: {filepath}")
        return None
    
    file_size = os.path.getsize(filepath)
    if file_size == 0:
        logger.error(f"El archivo está vacío: {filepath}")
        return None
    
    if file_size > 50 * 1024 * 1024:  # 50MB límite
        logger.error(f"Archivo demasiado grande: {file_size} bytes")
        return None
    
    logger.info(f"Procesando archivo: {filepath} ({file_size} bytes)")
    
    # Determinar extensión del archivo
    if file_extension is None:
        if '.' not in filepath:
            # Detectar tipo por header
            try:
                with open(filepath, 'rb') as f:
                    header = f.read(16)
                    if header.startswith(b'%PDF'):
                        file_extension = 'pdf'
                        logger.info("Detectado como PDF por header")
                    elif header.startswith(b'PK\x03\x04'):
                        file_extension = 'docx'
                        logger.info("Detectado como DOCX por header")
                    elif header.startswith(b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'):
                        file_extension = 'doc'
                        logger.info("Detectado como DOC por header")
                    else:
                        logger.error(f"Tipo de archivo no reconocido: {header[:8]}")
                        return None
            except Exception as e:
                logger.error(f"Error leyendo header del archivo: {e}")
                return None
        else:
            file_extension = filepath.rsplit('.', 1)[1].lower()
    
    text = ""
    
    try:
        if file_extension == 'pdf':
            text = _extract_from_pdf(filepath)
        elif file_extension in ['doc', 'docx']:
            text = _extract_from_word(filepath, file_extension)
        else:
            logger.error(f"Extensión no soportada: {file_extension}")
            return None
    
    except Exception as e:
        logger.error(f"Error general extrayendo texto: {e}")
        return None
    
    if not text or len(text.strip()) == 0:
        logger.warning("No se extrajo texto del archivo")
        return None
    
    # Validar calidad del texto
    text = text.strip()
    if len(text) < 10:
        logger.warning(f"Texto extraído muy corto: {len(text)} caracteres")
    
    # Verificar caracteres imprimibles
    printable_ratio = sum(1 for c in text if c.isprintable() or c.isspace()) / len(text)
    if printable_ratio < 0.8:
        logger.warning(f"Texto contiene muchos caracteres no imprimibles: {printable_ratio:.2%}")
    
    logger.info(f"Texto extraído exitosamente: {len(text)} caracteres")
    return text

def _extract_from_pdf(filepath):
    """
    Extraer texto de archivo PDF con manejo robusto de errores
    """
    text = ""
    
    try:
        with open(filepath, 'rb') as file:
            # Verificar que es un PDF válido
            file.seek(0)
            header = file.read(4)
            if not header.startswith(b'%PDF'):
                raise ValueError("No es un archivo PDF válido")
            
            file.seek(0)
            pdf_reader = PyPDF2.PdfReader(file)
            
            # Verificar si está encriptado
            if pdf_reader.is_encrypted:
                logger.warning("PDF está encriptado, intentando desencriptar")
                try:
                    pdf_reader.decrypt('')  # Intentar con contraseña vacía
                except:
                    logger.error("No se puede desencriptar el PDF")
                    return None
            
            num_pages = len(pdf_reader.pages)
            if num_pages == 0:
                logger.error("PDF no tiene páginas")
                return None
            
            logger.info(f"PDF tiene {num_pages} páginas")
            
            # Extraer texto de cada página
            for i, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                        logger.debug(f"Página {i+1}: {len(page_text)} caracteres")
                    else:
                        logger.warning(f"Página {i+1} no tiene texto extraíble")
                except Exception as page_error:
                    logger.warning(f"Error en página {i+1}: {page_error}")
                    continue
    
    except Exception as e:
        logger.error(f"Error procesando PDF: {e}")
        raise
    
    return text

def _extract_from_word(filepath, file_extension):
    """
    Extraer texto de archivo Word con manejo robusto de errores
    """
    text = ""
    
    try:
        # Verificar header del archivo
        with open(filepath, 'rb') as f:
            header = f.read(8)
            if file_extension == 'docx' and not header.startswith(b'PK'):
                logger.warning("El archivo no parece ser un DOCX válido")
            elif file_extension == 'doc' and not header.startswith(b'\xd0\xcf'):
                logger.warning("El archivo no parece ser un DOC válido")
        
        # Abrir documento
        doc = Document(filepath)
        
        # Extraer texto de párrafos
        paragraph_count = 0
        for paragraph in doc.paragraphs:
            para_text = paragraph.text.strip()
            if para_text:
                text += para_text + '\n'
                paragraph_count += 1
        
        logger.info(f"Extraídos {paragraph_count} párrafos")
        
        # Extraer texto de tablas
        table_count = 0
        if hasattr(doc, 'tables') and doc.tables:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            text += cell_text + ' '
                text += '\n'
                table_count += 1
            
            logger.info(f"Extraídas {table_count} tablas")
    
    except Exception as e:
        logger.error(f"Error procesando documento Word: {e}")
        raise
    
    return text

def test_improved_extraction():
    """
    Función de prueba para la extracción mejorada
    """
    import tempfile
    
    # Crear archivo de prueba
    test_content = """CURRÍCULUM VITAE
    
Nombre: Juan Pérez
Teléfono: +1234567890
Email: juan.perez@email.com

EXPERIENCIA LABORAL:
- Desarrollador Senior (2020-2023)
- Analista de Sistemas (2018-2020)

EDUCACIÓN:
- Ingeniería en Sistemas (2014-2018)

HABILIDADES:
- Python, JavaScript, SQL
- Gestión de proyectos
- Trabajo en equipo"""
    
    with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as temp_file:
        temp_file.write(test_content)
        temp_path = temp_file.name
    
    try:
        print("=== PRUEBA DE EXTRACCIÓN MEJORADA ===")
        result = extract_text_from_file_improved(temp_path, 'txt')
        
        if result:
            print(f"✅ Éxito: {len(result)} caracteres extraídos")
            print(f"Primeros 100 caracteres: {result[:100]}")
        else:
            print("❌ Error: No se extrajo texto")
    
    finally:
        try:
            os.unlink(temp_path)
        except:
            pass

if __name__ == "__main__":
    test_improved_extraction()